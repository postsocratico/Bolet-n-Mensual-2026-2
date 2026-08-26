
import streamlit as st
import requests
import pandas as pd
import html
from io import BytesIO
import datetime
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from bs4 import BeautifulSoup
import calendar
import time
import re
from dateutil import parser
import urllib.parse # Necesario para la herramienta de rescate
import cloudscraper  # Para bypass de Cloudflare en BID
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# ==========================================
# MAPEO DE FUENTES ORIGINALES PARA AUDITORÍA (CON ENLACES)
# ==========================================

FUENTES_FMI = {
    # === PUBLICACIONES INSTITUCIONALES (8 ENLACES) ===
    "F&D Magazine": {
        "nombre": "F&D Magazine",
        "origen": "JSON Next.js",
        "url": "https://www.imf.org/en/publications/fandd/issues",
        "enlace_id": 1
    },
    "Fiscal Monitor": {
        "nombre": "Fiscal Monitor",
        "origen": "JSON Next.js",
        "url": "https://www.imf.org/en/publications/fm",
        "enlace_id": 2
    },
    "Global Financial Stability Report": {
        "nombre": "Global Financial Stability Report",
        "origen": "JSON Next.js",
        "url": "https://www.imf.org/en/publications/gfsr",
        "enlace_id": 3
    },
    "IMF Annual Report": {
        "nombre": "IMF Annual Report",
        "origen": "JSON Next.js",
        "url": "https://www.imf.org/en/publications/areb",
        "enlace_id": 4
    },
    "Regional Economic Outlook": {
        "nombre": "Regional Economic Outlook",
        "origen": "JSON Next.js",
        "url": "https://www.imf.org/en/publications/reo",
        "enlace_id": 5
    },
    "World Economic Outlook": {
        "nombre": "World Economic Outlook",
        "origen": "JSON Next.js",
        "url": "https://www.imf.org/en/publications/weo",
        "enlace_id": 6
    },
    "Press Releases": {
        "nombre": "Press Releases",
        "origen": "API Coveo",
        "url": "https://www.imf.org/en/news/searchnews#sortCriteria=%40imfdate%20descending&cf-type=PRESSRES&df-date=past-3-month..now",
        "enlace_id": 7
    },
    "Country Reports (Article IV)": {
        "nombre": "Country Reports (Article IV)",
        "origen": "API Coveo",
        "url": "https://www.imf.org/en/search#sortCriteria=%40imfdate%20descending&cf-type=PUBS,COUNTRYREPS,ARTICLE4",
        "enlace_id": 8
    },
    
    # === INVESTIGACIÓN (NO son Publicaciones Institucionales) ===
    "Working Papers": {
        "nombre": "Working Papers",
        "origen": "Crossref API",
        "url": "https://www.imf.org/en/Publications/SPROLLs/working-papers",
        "enlace_id": "Inv-1"
    },
    "Blogs": {
        "nombre": "Blogs",
        "origen": "API Coveo",
        "url": "https://www.imf.org/en/Blogs",
        "enlace_id": "Inv-2"
    },
    
    # === DISCURSOS (NO son Publicaciones Institucionales) ===
    "Speeches": {
        "nombre": "Speeches",
        "origen": "API Coveo",
        "url": "https://www.imf.org/en/News/Speeches",
        "enlace_id": "Disc-1"
    },
}

# ==========================================
# ENLACES FIJOS PARA PUBLICACIONES INSTITUCIONALES DEL FMI
# ==========================================

ENLACES_FMI_PUB_INST = [
    {"id": 1, "nombre": "F&D Magazine", "origen": "JSON Next.js", "url": "https://www.imf.org/en/publications/fandd/issues"},
    {"id": 2, "nombre": "Fiscal Monitor", "origen": "JSON Next.js", "url": "https://www.imf.org/en/publications/fm"},
    {"id": 3, "nombre": "Global Financial Stability Report", "origen": "JSON Next.js", "url": "https://www.imf.org/en/publications/gfsr"},
    {"id": 4, "nombre": "IMF Annual Report", "origen": "JSON Next.js", "url": "https://www.imf.org/en/publications/areb"},
    {"id": 5, "nombre": "Regional Economic Outlook", "origen": "JSON Next.js", "url": "https://www.imf.org/en/publications/reo"},
    {"id": 6, "nombre": "World Economic Outlook", "origen": "JSON Next.js", "url": "https://www.imf.org/en/publications/weo"},
    {"id": 7, "nombre": "Press Releases", "origen": "API Coveo", "url": "https://www.imf.org/en/news/searchnews#sortCriteria=%40imfdate%20descending&cf-type=PRESSRES&df-date=past-3-month..now"},
    {"id": 8, "nombre": "Country Reports (Article IV) y Mission Concluding", "origen": "API Coveo", "url": "https://www.imf.org/en/search#sortCriteria=%40imfdate%20descending&cf-type=PUBS,COUNTRYREPS,ARTICLE4"},
]

# ==========================================
# FUNCIÓN PARA IDENTIFICAR FUENTE FMI
# ==========================================
def identificar_fuente_fmi(row):
    """
    Identifica la fuente original de un documento del FMI
    basado en su título y categoría.
    """
    titulo = row.get('Title', '')
    categoria = row.get('Categoría', '')
    organismo = row.get('Organismo', '')
    
    # Solo procesar FMI
    if organismo != 'FMI':
        return {
            "nombre": "No aplica",
            "origen": "No aplica",
            "url": "",
            "enlace_id": "N/A"
        }
    
    # === PUBLICACIONES INSTITUCIONALES ===
    if categoria == 'Publicaciones Institucionales':
        # Enlace 1: F&D Magazine
        if 'F&D' in titulo or 'Finance & Development' in titulo:
            return {
                "nombre": "F&D Magazine",
                "origen": "JSON Next.js",
                "url": "https://www.imf.org/en/publications/fandd/issues",
                "enlace_id": "1"
            }
        # Enlace 2: Fiscal Monitor
        elif 'Fiscal Monitor' in titulo:
            return {
                "nombre": "Fiscal Monitor",
                "origen": "JSON Next.js",
                "url": "https://www.imf.org/en/publications/fm",
                "enlace_id": "2"
            }
        # Enlace 3: Global Financial Stability Report
        elif 'Global Financial Stability' in titulo or 'GFSR' in titulo:
            return {
                "nombre": "Global Financial Stability Report",
                "origen": "JSON Next.js",
                "url": "https://www.imf.org/en/publications/gfsr",
                "enlace_id": "3"
            }
        # Enlace 4: IMF Annual Report
        elif 'Annual Report' in titulo and 'IMF' in titulo:
            return {
                "nombre": "IMF Annual Report",
                "origen": "JSON Next.js",
                "url": "https://www.imf.org/en/publications/areb",
                "enlace_id": "4"
            }
        # Enlace 5: Regional Economic Outlook
        elif 'Regional Economic Outlook' in titulo or 'REO' in titulo:
            return {
                "nombre": "Regional Economic Outlook",
                "origen": "JSON Next.js",
                "url": "https://www.imf.org/en/publications/reo",
                "enlace_id": "5"
            }
        # Enlace 6: World Economic Outlook
        elif 'World Economic Outlook' in titulo or 'WEO' in titulo:
            return {
                "nombre": "World Economic Outlook",
                "origen": "JSON Next.js",
                "url": "https://www.imf.org/en/publications/weo",
                "enlace_id": "6"
            }
        # Enlace 7: Press Releases
        elif 'Press Release' in titulo or 'IMF Staff Completes' in titulo or 'IMF Reaches' in titulo:
            return {
                "nombre": "Press Releases",
                "origen": "API Coveo",
                "url": "https://www.imf.org/en/news/searchnews#sortCriteria=%40imfdate%20descending&cf-type=PRESSRES&df-date=past-3-month..now",
                "enlace_id": "7"
            }
        # Enlace 8: Country Reports (Article IV) y Mission Concluding
        elif 'Article IV' in titulo or 'Staff Report' in titulo:
            return {
                "nombre": "Country Reports (Article IV)",
                "origen": "API Coveo",
                "url": "https://www.imf.org/en/search#sortCriteria=%40imfdate%20descending&cf-type=PUBS,COUNTRYREPS,ARTICLE4",
                "enlace_id": "8"
            }
        elif 'Concluding Statement' in titulo or 'Mission Concluding' in titulo:
            return {
                "nombre": "Mission Concluding",
                "origen": "API Coveo",
                "url": "https://www.imf.org/en/search#sortCriteria=%40imfdate%20descending&cf-type=PUBS,COUNTRYREPS,ARTICLE4",
                "enlace_id": "8"  # Mismo ID que Country Reports
            }
        # Si no se identifica
        else:
            return {
                "nombre": "No identificado",
                "origen": "Desconocido",
                "url": "",
                "enlace_id": "No-ID"
            }
    
    # === INVESTIGACIÓN ===
    elif categoria == 'Investigación':
        if 'Working Paper' in titulo or 'WP/' in titulo:
            return {
                "nombre": "Working Papers",
                "origen": "Crossref API",
                "url": "https://www.imf.org/en/Publications/SPROLLs/working-papers",
                "enlace_id": "Inv-1"
            }
        elif 'Blog' in titulo or 'blogs' in str(row.get('Link', '')).lower():
            return {
                "nombre": "Blogs",
                "origen": "API Coveo",
                "url": "https://www.imf.org/en/Blogs",
                "enlace_id": "Inv-2"
            }
        else:
            return {
                "nombre": "Investigación FMI",
                "origen": "Crossref API",
                "url": "https://www.imf.org/en/Research",
                "enlace_id": "Inv-Other"
            }
    
    # === DISCURSOS ===
    elif categoria == 'Discursos':
        if 'Speech' in titulo or 'Transcript' in titulo or 'Remarks' in titulo:
            return {
                "nombre": "Speeches",
                "origen": "API Coveo",
                "url": "https://www.imf.org/en/News/Speeches",
                "enlace_id": "Disc-1"
            }
        else:
            return {
                "nombre": "Discursos FMI",
                "origen": "API Coveo",
                "url": "https://www.imf.org/en/News/Speeches",
                "enlace_id": "Disc-Other"
            }
    
    # Default
    return {
        "nombre": "Otra fuente",
        "origen": "No identificado",
        "url": "",
        "enlace_id": "N/A"
    }

# ==========================================
# CONFIGURACIÓN INICIAL Y ESTILOS
# ==========================================
st.set_page_config(page_title="Boletín Mensual", layout="wide")

st.markdown("""
    <style>
    div.stButton > button, div.stDownloadButton > button {
        background-color: #00205B !important;
        color: white !important;
        border: none !important;
    }
    div.stButton > button:hover, div.stDownloadButton > button:hover {
        background-color: #00153D !important;
        color: white !important;
    }
    span[data-baseweb="tag"] {
        background-color: #00205B !important;
        color: white !important;
    }
.github-footer {
    position: fixed;
    right: 20px;
    bottom: 20px;
    background-color: rgba(255, 255, 255, 0.9);
    padding: 8px 12px;
    border-radius: 50px;
    border: 1px solid #d0d7de;
    z-index: 1000;
    display: flex;
    align-items: center;
    gap: 12px;
    font-family: 'Calibri', sans-serif;
    box-shadow: 0px 4px 12px rgba(0,0,0,0.1);
    transition: transform 0.2s, box-shadow 0.2s;
}
.github-footer:hover {
    transform: translateY(-2px);
    box-shadow: 0px 6px 16px rgba(0,0,0,0.15);
}
.github-footer a {
    text-decoration: none;
    color: #24292f;
    display: flex;
    align-items: center;
    gap: 6px;
    transition: color 0.2s;
}
.github-footer a:hover {
    color: #00205B;
}
.github-icon {
    width: 20px;
    height: 20px;
}
.separator {
    color: #d0d7de;
    font-weight: normal;
}
</style>
<div class="github-footer">
    <a href="https://github.com/sdiazprado" target="_blank">
        <img class="github-icon" src="https://github.githubassets.com/images/modules/logos_page/GitHub-Mark.png" alt="GitHub Logo">
        <span><strong>@sdiazprado</strong></span>
    </a>
    <span class="separator">|</span>
    <a href="https://github.com/barcevelasco" target="_blank">
        <img class="github-icon" src="https://github.githubassets.com/images/modules/logos_page/GitHub-Mark.png" alt="GitHub Logo">
        <span><strong>@barcevelasco</strong></span>
    </a>
</div>
""", unsafe_allow_html=True)

# ==========================================
# UTILIDADES DE FORMATO
# ==========================================
# ==========================================
# HERRAMIENTA DE RESCATE (TEXTO MANUAL)
# ==========================================
@st.cache_data(show_spinner=False)
def buscar_link_inteligente(titulo, organismo):
    """Cazador de DOIs de Doble Impacto (Estricto + Fuzzy). Cero Google."""
    import urllib.parse
    import requests
    import time
    import re

    # 1. Limpieza base
    titulo_raiz = re.split(r'[:\-]', titulo)[0].strip()
    titulo_limpio = re.sub(r'[^a-zA-Z0-9\s]', '', titulo_raiz)
    
    headers = {'User-Agent': 'mailto:bot_investigacion@banco.com'}
    time.sleep(0.5) 

    def consultar_api(query_param, texto_busqueda, modo_estricto=True):
        query_enc = urllib.parse.quote(texto_busqueda)
        url = f"https://api.crossref.org/works?{query_param}={query_enc}&select=URL,title,publisher&rows=4"
        
        try:
            res = requests.get(url, headers=headers, timeout=8)
            if res.status_code == 200:
                items = res.json().get('message', {}).get('items', [])
                
                for item in items:
                    url_oficial = item.get('URL')
                    if not url_oficial: continue
                        
                    pub = item.get('publisher', '').lower()
                    titulo_api = item.get('title', [''])[0].lower()
                    
                    if modo_estricto:
                        if 'oecd' in pub or 'organisation' in pub or organismo.lower() in pub:
                            return url_oficial
                    else:
                        titulo_comparar = titulo_limpio.lower()
                        if titulo_comparar in titulo_api or titulo_api in titulo_comparar:
                            return url_oficial
        except:
            pass
        return None

    link = consultar_api("query.title", titulo_limpio, modo_estricto=True)
    if link: return link

    time.sleep(0.5)
    link = consultar_api("query.bibliographic", titulo, modo_estricto=False)
    if link: return link

    return ""

def procesar_texto_pegado(texto_crudo, organismo_nombre):
    """Extrae Fecha y Título del texto pegado. Retorna DataFrame estandarizado."""
    rows = []
    lineas = [linea.strip() for linea in texto_crudo.split('\n') if linea.strip()]
    patron_fecha = r'(\d{1,2}\s+[A-Za-z]{3,}\s+\d{4})'
    
    i = 0
    while i < len(lineas):
        match_fecha = re.search(patron_fecha, lineas[i])
        if match_fecha:
            try:
                parsed_date = parser.parse(match_fecha.group(1))
            except:
                i += 1; continue
            
            titulo = ""
            if i >= 1:
                titulo = lineas[i-1]
                basura_menu = ['list view', 'grid view', 'z-a', 'a-z', 'oldest', 'most recent', 'most relevant', 'order by']
                if titulo.lower() in basura_menu and i >= 2: 
                    titulo = lineas[i-2]
            
            if titulo and len(titulo) > 10 and not any(b in titulo.lower() for b in ['search', 'filter', 'sort by', 'publications']):
                rows.append({
                    "Date": parsed_date, 
                    "Title": titulo,
                    "Link": "Pendiente",
                    "Organismo": organismo_nombre
                })
        i += 1
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.sort_values(by="Date", ascending=False).drop_duplicates(subset=['Title'])
    return df

def buscar_link_boe(titulo):
    """Busca silenciosamente en la web para obtener el Link Directo y Oficial del BoE"""
    import urllib.parse
    import requests
    from bs4 import BeautifulSoup
    import re
    
    # Extraemos solo el título limpio sin el autor para la búsqueda
    titulo_limpio = titulo.split(': ')[-1] if ': ' in titulo else titulo
    titulo_limpio = re.sub(r'[^a-zA-Z0-9\s]', '', titulo_limpio)
    
    # Usamos DuckDuckGo HTML para evadir bloqueos y obtener el link oficial sin usar Google
    query = f"site:bankofengland.co.uk/speech {titulo_limpio}"
    url = f"https://html.duckduckgo.com/html/?q={urllib.parse.quote(query)}"
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
    
    try:
        res = requests.get(url, headers=headers, timeout=8)
        soup = BeautifulSoup(res.text, 'html.parser')
        
        # Atrapamos el link real de los resultados
        for a in soup.find_all('a', class_='result__url'):
            href = a.get('href', '').strip()
            if 'bankofengland.co.uk/speech' in href:
                if not href.startswith('http'):
                    href = 'https://' + href
                return href
    except:
        pass
        
    # Fallback de emergencia (1 clic)
    google_query = urllib.parse.quote(query)
    return f"https://www.google.com/search?q={google_query}"

def procesar_texto_pegado_boe(texto_crudo):
    """Extractor especializado para el formato del Bank of England (BoE)"""
    rows = []
    lineas = [linea.strip() for linea in texto_crudo.split('\n') if linea.strip()]
    patron_fecha = r'(\d{1,2}\s+[A-Za-z]{3,}\s+\d{4})'
    
    i = 0
    while i < len(lineas):
        match_fecha = re.search(patron_fecha, lineas[i])
        if match_fecha:
            try:
                parsed_date = parser.parse(match_fecha.group(1))
            except:
                i += 1; continue
            
            # 1. Buscar Autor un renglón ARRIBA (ej. "Speech // Phil Evans")
            autor = ""
            if i >= 1 and "//" in lineas[i-1]:
                partes = lineas[i-1].split("//")
                if len(partes) > 1:
                    autor = clean_author_name(partes[1].strip())
            
            # 2. Buscar Título Completo dos renglones ABAJO
            titulo = ""
            if i + 2 < len(lineas):
                titulo_raw = lineas[i+2]
                # Le quitamos el sufijo redundante " - speech by Autor"
                titulo_raw = re.sub(r'(?i)\s*[\-–—]\s*speech\s+by\s+.*$', '', titulo_raw).strip()
                titulo = titulo_raw
            
            # 3. Ensamblar y Guardar
            if titulo:
                titulo_final = f"{autor}: {titulo}" if autor else titulo
                rows.append({
                    "Date": parsed_date, 
                    "Title": titulo_final,
                    "Link": "Pendiente",
                    "Organismo": "BoE (Inglaterra)"
                })
        i += 1
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.sort_values(by="Date", ascending=False).drop_duplicates(subset=['Title'])
    return df
def clean_author_name(name):
    if not name: return ""
    minusc = ['de', 'van', 'von', 'der', 'del', 'la']
    words = name.strip().split()
    
    # Capitaliza todo excepto las preposiciones europeas
    cleaned_words = [w.capitalize() if w.lower() not in minusc else w.lower() for w in words]
    if cleaned_words:
        cleaned_words[0] = cleaned_words[0].capitalize() # La primera siempre mayúscula
        
    cleaned = " ".join(cleaned_words)
    # Arreglar iniciales pegadas (ej. "J.M. Keynes" -> "J. M. Keynes")
    cleaned = re.sub(r'\b([A-Z])\.\s*([A-Z])', lambda m: f"{m.group(1)}. {m.group(2)}", cleaned)
    return cleaned



# ==========================================
# FUNCIONES DE EXTRACCIÓN (BACKEND)
# ==========================================

@st.cache_data(show_spinner=False)
def load_reportes_fem(start_date_str, end_date_str):
    """Extractor FEM - Versión Selenium Final (Scroll + Fallback de Fecha)"""
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    import time
    import re

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2025, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    # CAMBIO IMPORTANTE: Eliminar el filtro de tipos
    url = "https://es.weforum.org/publications/?years=2026"  # ← Sin filter de tipos
    
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36")

    try:
        driver = webdriver.Chrome(options=chrome_options)
        driver.get(url)
        time.sleep(8)
        driver.execute_script("window.scrollTo(0, 1000);")
        time.sleep(4)
        
        js_script = """
        let res = [];
        let seenLinks = new Set();
        
        document.querySelectorAll('a[href*="/publications/"]').forEach(el => {
            let title = el.innerText || el.textContent || "";
            let href = el.href;
            
            title = title.split('\\n')[0];
            title = title.replace(/Download PDF|Leer más|Read more|View details/gi, '').trim();
            
            if (title.length > 15 && !seenLinks.has(href) && !href.includes('/series/')) {
                seenLinks.add(href);
                
                let container = el.closest('article') || el.closest('div[class*="publication"]') || el.parentElement;
                let date = "";
                
                let dateEl = container ? container.querySelector('.date, time, [class*="date"], [class*="Date"]') : null;
                if (dateEl) {
                    date = dateEl.innerText || dateEl.textContent || "";
                }
                
                if (!date) {
                    let siblings = el.parentElement ? el.parentElement.querySelectorAll('div, span, p') : [];
                    for (let sib of siblings) {
                        let text = sib.innerText || "";
                        if (text.match(/\\d{1,2}\\s+[A-Za-z]{3,}\\s+\\d{4}/)) {
                            date = text;
                            break;
                        }
                    }
                }
                
                res.push({ t: title, l: href, d: date });
            }
        });
        return res;
        """
        
        extracted = driver.execute_script(js_script)
        driver.quit()

        print(f"   📚 Total de ítems encontrados: {len(extracted)}")
        
        meses_map = {
            'ene': 1, 'feb': 2, 'mar': 3, 'abr': 4, 'may': 5, 'jun': 6,
            'jul': 7, 'ago': 8, 'sep': 9, 'sept': 9, 'oct': 10, 'nov': 11, 'dic': 12,
            'enero': 1, 'febrero': 2, 'marzo': 3, 'abril': 4, 'mayo': 5, 'junio': 6,
            'julio': 7, 'agosto': 8, 'septiembre': 9, 'octubre': 10, 'noviembre': 11, 'diciembre': 12
        }

        for item in extracted:
            titulo = item['t']
            link = item['l']
            fecha_texto = item['d'].lower() if item['d'] else ""
            
            parsed_date = None
            
            match = re.search(r'(\d{1,2})\s+([a-z]{3,})\s+(\d{4})', fecha_texto)
            if match:
                dia = int(match.group(1))
                mes_str = match.group(2)[:3]
                año = int(match.group(3))
                mes_num = meses_map.get(mes_str, 1)
                try:
                    parsed_date = datetime.datetime(año, mes_num, dia)
                except:
                    parsed_date = datetime.datetime(año, mes_num, 1)
            
            if not parsed_date:
                m = re.search(r'/(\d{4})/(\d{2})/', link)
                if m:
                    parsed_date = datetime.datetime(int(m.group(1)), int(m.group(2)), 1)

            if parsed_date and start_date <= parsed_date <= end_date:
                if not any(r['Link'] == link for r in rows):
                    rows.append({
                        "Date": parsed_date, 
                        "Title": titulo, 
                        "Link": link, 
                        "Organismo": "FEM"
                    })
                    print(f"   ✅ {parsed_date.strftime('%d/%m/%Y')}: {titulo[:50]}...")
                    
    except Exception as e:
        print(f"❌ Error en load_reportes_fem: {e}")
        import traceback
        traceback.print_exc()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False).drop_duplicates(subset=['Link'])
    
    print(f"\n📋 TODOS los títulos encontrados ({len(df)} documentos):")
    for i, row in df.iterrows():
        print(f"   - {row['Date'].strftime('%d/%m/%Y')}: {row['Title'][:60]}...")
    
    print(f"📊 FEM Reportes - Total final: {len(df)} documentos")
    return df



# BID (Annual Reports en inglés)
@st.cache_data(show_spinner=False)
def load_reportes_bid_en(start_date_str, end_date_str):
    """
    Extrae Annual Reports del BID en inglés usando cloudscraper
    (mismo método que funciona para BID Investigación)
    """
    import cloudscraper
    from bs4 import BeautifulSoup
    import datetime
    import re
    import time
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BID Reportes: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    page = 0
    max_pages = 5
    
    # Crear scraper con la misma configuración que usas en BID Investigación
    scraper = cloudscraper.create_scraper(
        browser={
            'browser': 'chrome',
            'platform': 'windows',
            'mobile': False
        },
        delay=5
    )
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
        'Accept-Language': 'en-US,en;q=0.5',
        'Accept-Encoding': 'gzip, deflate, br',
        'Connection': 'keep-alive',
        'Upgrade-Insecure-Requests': '1',
    }
    
    meses_map = {
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12,
        'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
        'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12
    }
    
    while page < max_pages:
        url = f"https://publications.iadb.org/en?f%5B0%5D=type%3AAnnual%20Reports&page={page}"
        print(f"📄 Página {page+1}: {url}")
        
        try:
            response = scraper.get(url, headers=headers, timeout=30)
            
            if response.status_code != 200:
                print(f"   ❌ Error HTTP: {response.status_code}")
                break
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Buscar artículos
            items = soup.find_all('div', class_='views-row')
            
            if not items:
                print(f"   📭 No hay resultados en página {page+1}")
                # Guardar HTML para depuración
                with open(f"bid_reportes_page_{page}_debug.html", "w", encoding="utf-8") as f:
                    f.write(response.text)
                print(f"   💾 HTML guardado en bid_reportes_page_{page}_debug.html")
                break
            
            print(f"   📚 Artículos encontrados: {len(items)}")
            
            items_found = 0
            for item in items:
                try:
                    # Título y link
                    title_div = item.find('div', class_='views-field-field-title')
                    if not title_div:
                        continue
                    
                    a_tag = title_div.find('a')
                    if not a_tag:
                        continue
                    
                    titulo = a_tag.get_text(strip=True)
                    link = a_tag.get('href')
                    if link and not link.startswith('http'):
                        link = "https://publications.iadb.org" + link
                    
                    # Fecha
                    date_div = item.find('div', class_='views-field-field-date-issued-text')
                    if not date_div:
                        continue
                    
                    date_text = date_div.get_text(strip=True)
                    match = re.search(r'([A-Za-z]{3,9})\s+(\d{4})', date_text)
                    if not match:
                        continue
                    
                    mes_str = match.group(1).lower()[:3]
                    año = int(match.group(2))
                    mes_num = meses_map.get(mes_str, 1)
                    parsed_date = datetime.datetime(año, mes_num, 15)
                    
                    # Filtrar por fecha
                    if parsed_date < start_date or parsed_date > end_date:
                        continue
                    
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo,
                            "Link": link,
                            "Organismo": "BID (Reportes)"
                        })
                        items_found += 1
                        print(f"   ✅ {parsed_date.date()} - {titulo[:50]}...")
                        
                except Exception as e:
                    continue
            
            print(f"   📊 Documentos en página {page+1}: {items_found}")
            
            if items_found == 0 and page > 0:
                break
            
            page += 1
            time.sleep(2)
            
        except Exception as e:
            print(f"   ❌ Error: {e}")
            break
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.drop_duplicates(subset=['Link'])
        df = df.sort_values("Date", ascending=False)
    
    print(f"✅ BID Reportes - Total: {len(df)} documentos")
    return df


@st.cache_data(show_spinner=False)
def load_reportes_bpi(start_date_str, end_date_str):
    urls_api = [
        "https://www.bis.org/api/document_lists/bcbspubls.json",
        "https://www.bis.org/api/document_lists/cpmi_publs.json"
    ]
    urls_html = ["https://www.bis.org/ifc/publications.htm"]
    headers = {'User-Agent': 'Mozilla/5.0'}

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)

    rows = []

    for url in urls_api:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            data = res.json()
            lista_documentos = data.get("list", {})
            for path, doc_info in lista_documentos.items():
                titulo = html.unescape(doc_info.get("short_title", ""))
                if not titulo:
                    continue
                link = "https://www.bis.org" + doc_info.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"):
                    link += ".htm"
                date_str = doc_info.get("publication_start_date", "")
                parsed_date = None
                if date_str:
                    try:
                        parsed_date = parser.parse(date_str)
                    except:
                        pass
                if not parsed_date:
                    continue
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo,
                                "Link": link, "Organismo": "BPI"})
        except Exception as e:
            continue

    for url in urls_html:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            soup = BeautifulSoup(res.text, 'html.parser')
            content_div = soup.find('div', id='cmsContent')
            if not content_div:
                continue
            for p in content_div.find_all('p'):
                a_tag = p.find('a')
                if not a_tag:
                    continue
                titulo = a_tag.get_text(strip=True)
                href = a_tag.get('href', '')
                if not href or 'index.htm' in href:
                    continue
                link = "https://www.bis.org" + \
                    href if href.startswith('/') else href
                full_text = p.get_text(strip=True)
                date_str = full_text.replace(titulo, '').strip(', ')
                parsed_date = None
                if date_str:
                    try:
                        parsed_date = parser.parse(date_str)
                    except:
                        pass
                if not parsed_date:
                    match = re.search(r'\b(20\d{2})\b', titulo)
                    if match:
                        parsed_date = datetime.datetime(
                            int(match.group(1)), 1, 1)
                if not parsed_date:
                    continue
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo,
                                "Link": link, "Organismo": "BPI"})
        except Exception as e:
            continue

    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None:
            df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df

## Reportes BM 
@st.cache_data(show_spinner=False)
def load_reportes_bm(start_date_str, end_date_str):
    """Extractor para Reportes del BM usando API de DSpace filtrando por doctype."""
    base_url = (
        "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    )
    headers = {"User-Agent": "Mozilla/5.0"}
    scope_id = "06251f8a-62c2-59fb-add5-ec0993fc20d9"

    try:
        start_date = datetime.datetime.strptime(start_date_str, "%d.%m.%Y")
        end_date = datetime.datetime.strptime(end_date_str, "%d.%m.%Y")
        print(f"📅 BM Reportes: {start_date.date()} a {end_date.date()}")
    except Exception:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    page = 0
    max_pages = 10

    while page < max_pages:
        try:
            params = {
                "scope": scope_id,
                "sort": "dc.date.issued,DESC",
                "page": page,
                "size": 50,
                "f.doctype": "Report,equals",  # Filtro nativo de DSpace
            }
            res = requests.get(
                base_url, headers=headers, params=params, timeout=15
            )
            data = res.json()

            objects = (
                data.get("_embedded", {})
                .get("searchResult", {})
                .get("_embedded", {})
                .get("objects", [])
            )

            if not objects:
                print(f"📭 No hay más resultados en página {page}")
                break

            print(f"📄 Página {page + 1}: {len(objects)} objetos encontrados")

            items_found = 0
            for obj in objects:
                item = obj.get("_embedded", {}).get("indexableObject", {})
                meta = item.get("metadata", {})

                # Extraer Título
                title = (
                    meta.get("dc.title", [{"value": ""}])[0].get("value", "")
                )
                if not title:
                    continue

                # Extraer Fecha
                date_s = (
                    meta.get("dc.date.issued", [{"value": ""}])[0].get(
                        "value", ""
                    )
                )
                if not date_s:
                    continue

                try:
                    parsed_date = parser.parse(date_s)
                    if parsed_date.tzinfo is not None:
                        parsed_date = parsed_date.replace(tzinfo=None)
                except Exception:
                    continue

                # Filtrar rango de fechas
                if parsed_date < start_date or parsed_date > end_date:
                    continue

                # Link permanente
                link = (
                    meta.get("dc.identifier.uri", [{"value": ""}])[0].get(
                        "value", ""
                    )
                )
                if not link:
                    link = f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"

                if not any(r["Link"] == link for r in rows):
                    rows.append({
                        "Date": parsed_date,
                        "Title": title,
                        "Link": link,
                        "Organismo": "BM",
                    })
                    items_found += 1
                    print(f"   ✅ {parsed_date.date()} - {title[:60]}...")

            print(f"   📊 Documentos en página {page + 1}: {items_found}")

            if items_found == 0 and page > 1:
                break

            page += 1
            time.sleep(0.5)

        except Exception as e:
            print(f"⚠️ Error en página {page}: {e}")
            break

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=["Link"])

    print(f"✅ BM Reportes - Total: {len(df)} documentos")
    return df

@st.cache_data(show_spinner=False)
def load_reportes_cef(start_date_str, end_date_str):
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 1
    while True:
        url = f"https://www.fsb.org/publications/?dps_paged={page}"
        try:
            res = requests.get(url, headers=headers, timeout=15)
            soup = BeautifulSoup(res.text, 'html.parser')
            items = soup.find_all(
                'div', class_=lambda c: c and 'post-excerpt' in c)
            if not items:
                break
            items_found = 0
            for item in items:
                title_div = item.find('div', class_='post-title')
                if not title_div or not title_div.find('a'):
                    continue
                a_tag = title_div.find('a')
                titulo_raw = a_tag.get_text(strip=True)
                link = a_tag.get('href', '')
                date_div = item.find('div', class_='post-date')
                parsed_date = None
                if date_div:
                    try:
                        parsed_date = parser.parse(
                            date_div.get_text(strip=True))
                    except:
                        pass
                if not parsed_date:
                    continue
                if not any(r['Link'] == link for r in rows):
                    rows.append(
                        {"Date": parsed_date, "Title": titulo_raw, "Link": link, "Organismo": "CEF"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date):
                break
            page += 1
            time.sleep(0.5)
        except:
            break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


# -- OCDE -- REPORTES -- 
@st.cache_data(show_spinner=False)
def load_reportes_ocde(start_date_str, end_date_str):
    """Extractor OCDE - Reports (API oficial)"""
    import requests
    import datetime
    import re
    import time
    from dateutil import parser

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 OCDE Reportes: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    rows = []

    # API base de la OCDE
    base_url = "https://api.oecd.org/webcms/search/faceted-search"

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "application/json"
    }

    page = 0
    page_size = 50  # Número de resultados por página
    max_pages = 10  # Límite de seguridad
    documentos_procesados = 0

    print("📡 Solicitando Reportes a la API de la OCDE (con paginación)...")

    try:
        while page < max_pages:
            # Parámetros para buscar Reports en inglés
            params = {
                "siteName": "oecd",
                "interfaceLanguage": "en",
                "orderBy": "mostRecent",
                "pageSize": page_size,
                "page": page,
                "facets": "oecd-languages:en",
                "hiddenFacets": "oecd-content-types:publications/reports"  # <-- FILTRO PARA REPORTES
            }

            print(f"   📄 Procesando página {page + 1}...")
            response = requests.get(base_url, params=params, headers=headers, timeout=15)

            if response.status_code != 200:
                print(f"   ❌ Error en página {page + 1}: {response.status_code}")
                break

            data = response.json()

            # Buscar los resultados
            results = data.get("results", [])

            if not results:
                print(f"   📭 No hay más resultados en página {page + 1}")
                break

            documentos_en_pagina = 0
            fecha_mas_antigua = None

            for item in results:
                titulo = item.get("title", "") or item.get("name", "")
                link = item.get("url", "") or item.get("link", "")

                if not titulo or not link:
                    continue

                # Extraer fecha
                fecha_texto = item.get("publicationDateTime", "")
                parsed_date = None
                if fecha_texto:
                    try:
                        parsed_date = parser.parse(fecha_texto)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        continue

                if not parsed_date:
                    continue

                fecha_mas_antigua = parsed_date

                # Si el documento es más antiguo que start_date, paramos
                if parsed_date < start_date:
                    print(f"   ⏹️ Documento más antiguo que {start_date.strftime('%Y-%m')}, deteniendo paginación")
                    page = max_pages
                    break

                # Filtrar por rango de fechas
                if parsed_date >= start_date and parsed_date <= end_date:
                    # Limpiar título
                    titulo = re.sub(r'\s+', ' ', titulo).strip()

                    # Asegurar URL absoluta
                    if link.startswith('/'):
                        link = f"https://www.oecd.org{link}"

                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "OCDE"
                    })
                    documentos_en_pagina += 1
                    documentos_procesados += 1

            print(f"   📊 Página {page + 1}: {documentos_en_pagina} documentos en el rango")

            # Si no encontramos documentos en esta página y ya pasamos la fecha límite
            if documentos_en_pagina == 0 and fecha_mas_antigua and fecha_mas_antigua < start_date:
                print(f"   ⏹️ Fin de resultados para el mes solicitado")
                break

            # Si encontramos menos de page_size documentos, probablemente es la última página
            if len(results) < page_size:
                print(f"   📭 Última página alcanzada")
                break

            page += 1
            time.sleep(0.3)  # Pequeña pausa para no sobrecargar la API

        print(f"\n📊 Total Reportes OCDE encontrados: {documentos_procesados}")

    except Exception as e:
        print(f"❌ Error en load_reportes_ocde: {e}")
        import traceback
        traceback.print_exc()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])

    print(f"📊 OCDE Reportes - Total final: {len(df)}")
    return df


@st.cache_data(show_spinner=False)
def load_reportes_bpi(start_date_str, end_date_str):
    """
    Extractor BPI - Reportes (BCBS, CPMI, IFC, CGFS)
    """
    import requests
    import pandas as pd
    import datetime
    import html
    from dateutil import parser
    from bs4 import BeautifulSoup
    import re
    import time
    
    # Configuración de fechas
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BPI Reportes: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    # ========== 1. URLs EXISTENTES (BCBS, CPMI) ==========
    urls_api = [
        "https://www.bis.org/api/document_lists/bcbspubls.json",
        "https://www.bis.org/api/document_lists/cpmi_publs.json"
    ]
    
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []

    # API calls (BCBS y CPMI) - TU CÓDIGO EXISTENTE
    for url in urls_api:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            data = res.json()
            lista_documentos = data.get("list", {})
            for path, doc_info in lista_documentos.items():
                titulo = html.unescape(doc_info.get("short_title", ""))
                if not titulo:
                    continue
                link = "https://www.bis.org" + doc_info.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"):
                    link += ".htm"
                date_str = doc_info.get("publication_start_date", "")
                parsed_date = None
                if date_str:
                    try:
                        parsed_date = parser.parse(date_str)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        pass
                if not parsed_date:
                    continue
                if start_date <= parsed_date <= end_date:
                    rows.append({"Date": parsed_date, "Title": titulo,
                                "Link": link, "Organismo": "BPI"})
        except Exception as e:
            continue

    # ========== 2. IFC publications (HTML) - TU CÓDIGO EXISTENTE ==========
    urls_html = ["https://www.bis.org/ifc/publications.htm"]
    for url in urls_html:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            soup = BeautifulSoup(res.text, 'html.parser')
            content_div = soup.find('div', id='cmsContent')
            if not content_div:
                continue
            for p in content_div.find_all('p'):
                a_tag = p.find('a')
                if not a_tag:
                    continue
                titulo = a_tag.get_text(strip=True)
                href = a_tag.get('href', '')
                if not href or 'index.htm' in href:
                    continue
                link = "https://www.bis.org" + href if href.startswith('/') else href
                full_text = p.get_text(strip=True)
                date_str = full_text.replace(titulo, '').strip(', ')
                parsed_date = None
                if date_str:
                    try:
                        parsed_date = parser.parse(date_str)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        pass
                if not parsed_date:
                    match = re.search(r'\b(20\d{2})\b', titulo)
                    if match:
                        parsed_date = datetime.datetime(int(match.group(1)), 1, 1)
                if not parsed_date:
                    continue
                if start_date <= parsed_date <= end_date:
                    rows.append({"Date": parsed_date, "Title": titulo,
                                "Link": link, "Organismo": "BPI"})
        except Exception as e:
            continue

    # ========== 3. NUEVO: CGFS publications (API) ==========
    try:
        url_cgfs_api = "https://www.bis.org/api/document_lists/cgfs_publs.json"
        res = requests.get(url_cgfs_api, headers=headers, timeout=15)
        if res.status_code == 200:
            data = res.json()
            for path, doc in data.get("list", {}).items():
                titulo = html.unescape(doc.get("short_title", ""))
                if not titulo:
                    continue
                link = "https://www.bis.org" + doc.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"):
                    link += ".htm"
                try:
                    parsed_date = parser.parse(doc.get("publication_start_date", ""))
                    if parsed_date.tzinfo is not None:
                        parsed_date = parsed_date.replace(tzinfo=None)
                    if start_date <= parsed_date <= end_date:
                        rows.append({"Date": parsed_date, "Title": titulo,
                                    "Link": link, "Organismo": "BPI"})
                except:
                    continue
    except Exception as e:
        print(f"   ⚠️ Error en CGFS API: {e}")

    # ========== 4. NUEVO: CGFS HTML (para documento No 71 que no está en API) ==========
    try:
        url_cgfs_html = "https://www.bis.org/cgfs_publs/index.htm"
        res = requests.get(url_cgfs_html, headers=headers, timeout=15)
        if res.status_code == 200:
            soup = BeautifulSoup(res.text, 'html.parser')
            
            # Buscar filas de la tabla
            rows_html = soup.find_all('tr')
            
            for row in rows_html:
                try:
                    cells = row.find_all('td')
                    if len(cells) < 2:
                        continue
                    
                    # Fecha en primera celda
                    date_text = cells[0].get_text(strip=True)
                    match = re.search(r'(\d{1,2})\s+([A-Za-z]{3,})\s+(\d{4})', date_text)
                    if not match:
                        continue
                    
                    day = int(match.group(1))
                    mes_str = match.group(2).lower()
                    año = int(match.group(3))
                    
                    meses = {'jan':1, 'feb':2, 'mar':3, 'apr':4, 'may':5, 'jun':6,
                            'jul':7, 'aug':8, 'sep':9, 'oct':10, 'nov':11, 'dec':12}
                    mes = meses.get(mes_str[:3], 1)
                    parsed_date = datetime.datetime(año, mes, day)
                    
                    if parsed_date < start_date or parsed_date > end_date:
                        continue
                    
                    # Título y enlace en segunda celda
                    link_elem = cells[1].find('a')
                    if not link_elem:
                        continue
                    
                    titulo = link_elem.get_text(strip=True)
                    link = link_elem.get('href')
                    if link and not link.startswith('http'):
                        link = "https://www.bis.org" + link
                    
                    # Evitar duplicados con los de la API
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo,
                            "Link": link,
                            "Organismo": "BPI"
                        })
                        print(f"   ✅ CGFS HTML: {parsed_date.date()} - {titulo[:50]}...")
                        
                except Exception as e:
                    continue
                    
    except Exception as e:
        print(f"   ⚠️ Error scraping CGFS: {e}")

    # ========== 5. Crear DataFrame final ==========
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None:
            df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    
    print(f"✅ BPI Reportes - Total final: {len(df)} documentos")
    return df



# --- SECCIÓN: PUBLICACIONES INSTITUCIONALES ---

# --- Publicaciones Institucionales --- OCDE 

@st.cache_data(show_spinner=False)
def load_pub_inst_ocde(start_date_str, end_date_str):
    """Extractor OCDE - Publicaciones Institucionales (API oficial)"""
    import requests
    import datetime
    import re
    import time
    from dateutil import parser
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 OCDE Pub. Institucionales: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    
    # API base de la OCDE
    base_url = "https://api.oecd.org/webcms/search/faceted-search"
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "application/json"
    }
    
    page = 0
    page_size = 50
    max_pages = 10
    
    print("📡 Solicitando Publicaciones Institucionales a la API de la OCDE (con paginación)...")
    
    try:
        while page < max_pages:
            # Parámetros para buscar el sub-tema psi114
            params = {
                "siteName": "oecd",
                "interfaceLanguage": "en",
                "orderBy": "mostRecent",
                "pageSize": page_size,
                "page": page,
                "facets": "oecd-languages:en",
                "hiddenFacets": "oecd-policy-subissues:psi114"  # <-- FILTRO PARA PUB. INSTITUCIONALES
            }
            
            print(f"   📄 Procesando página {page + 1}...")
            response = requests.get(base_url, params=params, headers=headers, timeout=15)
            
            if response.status_code != 200:
                print(f"   ❌ Error en página {page + 1}: {response.status_code}")
                break
            
            data = response.json()
            
            # Buscar los resultados
            results = data.get("results", [])
            
            if not results:
                print(f"   📭 No hay más resultados en página {page + 1}")
                break
            
            documentos_en_pagina = 0
            fecha_mas_antigua = None
            
            for item in results:
                titulo = item.get("title", "") or item.get("name", "")
                link = item.get("url", "") or item.get("link", "")
                
                if not titulo or not link:
                    continue
                
                # Extraer fecha
                fecha_texto = item.get("publicationDateTime", "")
                
                parsed_date = None
                if fecha_texto:
                    try:
                        parsed_date = parser.parse(fecha_texto)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        continue
                
                if not parsed_date:
                    continue
                
                fecha_mas_antigua = parsed_date
                
                # Si el documento es más antiguo que start_date, paramos
                if parsed_date < start_date:
                    print(f"   ⏹️ Documento más antiguo que {start_date.strftime('%Y-%m')}, deteniendo paginación")
                    page = max_pages
                    break
                
                # Filtrar por rango de fechas
                if parsed_date >= start_date and parsed_date <= end_date:
                    # Limpiar título
                    titulo = re.sub(r'\s+', ' ', titulo).strip()
                    
                    # Asegurar URL absoluta
                    if link.startswith('/'):
                        link = f"https://www.oecd.org{link}"
                    
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "OCDE"
                    })
                    documentos_en_pagina += 1
            
            print(f"   📊 Página {page + 1}: {documentos_en_pagina} documentos en el rango")
            
            # Si no encontramos documentos en esta página y ya pasamos la fecha límite
            if documentos_en_pagina == 0 and fecha_mas_antigua and fecha_mas_antigua < start_date:
                print(f"   ⏹️ Fin de resultados para el mes solicitado")
                break
            
            # Si encontramos menos de page_size documentos, probablemente es la última página
            if len(results) < page_size:
                print(f"   📭 Última página alcanzada")
                break
            
            page += 1
            time.sleep(0.3)
        
        print(f"\n📊 Total documentos OCDE Pub. Institucionales encontrados: {len(rows)}")
        
    except Exception as e:
        print(f"❌ Error en load_pub_inst_ocde: {e}")
        import traceback
        traceback.print_exc()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"📊 OCDE Pub. Institucionales - Total final: {len(df)}")
    return df

# --- Publicaciones Institucionales --- OEI 
@st.cache_data(show_spinner=False)
def load_pub_inst_oei(start_date_str, end_date_str):
    """Extractor OEI (IEO-IMF) - Versión API Next.js con headers completos"""
    import requests
    import datetime
    import re
    from dateutil import parser
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 OEI: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    
    # Intentar obtener el build ID dinámicamente desde la página HTML
    build_id = "qchYZivFKVMGvRneSTtnM"  # Fallback
    try:
        headers_browser = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate, br',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
            'Sec-Fetch-Dest': 'document',
            'Sec-Fetch-Mode': 'navigate',
            'Sec-Fetch-Site': 'none',
            'Sec-Fetch-User': '?1',
            'Cache-Control': 'max-age=0',
        }
        res = requests.get("https://ieo.imf.org/en/Publications/annual-reports", headers=headers_browser, timeout=15)
        # Buscar el build ID en el HTML
        match = re.search(r'/_next/data/([a-zA-Z0-9]+)/en/publications/annual-reports\.json', res.text)
        if match:
            build_id = match.group(1)
            print(f"🔧 Build ID encontrado: {build_id}")
    except Exception as e:
        print(f"⚠️ Usando build ID por defecto: {build_id}")
    
    # URL del JSON
    url = f"https://ieo.imf.org/_next/data/{build_id}/en/publications/annual-reports.json"
    
    # Headers completos para simular un navegador real
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'application/json, text/plain, */*',
        'Accept-Language': 'en-US,en;q=0.5',
        'Accept-Encoding': 'gzip, deflate, br',
        'Referer': 'https://ieo.imf.org/en/Publications/annual-reports',
        'Origin': 'https://ieo.imf.org',
        'Connection': 'keep-alive',
        'Sec-Fetch-Dest': 'empty',
        'Sec-Fetch-Mode': 'cors',
        'Sec-Fetch-Site': 'same-origin',
        'Cache-Control': 'no-cache',
        'Pragma': 'no-cache',
    }
    
    try:
        print(f"📡 Consultando: {url}")
        response = requests.get(url, headers=headers, timeout=15)
        
        if response.status_code == 200:
            data = response.json()
            
            # ✅ RUTA CORRECTA según el JSON que analizamos
            try:
                # Los reportes están en componentProps.[id].fields.datasource.reports.results
                component_props = data.get('pageProps', {}).get('componentProps', {})
                
                # Buscar el componente ReportsListing
                reports_results = None
                for comp_id, comp_value in component_props.items():
                    if 'fields' in comp_value and 'datasource' in comp_value['fields']:
                        datasource = comp_value['fields']['datasource']
                        if 'reports' in datasource and 'results' in datasource['reports']:
                            reports_results = datasource['reports']['results']
                            print(f"✅ Componente encontrado: {comp_id}")
                            break
                
                if not reports_results:
                    print("⚠️ No se encontraron reportes en componentProps")
                    return pd.DataFrame()
                
                print(f"📚 Reportes encontrados: {len(reports_results)}")
                
                for report in reports_results:
                    # Extraer título
                    titulo = report.get('title', {}).get('jsonValue', {}).get('value', '')
                    
                    # Extraer fecha
                    fecha_texto = report.get('publicationDate', {}).get('jsonValue', {}).get('value', '')
                    
                    # Extraer link del PDF
                    completed_link = report.get('completedReportLink', {}).get('jsonValue', {}).get('value', {})
                    link = completed_link.get('href', '') if isinstance(completed_link, dict) else ''
                    
                    if not titulo or not fecha_texto:
                        continue
                    
                    # Parsear fecha
                    parsed_date = parser.parse(fecha_texto).replace(tzinfo=None)
                    
                    if parsed_date < start_date or parsed_date > end_date:
                        continue
                    
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "OEI"
                    })
                    print(f"   ✅ {parsed_date.strftime('%Y-%m-%d')}: {titulo}")
                
            except Exception as e:
                print(f"   ⚠️ Error procesando el JSON: {e}")
                import traceback
                traceback.print_exc()
        else:
            print(f"❌ Error en la API: {response.status_code}")
            print(f"   Respuesta: {response.text[:200] if response.text else 'Vacía'}")
            
    except Exception as e:
        print(f"❌ Error en load_pub_inst_oei: {e}")
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"📊 OEI - Total documentos: {len(df)}")
    return df
    
# ========== FUNCIÓN PARA CEMLA (PUBLICACIONES INSTITUCIONALES) ==========
@st.cache_data(show_spinner=False)
def load_pub_inst_cemla(start_date_str, end_date_str):
    """
    Extractor CEMLA - Publicaciones Institucionales (Novedades individuales)
    Filtra eventos y contenido no académico
    """
    import requests
    from bs4 import BeautifulSoup
    import datetime
    import re
    import pandas as pd
    import time
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    print("="*50)
    print("🔍 CEMLA PUBLICACIONES - Extrayendo novedades de boletines...")
    print(f"📅 Rango solicitado: {start_date_str} a {end_date_str}")
    print("="*50)

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"✅ Fechas parseadas: {start_date.date()} a {end_date.date()}")
    except Exception as e:
        print(f"⚠️ Error parseando fechas: {e}")
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now() + datetime.timedelta(days=365)

    # Palabras a excluir (eventos, cursos, etc. - no publicaciones académicas)
    palabras_excluir = [
        'reunión', 'reunion', 'virtual', 'curso', 'taller', 'seminario',
        'conferencia', 'webinar', 'congreso', 'foro', 'encuentro',
        'junta', 'comité', 'comite', 'próximas actividades', 'calendario',
        'convocatoria', 'premio', 'inscripción', 'registro'
    ]

    rows = []
    
    url = "https://www.cemla.org/comunicados.html"
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
    }

    try:
        # =========================================================
        # PASO 1: Obtener la lista de boletines
        # =========================================================
        print(f"📡 Solicitando página de boletines: {url}")
        response = requests.get(url, headers=headers, timeout=30, verify=False)
        print(f"   Status Code: {response.status_code}")
        
        if response.status_code != 200:
            print(f"❌ Error al acceder a la página")
            return pd.DataFrame()

        soup = BeautifulSoup(response.text, 'html.parser')
        
        meses = {
            'enero': 1, 'febrero': 2, 'marzo': 3, 'abril': 4,
            'mayo': 5, 'junio': 6, 'julio': 7, 'agosto': 8,
            'septiembre': 9, 'octubre': 10, 'noviembre': 11, 'diciembre': 12
        }
        
        # Encontrar todos los boletines en el rango de fechas
        boletines_a_procesar = []
        
        for ul in soup.find_all('ul', class_='iconlist'):
            for li in ul.find_all('li'):
                p = li.find('p')
                if not p:
                    continue
                
                a_tag = p.find('a')
                if not a_tag:
                    continue
                
                titulo_texto = a_tag.get_text(strip=True)
                link = a_tag.get('href', '')
                
                match = re.match(r'^([A-Za-z]+)\s+(\d{4})$', titulo_texto, re.IGNORECASE)
                if match:
                    mes_str, año = match.groups()
                    mes_num = meses.get(mes_str.lower(), 0)
                    
                    if mes_num:
                        fecha = datetime.datetime(int(año), mes_num, 1)
                        
                        if start_date <= fecha <= end_date:
                            boletines_a_procesar.append({
                                'fecha': fecha,
                                'titulo': titulo_texto,
                                'link': link
                            })
                            print(f"📌 Boletín encontrado: {fecha.strftime('%Y-%m')} - {titulo_texto}")
        
        print(f"✅ Total boletines en rango: {len(boletines_a_procesar)}")
        
        if not boletines_a_procesar:
            print("⚠️ No se encontraron boletines en el rango de fechas")
            return pd.DataFrame()
        
        # =========================================================
        # PASO 2: Procesar cada boletín y extraer sus novedades
        # =========================================================
        for boletin in boletines_a_procesar:
            print(f"\n🔍 Procesando boletín: {boletin['titulo']} ({boletin['link']})")
            
            try:
                time.sleep(1)
                
                res_boletin = requests.get(boletin['link'], headers=headers, timeout=30, verify=False)
                if res_boletin.status_code != 200:
                    print(f"  ⚠️ Error al acceder al boletín: {res_boletin.status_code}")
                    continue
                
                soup_boletin = BeautifulSoup(res_boletin.text, 'html.parser')
                
                # Buscar todas las novedades (divs con clase "ipost clearfix")
                novedades = soup_boletin.find_all('div', class_=lambda c: c and 'ipost' in c.split() if c else False)
                
                if not novedades:
                    print(f"  ⚠️ No se encontraron novedades en este boletín")
                    continue
                
                print(f"  📚 Novedades encontradas: {len(novedades)}")
                
                for novedad in novedades:
                    try:
                        # Extraer título
                        title_elem = novedad.find('div', class_='entry-title')
                        if not title_elem:
                            continue
                        
                        h3 = title_elem.find('h3')
                        if not h3:
                            continue
                        
                        titulo = h3.get_text(strip=True)
                        
                        # ===== FILTRO: Excluir eventos y contenido no académico =====
                        titulo_lower = titulo.lower()
                        es_excluido = any(palabra in titulo_lower for palabra in palabras_excluir)
                        
                        if es_excluido:
                            print(f"    ⏭️ Excluido (evento): {titulo[:60]}...")
                            continue
                        
                        # Extraer descripción y enlace
                        content_elem = novedad.find('div', class_='entry-content')
                        if not content_elem:
                            continue
                        
                        p = content_elem.find('p')
                        if not p:
                            continue
                        
                        # Extraer el enlace "Leer más..."
                        a_link = p.find('a', href=True)
                        if a_link:
                            link_novedad = a_link.get('href', '')
                            descripcion = p.get_text(strip=True).replace(a_link.get_text(strip=True), '').strip()
                        else:
                            link_novedad = boletin['link']
                            descripcion = p.get_text(strip=True)
                        
                        # Limpiar título
                        titulo = re.sub(r'\s+', ' ', titulo).strip()
                        
                        # Solo agregar si el título es significativo
                        if titulo and len(titulo) > 10:
                            rows.append({
                                'Date': boletin['fecha'],
                                'Title': titulo,
                                'Link': link_novedad if link_novedad else boletin['link'],
                                'Organismo': "CEMLA"
                            })
                            print(f"    ✅ {titulo[:60]}...")
                    
                    except Exception as e:
                        print(f"    ⚠️ Error procesando novedad: {e}")
                        continue
                        
            except Exception as e:
                print(f"  ❌ Error procesando boletín: {e}")
                continue

    except Exception as e:
        print(f"❌ Error general: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.drop_duplicates(subset=['Link'], keep='first')
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ CEMLA PUBLICACIONES - Total novedades: {len(df)} documentos")

    return df

# -- G20 --
@st.cache_data(show_spinner=False)
def load_pub_inst_g20(start_date_str, end_date_str):
    """Extrae documentos del G20 desde la página de News and Media"""
    import requests
    from bs4 import BeautifulSoup
    import datetime
    import re
    import pandas as pd

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 G20: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")

    url = "https://g20.org/media/"
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
    
    # Palabras clave que queremos incluir
    keywords_incluir = [
        'chair summary', 'declarations', 'g-20 note', 'presidency note',
        'chair\'s summary', 'chair summary', 'g20 note', 'presidency note'
    ]
    
    # Palabras clave para excluir
    keywords_excluir = [
        'agriculture', 'cultura', 'cultural', 'food security', 
        'farming', 'rural', 'agri'
    ]
    
    rows = []

    try:
        print(f"📡 Solicitando página: {url}")
        res = requests.get(url, headers=headers, timeout=15)
        
        if res.status_code != 200:
            print(f"❌ Error al acceder a la página: {res.status_code}")
            return pd.DataFrame()

        soup = BeautifulSoup(res.text, 'html.parser')
        
        # Buscar la sección de Press Releases
        press_section = None
        for section in soup.find_all('section', class_='paragraphsection'):
            toptitle = section.find('h2', class_='toptitle')
            if toptitle and 'Press Releases' in toptitle.get_text():
                press_section = section
                break
        
        if not press_section:
            print("⚠️ No se encontró la sección de Press Releases")
            return pd.DataFrame()
        
        # Buscar todos los artículos (h2 seguido de p con fecha)
        articles = press_section.find_all(['h2', 'p'])
        
        i = 0
        while i < len(articles) - 1:
            if articles[i].name == 'h2':
                h2 = articles[i]
                a_tag = h2.find('a')
                
                if a_tag and a_tag.get('href'):
                    titulo = a_tag.get_text(strip=True)
                    link = a_tag.get('href')
                    
                    if not titulo:
                        i += 1
                        continue
                    
                    if i + 1 < len(articles) and articles[i + 1].name == 'p':
                        p_text = articles[i + 1].get_text(strip=True)
                        
                        match = re.search(r'([A-Za-z]+ \d{1,2},? \d{4})', p_text)
                        if match:
                            fecha_str = match.group(1)
                            try:
                                fecha_str = fecha_str.replace(',', '')
                                parsed_date = datetime.datetime.strptime(fecha_str, '%B %d %Y')
                            except:
                                try:
                                    parsed_date = datetime.datetime.strptime(fecha_str, '%b %d %Y')
                                except:
                                    parsed_date = None
                            
                            if parsed_date:
                                if parsed_date < start_date or parsed_date > end_date:
                                    i += 2
                                    continue
                                
                                titulo_lower = titulo.lower()
                                incluir = any(kw in titulo_lower for kw in keywords_incluir)
                                excluir = any(kw in titulo_lower for kw in keywords_excluir)
                                
                                if excluir or not incluir:
                                    i += 2
                                    continue
                                
                                if link.startswith('/'):
                                    link = f"https://g20.org{link}"
                                
                                rows.append({
                                    "Date": parsed_date,
                                    "Title": titulo,
                                    "Link": link,
                                    "Organismo": "G20"
                                })
                                print(f"   ✅ Agregado: {titulo[:60]}... ({parsed_date.date()})")
                    else:
                        print(f"   ⚠️ No hay párrafo después del h2")
                else:
                    print(f"   ⚠️ h2 sin enlace válido")
            
            i += 1
            
    except Exception as e:
        print(f"❌ Error general: {e}")
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.drop_duplicates(subset=['Link'])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ TOTAL G20: {len(df)} documentos")
    else:
        print("⚠️ No se encontraron documentos del G20")

    return df

# -- CEF -- #

@st.cache_data(show_spinner=False)
def load_pub_inst_cef(start_date_str, end_date_str):
    url = "https://www.fsb.org/publications/key-regular-publications/"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
    rows = []
    try:
        res = requests.get(url, headers=headers, timeout=15)
        soup = BeautifulSoup(res.text, 'html.parser')
        for section in soup.find_all('div', class_='wp-bootstrap-blocks-row'):
            h2 = section.find('h2')
            if not h2:
                continue
            base_title = h2.get_text(strip=True)
            # Latest
            latest_btn = section.find('button', class_='btn-primary')
            if latest_btn and latest_btn.find('a'):
                a_tag = latest_btn.find('a')
                link = "https://www.fsb.org" + \
                    a_tag['href'] if a_tag['href'].startswith(
                        '/') else a_tag['href']
                date_match = re.search(r'\((.*?)\)', a_tag.get_text())
                parsed_date = parser.parse(
                    date_match.group(1)) if date_match else None
                if parsed_date and parsed_date >= start_date:
                    rows.append(
                        {"Date": parsed_date, "Title": f"{base_title}: Latest Report", "Link": link, "Organismo": "CEF"})
            # Previous
            dropdown = section.find('div', class_='dropdown-menu')
            if dropdown:
                for l in dropdown.find_all('a'):
                    year_text = l.get_text(strip=True)
                    try:
                        parsed_date = datetime.datetime(int(year_text), 1, 1)
                    except:
                        parsed_date = None
                    if parsed_date and parsed_date >= start_date:
                        rows.append(
                            {"Date": parsed_date, "Title": f"{base_title} ({year_text})", "Link": l['href'], "Organismo": "CEF"})
    except:
        pass
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

# ========== FUNCIÓN UNIVERSAL PARA NOTICIAS DEL FMI (API COVEO) ==========
@st.cache_data(show_spinner=False)
def load_fmi_news_all(start_date_str, end_date_str):
    """
    Extrae TODAS las noticias del FMI usando la API de Coveo.
    Incluye Press Releases, Mission Concluding, Statements, News Articles, etc.
    """
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 FMI News (API Coveo): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")

    rows = []
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"

    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Origin": "https://www.imf.org",
        "Referer": "https://www.imf.org/",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    # Filtro para capturar TODO excepto discursos
    payload = {
        "aq": "@imftype==(\"News Article\",\"Press Release\",\"Communique\",\"Mission Concluding Statement\",\"News Brief\",\"Public Information Notice\",\"Statements at Donor Meeting\",\"Views and Commentaries\",\"Blog Page\",\"IMF Staff Country Reports\") AND NOT @imftype==(\"Speech\",\"Transcript\") AND @syslanguage==\"English\"",
        "numberOfResults": 300,
        "sortCriteria": "@imfdate descending"
    }

    try:
        print("📡 Solicitando noticias del FMI a la API de Coveo...")
        # Deshabilitar verificación SSL para evitar errores locales
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        res = requests.post(url, headers=headers, json=payload, timeout=15, verify=False)

        if res.status_code == 200:
            data = res.json()
            print(f"✅ Respuesta recibida. Total en API: {data.get('totalCount', 0)} resultados")

            for item in data.get("results", []):
                titulo = item.get("title", "").strip()
                link = item.get("clickUri", "")
                content_type = item.get("raw", {}).get("imftype", "Unknown")

                raw_date = item.get("raw", {}).get("date")
                parsed_date = None
                if raw_date:
                    try:
                        parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                    except:
                        pass

                if not titulo or not link or not parsed_date:
                    continue

                if start_date <= parsed_date <= end_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo,
                            "Link": link,
                            "Organismo": "FMI"
                        })
                        print(f"   ✅ [{content_type[:25]}] {parsed_date.strftime('%Y-%m-%d')}: {titulo[:60]}...")
        else:
            print(f"❌ Error en la API: {res.status_code}")
            print(f"   Respuesta: {res.text[:200]}")

    except Exception as e:
        print(f"❌ Error: {e}")

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ TOTAL FMI News (API): {len(df)} documentos")
    else:
        print("⚠️ No se encontraron documentos")

    return df


# -- BPI -- Publicaciones Institucionales 

@st.cache_data(show_spinner=False)
def load_pub_inst_bpi(start_date_str, end_date_str):
    urls_api = ["https://www.bis.org/api/document_lists/annualeconomicreports.json",
                "https://www.bis.org/api/document_lists/quarterlyreviews.json"]
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
    rows = []
    for url in urls_api:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            data = res.json()
            for path, doc in data.get("list", {}).items():
                titulo = html.unescape(doc.get("short_title", ""))
                link = "https://www.bis.org" + doc.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"):
                    link += ".htm"
                try:
                    parsed_date = parser.parse(
                        doc.get("publication_start_date", ""))
                except:
                    continue
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo,
                                "Link": link, "Organismo": "BPI"})
        except:
            continue
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_country_reports_fmi(start_date_str, end_date_str):
    """Extractor FMI - Country Reports + Article IV (Conexión Directa a Coveo API) con filtro anti-Coming Soon"""
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 FMI Country Reports: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now() + datetime.timedelta(days=365)
        print(f"⚠️ Error en fechas, usando rango por defecto")

    rows = []

    # 1. EL ENDPOINT Y LA LLAVE MAESTRA QUE DESCUBRISTE
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"

    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    # 2. EL PAYLOAD (Falsificamos la petición del buscador)
    payload = {
        "aq": "@imfseries==\"IMF Staff Country Reports\" OR @imftype==\"Article IV Staff Reports\"",  # Filtro estricto por la Serie
        "numberOfResults": 150,  # Cantidad a traer (Suficiente para un mes)
        "sortCriteria": "@imfdate descending"  # Los más recientes primero
    }

    try:
        print("📡 Solicitando Country Reports + Article IV a la API de Coveo...")
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        # Hacemos un POST directo a la base de datos de Coveo
        res = requests.post(url, headers=headers, json=payload, timeout=15, verify=False)


        if res.status_code == 200:
            data = res.json()
            print(f"✅ Respuesta recibida. Total en API: {data.get('totalCount', 0)} resultados")

            # 3. EXTRACCIÓN (Limpia y sin HTML)
            for item in data.get("results", []):
                titulo = item.get("title", "")
                link = item.get("clickUri", "")

                # La fecha viene en timestamp (milisegundos). Lo dividimos entre 1000 para segundos.
                raw_date = item.get("raw", {}).get("date")

                # 🚫 FILTRO: Saltar documentos "Coming Soon"
                if "coming soon" in titulo.lower():
                    print(f"   ⏭️ Country Reports - Excluido 'Coming Soon': {titulo[:60]}...")
                    continue  # Salta este documento y pasa al siguiente

                # Validación básica
                if not titulo or not link or not raw_date:
                    continue

                # Parsear fecha
                try:
                    parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                except:
                    continue

                # ✅ FILTRO DE FECHAS MEJORADO (rango completo)
                if parsed_date < start_date or parsed_date > end_date:
                    continue

                # Evitar duplicados
                if not any(r['Link'] == link for r in rows):
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "FMI"
                    })
                    
            print(f"✅ Total de documentos filtrados: {len(rows)}")
        else:
            print(f"❌ Error en la API: {res.status_code}")

    except Exception as e:
        print(f"❌ Error en load_country_reports_fmi: {e}")

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_press_releases_fmi(start_date_str, end_date_str):
    """Extractor FMI - Press Releases (Historial completo vía Coveo API)"""
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        print(f"📅 PRENSA - Rango de fechas: {start_date.date()} a {end_date_str}")
    except:
        start_date = datetime.datetime(2000, 1, 1)

    rows = []
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"
    
    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Origin": "https://www.imf.org",
        "Referer": "https://www.imf.org/",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    payload = {
        "aq": "@imftype==\"Press Release\" AND @syslanguage==\"English\"",
        "numberOfResults": 150,
        "sortCriteria": "@imfdate descending"
    }

    try:
        # 🔧 VERIFY=FALSE es la clave para evitar el error SSL
        res = requests.post(url, headers=headers, json=payload, timeout=15, verify=False)
        
        if res.status_code == 200:
            data = res.json()
            total_raw = data.get('totalCount', 0)
            print(f"📡 PRENSA - Total resultados de la API: {total_raw}")
            
            for item in data.get("results", []):
                titulo = item.get("title", "")
                link = item.get("clickUri", "")
                raw_date = item.get("raw", {}).get("date")
                
                # 🚫 FILTRO: Saltar documentos "Coming Soon"
                if "coming soon" in titulo.lower():
                    print(f"   ⏭️ PRENSA - Excluido 'Coming Soon': {titulo[:60]}...")
                    continue
                
                parsed_date = None
                if raw_date:
                    try:
                        parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                    except:
                        pass

                if not titulo or not link or not parsed_date:
                    continue

                if parsed_date >= start_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "FMI"})
                        print(f"   ✅ PRENSA - Agregado: {parsed_date.strftime('%Y-%m-%d')} - {titulo[:60]}...")
            
            print(f"📊 PRENSA - Total documentos filtrados: {len(rows)}")
        else:
            print(f"❌ PRENSA - Error en API: {res.status_code}")
    except Exception as e:
        print(f"❌ PRENSA - Excepción: {e}")

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_country_reports_elibrary(start_date_str, end_date_str):
    """Extractor FMI - Country Reports (Bypass de Tapestry 5 AJAX Lazy-Loading)"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8'
    }

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)

    rows = []
    base_domain = "https://www.elibrary.imf.org"
    url_overview = f"{base_domain}/view/journals/002/002-overview.xml"

    try:
        # FASE 1: Extraer los tokens dinámicos de AJAX para los años recientes
        res = requests.get(url_overview, headers=headers, timeout=15)
        if res.status_code != 200:
            return pd.DataFrame()

        soup = BeautifulSoup(res.text, 'html.parser')

        ajax_links = []
        current_year = datetime.datetime.now().year
        # Buscamos los enlaces de expansión para el año actual y el anterior
        target_years = [str(current_year), str(current_year - 1)]

        for li in soup.find_all('div', attrs={'data-toc-role': 'li'}):
            label_div = li.find('div', class_='label')
            if not label_div:
                continue

            texto_label = label_div.get_text()
            if any(year in texto_label for year in target_years):
                a_tag = li.find('a', class_='ajax-control')
                if a_tag and a_tag.has_attr('href'):
                    ajax_links.append(base_domain + a_tag['href'])

        # FASE 2: Interceptar y "deshidratar" las respuestas AJAX de Tapestry
        headers_ajax = headers.copy()
        # Engañamos al framework
        headers_ajax['X-Requested-With'] = 'XMLHttpRequest'
        headers_ajax['Accept'] = 'application/json, text/javascript, */*; q=0.01'

        for ajax_url in ajax_links:
            try:
                res_ajax = requests.get(
                    ajax_url, headers=headers_ajax, timeout=15)
                if res_ajax.status_code != 200:
                    continue

                data = res_ajax.json()

                # Extraemos el HTML inyectado dentro del nodo "zones"
                html_fragment = ""
                if "zones" in data:
                    for zone_id, html_content in data["zones"].items():
                        html_fragment += html_content

                if not html_fragment:
                    continue

                # FASE 3: Parsear el HTML revelado
                soup_fragment = BeautifulSoup(html_fragment, 'html.parser')

                for a_tag in soup_fragment.find_all('a', href=True):
                    href = a_tag['href']
                    titulo = a_tag.get_text(strip=True)

                    # Filtro de sanidad: debe ser un artículo real
                    if '/view/journals/002/' in href and len(titulo) > 15:
                        link_real = base_domain + \
                            href if href.startswith('/') else href

                        # Buscamos la fecha subiendo hasta 3 niveles en el DOM
                        date_str = ""
                        for padre in a_tag.find_parents(['div', 'li'], limit=3):
                            texto_padre = padre.get_text(
                                separator=" ", strip=True)

                            # Caza fechas en formatos "Mar 05, 2026" o "05 March 2026"
                            match = re.search(
                                r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2}?,?\s*\d{4}', texto_padre)
                            if not match:
                                match = re.search(
                                    r'\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}', texto_padre)

                            if match:
                                date_str = match.group(0)
                                break  # Encontramos la fecha, salimos del bucle

                        parsed_date = None
                        if date_str:
                            try:
                                parsed_date = parser.parse(date_str)
                                if parsed_date.tzinfo is not None:
                                    parsed_date = parsed_date.replace(
                                        tzinfo=None)
                            except:
                                pass

                        # Evaluación final
                        if parsed_date and parsed_date >= start_date:
                            if not any(r['Link'] == link_real for r in rows):
                                rows.append(
                                    {"Date": parsed_date, "Title": titulo, "Link": link_real, "Organismo": "FMI"})
            except:
                continue  # Aislamiento de fallos

    except Exception as e:
        pass

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

## FMI - Publiccaciones Institucionales - INICIO

## FMI - F&D Magazine (CORREGIDO)
## FMI - F&D Magazine (VERSIÓN CLOUDSCRAPER - SIN SELENIUM)
@st.cache_data(show_spinner=False)
## FMI - F&D Magazine (VERSIÓN CON API DIRECTA)
@st.cache_data(show_spinner=False)
def load_pub_inst_fandd(start_date_str, end_date_str):
    """
    Extrae ediciones completas de la revista F&D Magazine del FMI
    Usa la API de Next.js directamente para evitar bloqueos
    """
    import requests
    import json
    import re
    import datetime
    import pandas as pd
    from dateutil import parser
    import time

    print("="*50)
    print("📘 CARGANDO F&D MAGAZINE (API directa)")
    print(f"   Fechas: {start_date_str} a {end_date_str}")
    print("="*50)

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"   Rango parseado: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"   ⚠️ Error en fechas, usando rango por defecto")

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'application/json, text/plain, */*',
        'Accept-Language': 'en-US,en;q=0.9',
        'Referer': 'https://www.imf.org/en/publications/fandd/issues',
        'Origin': 'https://www.imf.org',
    }

    rows = []

    try:
        # ========== PASO 1: Obtener el build ID ==========
        print("   📡 Obteniendo build ID...")
        res = requests.get("https://www.imf.org/en/publications/fandd/issues", headers=headers, timeout=15)
        
        if res.status_code != 200:
            print(f"   ❌ Error al acceder: {res.status_code}")
            return pd.DataFrame()
        
        # Extraer build ID del HTML
        match = re.search(r'"buildId":"([^"]+)"', res.text)
        if not match:
            print("   ❌ No se encontró build ID")
            return pd.DataFrame()
        
        build_id = match.group(1)
        print(f"   ✅ Build ID: {build_id}")
        
        # ========== PASO 2: Llamar a la API de Next.js ==========
        url = f"https://www.imf.org/_next/data/{build_id}/en/publications/fandd/issues.json"
        print(f"   📡 Solicitando API: {url}")
        
        res = requests.get(url, headers=headers, timeout=15)
        
        if res.status_code != 200:
            print(f"   ❌ Error en API: {res.status_code}")
            return pd.DataFrame()
        
        data = res.json()
        print("   ✅ JSON recibido")
        
        # ========== PASO 3: Extraer los issues ==========
        results = []
        try:
            page_props = data.get('pageProps', {}).get('props', {})
            component_props = page_props.get('componentProps', {})
            
            for comp_id, comp_data in component_props.items():
                if isinstance(comp_data, dict) and 'issueList' in comp_data:
                    issue_list = comp_data['issueList']
                    if isinstance(issue_list, dict) and 'results' in issue_list:
                        results = issue_list['results']
                        print(f"   ✅ Encontrados {len(results)} issues")
                        break
        except Exception as e:
            print(f"   ⚠️ Error navegando en el JSON: {e}")
        
        if not results:
            print("   ❌ No se encontraron issues")
            return pd.DataFrame()

        meses_map = {
            'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
            'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12,
            'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
            'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
        }

        for issue in results:
            issue_title = issue.get('issueTitle', {}).get('jsonValue', {}).get('value', '')
            issue_label = issue.get('issueLabel', {}).get('jsonValue', {}).get('value', '')
            
            # ========== PRIORIZAR redirectUrl ==========
            redirect_url = issue.get('redirectUrl', {}).get('jsonValue', {}).get('value', {})
            if isinstance(redirect_url, dict):
                issue_url = redirect_url.get('href', '')
            else:
                issue_url = redirect_url
            
            # Si no hay redirectUrl, usar el url de la página
            if not issue_url:
                issue_url = issue.get('url', {}).get('url', '')
                if not issue_url and issue.get('url', {}).get('path'):
                    issue_url = "https://www.imf.org" + issue.get('url', {}).get('path', '')
            
            fecha_texto = issue_label if issue_label else issue_title
            
            match_date = re.search(r'([A-Za-z]+)\s+(\d{4})', fecha_texto, re.IGNORECASE)
            if not match_date:
                print(f"   ⚠️ No se pudo parsear fecha: '{fecha_texto}'")
                continue
            
            mes_str = match_date.group(1).lower()
            año = int(match_date.group(2))
            mes_num = meses_map.get(mes_str, 1)
            
            issue_date = datetime.datetime(año, mes_num, 15)
            
            if issue_date < start_date or issue_date > end_date:
                print(f"   ⏭️ Fuera de rango: {issue_date.strftime('%Y-%m')} - {issue_title}")
                continue
            
            title_clean = re.sub(r'\s+', ' ', issue_title).strip()
            if not title_clean:
                title_clean = fecha_texto
            
            titulo_final = f"F&D: {issue_label} - {title_clean}" if issue_label else f"F&D: {title_clean}"
            
            print(f"   📎 {issue_date.strftime('%Y-%m')}: PDF = {issue_url[:80]}...")
            
            rows.append({
                "Date": issue_date,
                "Title": titulo_final,
                "Link": issue_url if issue_url else f"https://www.imf.org/en/publications/fandd/issues/{año}/{mes_num:02d}",
                "Organismo": "FMI"
            })
            print(f"   ✅ AGREGADO: {issue_date.strftime('%Y-%m-%d')} - {titulo_final[:60]}...")
        
    except Exception as e:
        print(f"   ❌ Error general: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.drop_duplicates(subset=['Link'])
        df = df.sort_values("Date", ascending=False)
    
    print(f"   📊 TOTAL F&D: {len(df)} ediciones")
    print("="*50)
    return df

## FMI - 

@st.cache_data(show_spinner=False)
def load_pub_inst_fmi(start_date_str, end_date_str):
    """Extractor FMI - Vía directa por API Next.js (El Regalo) + filtro anti-Coming Soon"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'application/json, text/plain, */*'
    }

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)

    rows = []

    # 1. CAZADOR DE BUILD ID (Para que tu código no caduque nunca)
    build_id = "OPXKbpp2La91iW-gTVkBX"  # Tu regalo como plan de respaldo
    try:
        res_html = requests.get(
            "https://www.imf.org/en/publications", headers=headers, timeout=15)
        # Buscamos el código dinámico oculto en la página principal
        match = re.search(r'"buildId":"([^"]+)"', res_html.text)
        if match:
            build_id = match.group(1)
    except:
        pass

    # 2. CONSTRUCCIÓN DE LOS ENLACES JSON DIRECTOS
    endpoints_json = [
        f"https://www.imf.org/_next/data/{build_id}/en/publications/fm.json",
        f"https://www.imf.org/_next/data/{build_id}/en/publications/weo.json",
        f"https://www.imf.org/_next/data/{build_id}/en/publications/gfsr.json"
    ]

    for url in endpoints_json:
        try:
            # Ahora pedimos el JSON limpio, evadiendo el HTML
            res = requests.get(url, headers=headers, timeout=15)
            if res.status_code != 200:
                continue
            data = res.json()

            # Buscador recursivo dentro del JSON
            def extraer_issues(obj):
                if isinstance(obj, dict):
                    if "issuePage" in obj and isinstance(obj["issuePage"], dict) and "results" in obj["issuePage"]:
                        for r in obj["issuePage"]["results"]:
                            yield r
                    for k, v in obj.items():
                        yield from extraer_issues(v)
                elif isinstance(obj, list):
                    for item in obj:
                        yield from extraer_issues(item)

            for issue in extraer_issues(data):
                titulo = issue.get("title", {}).get(
                    "jsonValue", {}).get("value", "")
                link_raw = issue.get("url", {}).get(
                    "url", "") or issue.get("url", {}).get("path", "")
                if not titulo or not link_raw:
                    continue

                # 🚫 FILTRO CRÍTICO: Excluir documentos "Coming Soon"
                if "coming soon" in titulo.lower():
                    print(f"   ⏭️ Excluido por 'Coming Soon': {titulo[:60]}...")
                    continue

                link_real = link_raw if link_raw.startswith(
                    "http") else "https://www.imf.org" + link_raw

                d_str = issue.get("publicationDate", {}).get(
                    "jsonValue", {}).get("value", "")
                if d_str:
                    try:
                        parsed_date = parser.parse(d_str)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                        if parsed_date >= start_date and not any(r['Link'] == link_real for r in rows):
                            rows.append(
                                {"Date": parsed_date, "Title": titulo, "Link": link_real, "Organismo": "FMI"})
                    except:
                        pass
        except:
            continue

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_pub_inst_bm(start_date_str, end_date_str):
    """Extractor para Publicaciones Institucionales (Colecciones Específicas) del BM"""
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}

    # IDs exactos de las 3 colecciones
    scopes = [
        '4c48a649-7773-4d0f-b441-f5fc7e8d67f8',  # Business Ready
        '09c5e8fc-187f-5c2f-a077-3e03044c7b62',  # Perspectivas económicas mundiales
        '3d9bbbf6-c007-5043-b655-04d8a1cfbfb2'  # Tercera colección
    ]

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)

    rows = []

    # Iteramos sobre cada una de las colecciones
    for scope in scopes:
        page = 0
        while True:
            try:
                # Al pasarle el 'scope', la API restringe la búsqueda SOLO a esa colección
                params = {
                    'scope': scope,
                    'sort': 'dc.date.issued,DESC',
                    'page': page,
                    'size': 20
                }
                res = requests.get(base_url, headers=headers,
                                   params=params, timeout=15)
                data = res.json()

                objects = data.get('_embedded', {}).get(
                    'searchResult', {}).get('_embedded', {}).get('objects', [])
                if not objects:
                    break

                items_found = 0
                for obj in objects:
                    item = obj.get('_embedded', {}).get('indexableObject', {})
                    meta = item.get('metadata', {})

                    # Extraer Título
                    title = meta.get('dc.title', [{'value': ''}])[
                        0].get('value', '')
                    date_s = meta.get('dc.date.issued', [{'value': ''}])[
                        0].get('value', '')

                    parsed_date = None
                    if date_s:
                        try:
                            parsed_date = parser.parse(date_s)
                        except:
                            pass

                    if not parsed_date or parsed_date < start_date:
                        continue

                    link = meta.get('dc.identifier.uri', [{'value': ''}])[
                        0].get('value', '')
                    if not link:
                        link = f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"

                    if not any(r['Link'] == link for r in rows):
                        rows.append(
                            {"Date": parsed_date, "Title": title, "Link": link, "Organismo": "BM"})
                        items_found += 1

                if items_found == 0:
                    break
                page += 1
                if page > 3:
                    break  # Límite de seguridad
                time.sleep(0.2)
            except:
                break

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None:
            df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df

    # --- SECCIÓN: INVESTIGACIÓN ---

## - Working Papers - FMI
@st.cache_data(show_spinner=False)
def load_working_papers_fmi(start_date_str, end_date_str):
    """
    Extractor FMI - Working Papers usando la API de búsqueda del FMI (títulos completos)
    """
    import requests
    import datetime
    import re
    from dateutil import parser
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 FMI Working Papers (API de búsqueda FMI): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'application/json, text/plain, */*',
        'Accept-Language': 'en-US,en;q=0.9',
        'Referer': 'https://www.imf.org/en/publications',
        'Origin': 'https://www.imf.org',
    }
    
    # ========== USAR LA API DE BÚSQUEDA DEL FMI ==========
    # Esta API devuelve los títulos completos de los working papers
    # Nota: Solo funciona para publicaciones de los últimos meses
    
    # Construir la URL con los parámetros de búsqueda
    # Filtrar por fecha y tipo de publicación
    date_from = start_date.strftime('%Y-%m-%d')
    date_to = end_date.strftime('%Y-%m-%d')
    
    # API de búsqueda del FMI
    search_url = "https://www.imf.org/api/search/publications"
    
    params = {
        "locale": "en",
        "type": "WRKNGPPRS",  # Working Papers
        "dateFrom": date_from,
        "dateTo": date_to,
        "pageSize": 200,
        "sortBy": "date",
        "sortOrder": "desc"
    }
    
    try:
        print(f"📡 Solicitando Working Papers a la API de búsqueda del FMI...")
        response = requests.get(search_url, headers=headers, params=params, timeout=30, verify=False)
        
        if response.status_code == 200:
            data = response.json()
            items = data.get('results', []) or data.get('items', []) or data.get('publications', [])
            
            if not items:
                # Intentar otra estructura de la API
                items = data.get('data', []) or data.get('list', [])
            
            print(f"📚 Documentos encontrados en la API: {len(items)}")
            
            for item in items:
                # Extraer título
                titulo = item.get('title', '')
                if not titulo:
                    titulo = item.get('name', '')
                if not titulo:
                    titulo = item.get('publicationTitle', '')
                if not titulo:
                    titulo = item.get('heading', '')
                
                # Extraer URL
                link = item.get('url', '')
                if not link:
                    link = item.get('link', '')
                if not link:
                    link = item.get('publicationUrl', '')
                if link and link.startswith('/'):
                    link = "https://www.imf.org" + link
                elif link and not link.startswith('http'):
                    link = f"https://www.imf.org/en/Publications/WP/Issues/{link}"
                
                # Extraer DOI
                doi = item.get('doi', '')
                if not doi:
                    doi = item.get('identifier', '')
                if doi and not link:
                    link = f"https://doi.org/{doi}"
                
                # Extraer fecha
                fecha_texto = item.get('publicationDate', '')
                if not fecha_texto:
                    fecha_texto = item.get('date', '')
                if not fecha_texto:
                    fecha_texto = item.get('issued', '')
                
                parsed_date = None
                if fecha_texto:
                    try:
                        parsed_date = parser.parse(fecha_texto)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        pass
                
                if not parsed_date:
                    # Intentar extraer fecha de la URL
                    if link:
                        match = re.search(r'/(\d{4})/(\d{2})/(\d{2})', link)
                        if match:
                            try:
                                parsed_date = datetime.datetime(int(match.group(1)), int(match.group(2)), int(match.group(3)))
                            except:
                                pass
                
                if not titulo or not link or not parsed_date:
                    continue
                
                if start_date <= parsed_date <= end_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo,
                            "Link": link,
                            "Organismo": "FMI"
                        })
                        print(f"   ✅ {parsed_date.strftime('%Y-%m-%d')}: {titulo[:80]}...")
                        
        else:
            print(f"⚠️ Error en API de búsqueda FMI: {response.status_code}")
            print("   Intentando usar Crossref como fallback...")
            
            # ========== FALLBACK: Usar Crossref ==========
            return load_working_papers_fmi_crossref(start_date_str, end_date_str)
            
    except Exception as e:
        print(f"❌ Error en load_working_papers_fmi: {e}")
        import traceback
        traceback.print_exc()
        print("   Intentando usar Crossref como fallback...")
        return load_working_papers_fmi_crossref(start_date_str, end_date_str)
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"\n📊 FMI Working Papers - Total final: {len(df)}")
    return df


def load_working_papers_fmi_crossref(start_date_str, end_date_str):
    """
    Fallback: Usar Crossref API para Working Papers del FMI
    """
    import requests
    import datetime
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    url = "https://api.crossref.org/works"
    
    params = {
        "filter": f"from-pub-date:{start_date.strftime('%Y-%m-%d')},until-pub-date:{end_date.strftime('%Y-%m-%d')},prefix:10.5089",
        "rows": 200,
        "sort": "published-online",
        "order": "desc"
    }
    
    try:
        print("   📡 Usando Crossref API como fallback...")
        response = requests.get(url, headers={'User-Agent': 'Mozilla/5.0'}, params=params, timeout=30, verify=False)
        
        if response.status_code == 200:
            data = response.json()
            items = data.get('message', {}).get('items', [])
            
            for item in items:
                titulo = item.get('title', [''])[0] if item.get('title') else ''
                doi = item.get('DOI', '')
                if not doi or not titulo:
                    continue
                
                link = f"https://doi.org/{doi}"
                
                pub_date = item.get('published-print', {}) or item.get('published-online', {})
                date_parts = pub_date.get('date-parts', [[]])[0]
                parsed_date = None
                if len(date_parts) >= 3:
                    try:
                        parsed_date = datetime.datetime(date_parts[0], date_parts[1], date_parts[2])
                    except:
                        pass
                elif len(date_parts) >= 2:
                    try:
                        parsed_date = datetime.datetime(date_parts[0], date_parts[1], 1)
                    except:
                        pass
                
                if not parsed_date or start_date > parsed_date or end_date < parsed_date:
                    continue
                
                if 'working paper' not in titulo.lower() and 'working paper' not in str(item.get('container-title', [''])[0]).lower():
                    continue
                
                if "coming soon" in titulo.lower():
                    continue
                
                if not any(r['Link'] == link for r in rows):
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "FMI"
                    })
    except Exception as e:
        print(f"   ❌ Error en fallback de Crossref: {e}")
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    return df


# --- SECCIÓN: INVESTIGACIÓN --- 
@st.cache_data(show_spinner=False)
def load_investigacion_bpi(start_date_str, end_date_str):
    """Extractor Investigación BPI (BIS Papers & Working Papers) vía API JSON"""
    import requests
    import pandas as pd
    import datetime
    import html
    from dateutil import parser
    
    # Configuración de fechas
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime.now() - datetime.timedelta(days=365)
        end_date = datetime.datetime.now()

    # Los dos endpoints JSON de investigación del BPI
    urls = [
        "https://www.bis.org/api/document_lists/bispapers.json",
        "https://www.bis.org/api/document_lists/wppubls.json"
    ]
    
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []

    for url in urls:
        try:
            response = requests.get(url, headers=headers, timeout=10)
            if response.status_code == 200:
                data = response.json()
                
                # Usamos la lógica de extracción exacta de tu función de discursos
                for path, doc in data.get("list", {}).items():
                    title = html.unescape(doc.get("short_title", ""))
                    date_str = doc.get("publication_start_date", "")
                    
                    link = "https://www.bis.org" + path + (".htm" if not path.endswith(".htm") else "")
                    
                    if title and date_str:
                        try:
                            # Convertimos a datetime para aplicar el filtro de tu app
                            p_date = parser.parse(date_str).replace(tzinfo=None)
                            
                            if start_date <= p_date <= end_date:
                                rows.append({
                                    "Date": p_date,
                                    "Title": title,
                                    "Link": link,
                                    "Organismo": "BPI"
                                })
                        except:
                            continue
        except Exception as e:
            print(f"Error BPI Investigación en {url}: {e}")

    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df = df.sort_values("Date", ascending=False)
        
    return df

## BID - Inglés 
@st.cache_data(show_spinner=False)
def load_investigacion_bid_en(start_date_str, end_date_str):
    """
    Extrae Working Papers del BID en inglés usando undetected-chromedriver
    """
    import undetected_chromedriver as uc
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from bs4 import BeautifulSoup
    import datetime
    import time
    import re
    import ssl
    import urllib3
    
    # 🔧 SOLUCIÓN PARA REDES CORPORATIVAS
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    ssl._create_default_https_context = ssl._create_unverified_context
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BID Inglés: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")
    
    rows = []
    page = 0
    max_pages = 3
    
    options = uc.ChromeOptions()
    options.add_argument('--headless=new')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    options.add_argument('--window-size=1920,1080')
    options.add_argument('--disable-blink-features=AutomationControlled')
    options.add_argument('--disable-gpu')
    
    meses_en = {
        'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
        'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12,
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
    }
    
    try:
        print("🔍 Iniciando BID Inglés...")
        driver = uc.Chrome(options=options, version_main=146)
        time.sleep(2)
        
        while page < max_pages:
            url = f"https://publications.iadb.org/en?f%5B0%5D=type%3AWorking%20Papers&page={page}"
            print(f"📄 Página {page+1}: {url}")
            
            driver.get(url)
            time.sleep(10)
            
            if "Just a moment" in driver.page_source:
                print("   ⚠️ Cloudflare detectado, esperando...")
                time.sleep(15)
            
            try:
                WebDriverWait(driver, 45).until(
                    EC.presence_of_element_located((By.CLASS_NAME, "views-row"))
                )
            except:
                pass
            
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            items = soup.find_all('div', class_='views-row')
            
            if not items:
                print(f"   📭 No hay artículos en página {page+1}")
                break
            
            print(f"   📚 Artículos: {len(items)}")
            
            for item in items:
                try:
                    title_container = item.find('div', class_='views-field-field-title')
                    if title_container:
                        a_tag = title_container.find('a')
                        if a_tag:
                            titulo = a_tag.get_text(strip=True)
                            link = a_tag.get('href', '')
                            if link and not link.startswith('http'):
                                link = "https://publications.iadb.org" + link
                    else:
                        continue
                    
                    if not titulo or len(titulo) < 10:
                        continue
                    
                    date_container = item.find('div', class_='views-field-field-date-issued-text')
                    if date_container:
                        date_text = date_container.get_text(strip=True)
                        match = re.search(r'([A-Za-z]+)\s+(\d{4})', date_text)
                        if match:
                            mes_str = match.group(1).lower()
                            año = int(match.group(2))
                            mes_num = meses_en.get(mes_str, 1)
                            parsed_date = datetime.datetime(año, mes_num, 1)
                        else:
                            continue
                    else:
                        continue
                    
                    # Filtrar por año y mes
                    if parsed_date.year < start_date.year or parsed_date.year > end_date.year:
                        continue
                    if parsed_date.year == start_date.year and parsed_date.month < start_date.month:
                        continue
                    if parsed_date.year == end_date.year and parsed_date.month > end_date.month:
                        continue
                    
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo,
                            "Link": link,
                            "Organismo": "BID (Inglés)"
                        })
                        print(f"   ✅ {parsed_date.strftime('%Y-%m')}: {titulo[:50]}...")
                        
                except Exception as e:
                    print(f"   ⚠️ Error procesando artículo: {e}")
                    continue
            
            page += 1
            time.sleep(3)
        
        driver.quit()
        
    except Exception as e:
        print(f"❌ Error en BID Inglés: {e}")
        import traceback
        traceback.print_exc()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ BID Inglés: {len(df)} documentos")
    else:
        print("\n⚠️ No se encontraron documentos del BID (Inglés)")
    
    return df

## BID ESPAÑOL 

@st.cache_data(show_spinner=False)
def load_investigacion_bid(start_date_str, end_date_str):
    import undetected_chromedriver as uc
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from bs4 import BeautifulSoup
    import datetime
    import time
    import re
    import ssl
    import urllib3
    
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    ssl._create_default_https_context = ssl._create_unverified_context
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BID Español: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    page = 0
    max_pages = 3
    
    options = uc.ChromeOptions()
    options.add_argument('--headless=new')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    options.add_argument('--window-size=1920,1080')
    options.add_argument('--disable-blink-features=AutomationControlled')
    options.add_argument('--disable-gpu')
    
    meses_es = {
        'enero': 1, 'febrero': 2, 'marzo': 3, 'abril': 4, 'mayo': 5, 'junio': 6,
        'julio': 7, 'agosto': 8, 'septiembre': 9, 'octubre': 10, 'noviembre': 11, 'diciembre': 12,
    }
    
    try:
        print("🔍 Iniciando BID Español...")
        driver = uc.Chrome(options=options, version_main=146)
        time.sleep(2)
        
        while page < max_pages:
            url = f"https://publications.iadb.org/es?f%5B0%5D=type%3A4633&f%5B1%5D=type%3ADocumentos%20de%20Trabajo&page={page}"
            print(f"📄 Página {page+1}: {url}")
            
            driver.get(url)
            time.sleep(10)
            
            if "Just a moment" in driver.page_source:
                print("   ⚠️ Cloudflare detectado, esperando...")
                time.sleep(15)
            
            try:
                WebDriverWait(driver, 45).until(
                    EC.presence_of_element_located((By.CLASS_NAME, "views-row"))
                )
            except:
                pass
            
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            items = soup.find_all('div', class_='views-row')
            
            if not items:
                print(f"   📭 No hay artículos en página {page+1}")
                break
            
            print(f"   📚 Artículos: {len(items)}")
            
            for item in items:
                try:
                    title_container = item.find('div', class_='views-field-field-title')
                    if title_container:
                        a_tag = title_container.find('a')
                        if a_tag:
                            titulo = a_tag.get_text(strip=True)
                            link = a_tag.get('href', '')
                            if link and not link.startswith('http'):
                                link = "https://publications.iadb.org" + link
                    else:
                        continue
                    
                    if not titulo or len(titulo) < 10:
                        continue
                    
                    date_container = item.find('div', class_='views-field-field-date-issued-text')
                    if date_container:
                        date_text = date_container.get_text(strip=True)
                        match = re.search(r'([A-Za-z]+)\s+(\d{4})', date_text)
                        if match:
                            mes_str = match.group(1).lower()
                            año = int(match.group(2))
                            mes_num = meses_es.get(mes_str, 1)
                            parsed_date = datetime.datetime(año, mes_num, 1)
                        else:
                            continue
                    else:
                        continue
                    
                    if parsed_date.year < start_date.year or parsed_date.year > end_date.year:
                        continue
                    if parsed_date.year == start_date.year and parsed_date.month < start_date.month:
                        continue
                    if parsed_date.year == end_date.year and parsed_date.month > end_date.month:
                        continue
                    
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo,
                            "Link": link,
                            "Organismo": "BID"
                        })
                        print(f"   ✅ {parsed_date.strftime('%Y-%m')}: {titulo[:50]}...")
                        
                except Exception as e:
                    continue
            
            page += 1
            time.sleep(3)
        
        driver.quit()
        
    except Exception as e:
        print(f"❌ Error: {e}")
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ BID Español: {len(df)} documentos")
    
    return df

# ========== INVESTIGACIÓN CEMLA (Latin American Journal of Central Banking) ==========
@st.cache_data(show_spinner=False)
def load_investigacion_cemla(start_date_str, end_date_str):
    """
    Extractor CEMLA - Latin American Journal of Central Banking
    Extrae fecha COMPLETA (con día) si está disponible en Crossref
    """
    import requests
    import datetime
    from dateutil import parser
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    print("="*60)
    print("🔍 CEMLA INVESTIGACIÓN - Buscando fechas completas (con día)")
    print("="*60)
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 Rango: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    issn = "2666-1438"
    base_url = "https://api.crossref.org/works"
    
    # Buscar mes por mes para tener mejor control
    current = start_date.replace(day=1)
    
    while current <= end_date:
        year = current.year
        month = current.month
        
        # Último día del mes
        if month == 12:
            last_day = 31
        elif month in [4, 6, 9, 11]:
            last_day = 30
        else:
            last_day = 28 if year % 4 != 0 else 29
        
        fecha_inicio = f"{year}-{month:02d}-01"
        fecha_fin = f"{year}-{month:02d}-{last_day}"
        
        print(f"\n📆 Buscando {year}-{month:02d}...")
        
        params = {
            "filter": f"from-pub-date:{fecha_inicio},until-pub-date:{fecha_fin},issn:{issn}",
            "rows": 50,
            "sort": "published-online",
            "order": "desc"
        }
        
        try:
            response = requests.get(base_url, params=params, timeout=30, verify=False)
            
            if response.status_code == 200:
                data = response.json()
                items = data.get('message', {}).get('items', [])
                
                if items:
                    print(f"   📚 Artículos: {len(items)}")
                    
                    for item in items:
                        titulo = item.get('title', [''])[0] if item.get('title') else ''
                        doi = item.get('DOI', '')
                        link = f"https://doi.org/{doi}" if doi else ''
                        
                        if not titulo or not link:
                            continue
                        
                        # ========== INTENTAR OBTENER FECHA COMPLETA ==========
                        fecha_completa = None
                        
                        # 1. Probar con 'published-online' (puede tener día)
                        pub_online = item.get('published-online', {})
                        if pub_online:
                            date_parts = pub_online.get('date-parts', [[]])[0]
                            if len(date_parts) >= 3:
                                try:
                                    fecha_completa = datetime.datetime(date_parts[0], date_parts[1], date_parts[2])
                                    print(f"      📅 Online: {fecha_completa.strftime('%Y-%m-%d')}")
                                except:
                                    pass
                        
                        # 2. Probar con 'issued' (fecha de publicación)
                        if not fecha_completa:
                            issued = item.get('issued', {})
                            if issued:
                                date_parts = issued.get('date-parts', [[]])[0]
                                if len(date_parts) >= 3:
                                    try:
                                        fecha_completa = datetime.datetime(date_parts[0], date_parts[1], date_parts[2])
                                        print(f"      📅 Issued: {fecha_completa.strftime('%Y-%m-%d')}")
                                    except:
                                        pass
                        
                        # 3. Probar con 'posted-online'
                        if not fecha_completa:
                            posted = item.get('posted-online', {})
                            if posted:
                                date_parts = posted.get('date-parts', [[]])[0]
                                if len(date_parts) >= 3:
                                    try:
                                        fecha_completa = datetime.datetime(date_parts[0], date_parts[1], date_parts[2])
                                        print(f"      📅 Posted: {fecha_completa.strftime('%Y-%m-%d')}")
                                    except:
                                        pass
                        
                        # 4. Fallback: usar el primer día del mes (si no hay día)
                        if not fecha_completa:
                            fecha_completa = datetime.datetime(year, month, 1)
                            print(f"      ⚠️ Fallback: {fecha_completa.strftime('%Y-%m-%d')} (sin día específico)")
                        
                        # Filtrar por rango
                        if start_date <= fecha_completa <= end_date:
                            rows.append({
                                "Date": fecha_completa,
                                "Title": titulo,
                                "Link": link,
                                "Organismo": "CEMLA"
                            })
                            print(f"      ✅ AGREGADO: {fecha_completa.strftime('%Y-%m-%d')}")
                        else:
                            print(f"      ⏭️ Fuera de rango: {fecha_completa.strftime('%Y-%m-%d')}")
                            
                else:
                    print(f"   📭 Sin artículos")
                    
            else:
                print(f"   ❌ Error: {response.status_code}")
                
        except Exception as e:
            print(f"   ❌ Error: {e}")
        
        # Siguiente mes
        if current.month == 12:
            current = current.replace(year=current.year + 1, month=1)
        else:
            current = current.replace(month=current.month + 1)
        
        time.sleep(0.5)
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"\n{'='*60}")
    print(f"📊 CEMLA Investigación - Total: {len(df)} documentos")
    if not df.empty:
        print("\n📅 Primeros 5 documentos con sus fechas:")
        for i, row in df.head(5).iterrows():
            print(f"   {row['Date'].strftime('%Y-%m-%d')}: {row['Title'][:60]}...")
    print(f"{'='*60}")
    
    return df


@st.cache_data(show_spinner=False)
def load_investigacion_fmi(start_date_str, end_date_str):
    """Extractor FMI - Blogs de Investigación (Vía Coveo API) - Versión mejorada"""
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 FMI Blogs - Rango solicitado: {start_date.date()} a {end_date.date()}")
    except Exception as e:
        print(f"⚠️ Error en fechas: {e}")
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"📅 Usando rango por defecto: {start_date.date()} a {end_date.date()}")

    rows = []
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"
    
    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Origin": "https://www.imf.org",
        "Referer": "https://www.imf.org/",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    payload = {
        "aq": "@imftype==\"IMF Blog Page\" AND @syslanguage==\"English\"",
        "numberOfResults": 250,  # Aumentado para capturar más
        "sortCriteria": "@imfdate descending"
    }

    try:
        print("📡 Solicitando blogs a la API de Coveo...")
        res = requests.post(url, headers=headers, json=payload, timeout=15, verify=False)
        
        if res.status_code == 200:
            data = res.json()
            total_api = data.get('totalCount', 0)
            print(f"✅ Total de blogs en la API: {total_api}")
            
            documentos_filtrados = 0
            for item in data.get("results", []):
                titulo = item.get("title", "").strip()
                link = item.get("clickUri", "")
                
                # === MEJORA: Extraer fecha de múltiples formatos ===
                parsed_date = None
                raw_data = item.get("raw", {})
                
                # Formato 1: timestamp en milisegundos (el más común)
                raw_date = raw_data.get("date")
                if raw_date:
                    try:
                        parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                    except:
                        pass
                
                # Formato 2: fecha como string ISO
                if not parsed_date:
                    date_str = raw_data.get("date") or raw_data.get("publisheddate") or raw_data.get("publicationdate")
                    if date_str and isinstance(date_str, str):
                        try:
                            parsed_date = parser.parse(date_str)
                        except:
                            pass
                
                # Formato 3: intentar con cualquier campo que parezca fecha
                if not parsed_date:
                    for key in ['date', 'publisheddate', 'publicationdate', 'createddate', 'lastmodified']:
                        val = raw_data.get(key)
                        if val:
                            try:
                                if isinstance(val, (int, float)):
                                    parsed_date = datetime.datetime.fromtimestamp(val / 1000.0)
                                elif isinstance(val, str):
                                    parsed_date = parser.parse(val)
                                if parsed_date:
                                    break
                            except:
                                continue
                
                if not titulo or not link or not parsed_date:
                    continue
                
                # Depuración: mostrar las fechas que se están procesando
                print(f"   📅 Procesando: {parsed_date.strftime('%Y-%m-%d')} - {titulo[:50]}...")
                
                # Filtrar por el rango de fechas
                if start_date <= parsed_date <= end_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date, 
                            "Title": titulo, 
                            "Link": link, 
                            "Organismo": "FMI"
                        })
                        documentos_filtrados += 1
                        print(f"      ✅ AGREGADO: {parsed_date.strftime('%Y-%m-%d')}")
            
            print(f"\n📊 Total de blogs en el rango {start_date.date()} a {end_date.date()}: {documentos_filtrados}")
            
        else:
            print(f"❌ Error en la API: {res.status_code}")
            
    except Exception as e:
        print(f"❌ Error en load_investigacion_fmi: {e}")
        import traceback
        traceback.print_exc()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    return df

@st.cache_data(show_spinner=False)
def load_investigacion_bm(start_date_str, end_date_str):
    """Extractor para Investigación del BM (Filtra y excluye los que son 'Reports')"""
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}

    # ID exacto de la comunidad de Investigación
    scope_id = '06251f8a-62c2-59fb-add5-ec0993fc20d9'

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)

    rows, page = [], 0
    while True:
        try:
            params = {
                'scope': scope_id,
                'sort': 'dc.date.issued,DESC',
                'page': page,
                'size': 20
            }
            res = requests.get(base_url, headers=headers,
                               params=params, timeout=15)
            data = res.json()

            objects = data.get('_embedded', {}).get(
                'searchResult', {}).get('_embedded', {}).get('objects', [])
            if not objects:
                break

            items_found = 0
            for obj in objects:
                item = obj.get('_embedded', {}).get('indexableObject', {})
                meta = item.get('metadata', {})

                # Extraer Título y Fecha
                title = meta.get('dc.title', [{'value': ''}])[
                    0].get('value', '')
                date_s = meta.get('dc.date.issued', [{'value': ''}])[
                    0].get('value', '')

                parsed_date = None
                if date_s:
                    try:
                        parsed_date = parser.parse(date_s)
                    except:
                        pass

                if not parsed_date or parsed_date < start_date:
                    continue

                # --- NUEVO FILTRO ANTI-REPORTES ---
                # Buscamos en el abstract o en la descripción general
                abstract_list = meta.get('dc.description.abstract', [])
                desc_list = meta.get('dc.description', [])

                description = ""
                if abstract_list:
                    description = abstract_list[0].get('value', '').lower()
                elif desc_list:
                    description = desc_list[0].get('value', '').lower()

                # Si la palabra exacta "report" está en la descripción, lo saltamos
                # Usamos \b para que sea la palabra exacta y no algo como "reporting"
                if re.search(r'\breport\b', description):
                    continue
                # ----------------------------------

                # Link permanente
                link = meta.get('dc.identifier.uri', [{'value': ''}])[
                    0].get('value', '')
                if not link:
                    link = f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"

                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": title,
                                "Link": link, "Organismo": "BM"})
                    items_found += 1

            if items_found == 0:
                break
            page += 1
            if page > 3:
                break  # Límite para evitar búsquedas infinitas
            time.sleep(0.2)
        except:
            break

    df = pd.DataFrame(rows)
    if not df.empty:
        # 🔧 1. Convertir a datetime y FORZAR eliminación de timezone
        df["Date"] = pd.to_datetime(df["Date"], errors='coerce', utc=False)
        if df["Date"].dt.tz is not None:
            df["Date"] = df["Date"].dt.tz_localize(None)
        
        # 🔧 2. Eliminar filas con fecha inválida
        df = df.dropna(subset=['Date'])
        
        # 🔧 3. Eliminar duplicados por Link (importante)
        df = df.drop_duplicates(subset=['Link'])
        
        # 🔧 4. Ordenar por fecha descendente
        df = df.sort_values("Date", ascending=False)
        
        # 🔧 5. DEPURACIÓN (opcional, pero ayuda a detectar problemas)
        print(f"📊 BM Investigación - Total después de limpieza: {len(df)}")
        if not df.empty:
            print(f"   📅 Meses en los datos: {sorted(df['Date'].dt.month.unique())}")
            print(f"   📅 Años en los datos: {sorted(df['Date'].dt.year.unique())}")
    
    return df

## OCDE - INVESTIGACION

@st.cache_data(show_spinner=False)
def load_investigacion_ocde(start_date_str, end_date_str):
    """Extractor OCDE - Working Papers (API oficial con paginación)"""
    import requests
    import datetime
    import re
    from dateutil import parser
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 OCDE Investigación: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    
    # API base de la OCDE
    base_url = "https://api.oecd.org/webcms/search/faceted-search"
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "application/json"
    }
    
    page = 0
    page_size = 50  # Número de resultados por página
    max_pages = 10  # Límite de seguridad (500 documentos máximo)
    documentos_procesados = 0
    
    print("📡 Solicitando Working Papers a la API de la OCDE (con paginación)...")
    
    try:
        while page < max_pages:
            params = {
                "siteName": "oecd",
                "interfaceLanguage": "en",
                "orderBy": "mostRecent",
                "pageSize": page_size,
                "page": page,
                "facets": "oecd-languages:en",
                "hiddenFacets": "oecd-content-types:publications/working-papers"
            }
            
            print(f"   📄 Procesando página {page + 1}...")
            response = requests.get(base_url, params=params, headers=headers, timeout=15)
            
            if response.status_code != 200:
                print(f"   ❌ Error en página {page + 1}: {response.status_code}")
                break
            
            data = response.json()
            
            # Buscar los resultados
            results = []
            if "results" in data:
                results = data["results"]
            else:
                print(f"   ⚠️ Estructura inesperada en página {page + 1}")
                break
            
            if not results:
                print(f"   📭 No hay más resultados en página {page + 1}")
                break
            
            # Contar cuántos documentos del mes encontramos en esta página
            documentos_en_pagina = 0
            fecha_mas_reciente = None
            fecha_mas_antigua = None
            
            for item in results:
                titulo = item.get("title", "") or item.get("name", "")
                link = item.get("url", "") or item.get("link", "")
                
                if not titulo or not link:
                    continue
                
                # Extraer fecha del campo publicationDateTime
                fecha_texto = item.get("publicationDateTime", "")
                
                parsed_date = None
                if fecha_texto:
                    try:
                        parsed_date = parser.parse(fecha_texto)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        continue
                
                if not parsed_date:
                    continue
                
                # Actualizar fechas extremas
                if fecha_mas_reciente is None or parsed_date > fecha_mas_reciente:
                    fecha_mas_reciente = parsed_date
                if fecha_mas_antigua is None or parsed_date < fecha_mas_antigua:
                    fecha_mas_antigua = parsed_date
                
                # Si el documento es más antiguo que start_date, podemos parar porque
                # los resultados están ordenados por fecha descendente
                if parsed_date < start_date:
                    # Ya no hay más documentos del mes en esta página ni en las siguientes
                    print(f"   ⏹️ Documento más antiguo que {start_date.strftime('%Y-%m')}, deteniendo paginación")
                    # Salimos del while principal
                    page = max_pages
                    break
                
                # Filtrar por rango de fechas
                if parsed_date >= start_date and parsed_date <= end_date:
                    # Limpiar título
                    titulo = re.sub(r'\s+', ' ', titulo).strip()
                    
                    # Asegurar URL absoluta
                    if link.startswith('/'):
                        link = f"https://www.oecd.org{link}"
                    
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "OCDE"
                    })
                    documentos_en_pagina += 1
                    documentos_procesados += 1
            
            print(f"   📊 Página {page + 1}: {documentos_en_pagina} documentos en el rango")
            
            # Si no encontramos documentos en esta página y ya pasamos la fecha límite, paramos
            if documentos_en_pagina == 0 and fecha_mas_antigua and fecha_mas_antigua < start_date:
                print(f"   ⏹️ Fin de resultados para el mes solicitado")
                break
            
            # Si encontramos menos de page_size documentos, probablemente es la última página
            if len(results) < page_size:
                print(f"   📭 Última página alcanzada")
                break
            
            page += 1
            # Pequeña pausa para no sobrecargar la API
            time.sleep(0.3)
        
        print(f"\n📊 Total documentos OCDE encontrados: {documentos_procesados}")
        
    except Exception as e:
        print(f"❌ Error en load_investigacion_ocde: {e}")
        import traceback
        traceback.print_exc()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"📊 OCDE Investigación - Total final: {len(df)}")
    return df


# --- SECCIÓN: DISCURSOS ---

## -- Banco de Inglaterra -- Bank of England (BoE)
@st.cache_data(show_spinner=False)
def load_discursos_boe(start_date_str, end_date_str):
    """Extractor Automático BoE - Vía RSS con formato consistente 'Autor: Título'"""
    import requests
    from bs4 import BeautifulSoup
    import pandas as pd
    import datetime
    import re
    from dateutil import parser

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2025, 1, 1)
        end_date = datetime.datetime.now()

    url = "https://www.bankofengland.co.uk/rss/speeches"
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []

    def extract_author_from_title(title):
        """Extrae el nombre del autor del título en varios formatos"""
        autor = ""
        titulo_limpio = title
        
        # Patrón 1: "Título − speech by Autor" (con guión largo o corto)
        match = re.search(r'(?i)\s*[\-–—]\s*speech\s+by\s+(.+?)$', title)
        if match:
            autor = clean_author_name(match.group(1).strip())
            # Eliminar TODO desde el guión hasta el final
            titulo_limpio = re.sub(r'(?i)\s*[\-–—]\s*speech\s+by\s+.*$', '', title).strip()
            return autor, titulo_limpio
        
        # Patrón 2: "Speech by Autor: Título" o "Speech by Autor - Título"
        match = re.search(r'(?i)^speech\s+by\s+([^:—-]+)[:—-]\s*(.+)$', title)
        if match:
            autor = clean_author_name(match.group(1).strip())
            titulo_limpio = match.group(2).strip()
            return autor, titulo_limpio
        
        # Patrón 3: "Autor: Título" (ya está bien formateado)
        match = re.search(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*:\s*(.+)$', title)
        if match:
            autor = clean_author_name(match.group(1))
            titulo_limpio = match.group(2)
            return autor, titulo_limpio
        
        # Patrón 4: "Título by Autor" (sin "speech")
        match = re.search(r'(?i)\s+by\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)$', title)
        if match:
            autor = clean_author_name(match.group(1))
            titulo_limpio = re.sub(r'(?i)\s+by\s+.*$', '', title).strip()
            return autor, titulo_limpio
        
        return None, title

    try:
        res = requests.get(url, headers=headers, timeout=10)
        if res.status_code == 200:
            soup = BeautifulSoup(res.content, "xml")
            items = soup.find_all("item")

            for item in items:
                titulo_raw = item.find("title").text if item.find("title") else ""
                link = item.find("link").text if item.find("link") else ""
                fecha_raw = item.find("pubDate").text if item.find("pubDate") else ""

                if not titulo_raw or not link or not fecha_raw:
                    continue

                try:
                    parsed_date = parser.parse(fecha_raw)
                    if parsed_date.tzinfo is not None:
                        parsed_date = parsed_date.replace(tzinfo=None)
                except:
                    continue

                if start_date <= parsed_date <= end_date:
                    # Extraer autor y título limpio
                    autor, titulo_limpio = extract_author_from_title(titulo_raw)
                    
                    # LIMPIEZA DIRECTA: eliminar específicamente " − speech" o " −" al final
                    # Primero, eliminar " − speech" (con el guión especial)
                    titulo_limpio = titulo_limpio.replace(' − speech', '').replace(' - speech', '').replace('— speech', '')
                    # Luego, eliminar " −" solitario al final
                    titulo_limpio = titulo_limpio.replace(' −', '').replace(' -', '').replace('—', '')
                    # Eliminar espacios sobrantes al final
                    titulo_limpio = titulo_limpio.rstrip()
                    
                    # Construir título final en formato "Autor: Título"
                    if autor:
                        titulo_final = f"{autor}: {titulo_limpio}"
                    else:
                        titulo_final = titulo_limpio
                    
                    # Limpieza final de espacios múltiples
                    titulo_final = re.sub(r'\s+', ' ', titulo_final).strip()
                    titulo_final = titulo_final.strip('"').strip("'").strip()
                    
                    if not any(r['Link'] == link for r in rows):
                        rows.append({
                            "Date": parsed_date,
                            "Title": titulo_final,
                            "Link": link,
                            "Organismo": "BoE (Inglaterra)"
                        })
    except Exception as e:
        print(f"Error en load_discursos_boe: {e}")

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values(by="Date", ascending=False)
    return df

## -- FMI - Discursos 

@st.cache_data(show_spinner=False)
def load_discursos_fmi(start_date_str, end_date_str):
    """
    Extractor FMI - Discursos y Transcripts (Coveo API)
    """
    import datetime
    import requests
    import pandas as pd
    import re
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 FMI Discursos y Transcripts: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2025, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"

    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
    }

    # Incluir tanto SPEECHES como TRANSCRIPTS
    payload = {
        "aq": "@imftype==(\"Speech\",\"Transcript\") AND @syslanguage==\"English\"",
        "numberOfResults": 150,
        "sortCriteria": "@imfdate descending"
    }

    def limpiar_titulo(titulo, speaker_name):
        """Limpia el título: elimina comillas, sufijos redundantes y el nombre del autor si está repetido"""
        if not titulo:
            return titulo
        
        # 1. Eliminar comillas
        titulo = titulo.strip('"').strip("'")
        titulo = titulo.replace('\\"', '')
        titulo = titulo.replace('"', '')
        titulo = titulo.replace("'", "")
        
        # 2. Eliminar sufijos comunes que son redundantes
        sufijos_redundantes = [
            r'\s*[-–—]\s*(?:Keynote\s+)?Speech\s+by\s+.*$',
            r'\s*[-–—]\s*(?:Opening\s+)?Remarks\s+by\s+.*$',
            r'\s*[-–—]\s*(?:Press\s+)?Briefing\s+Transcript.*$',
            r'\s*[-–—]\s*Statement\s+by\s+.*$',
            r'\s*[-–—]\s*Address\s+by\s+.*$',
            r'\s*[-–—]\s*Transcript:.*$',
        ]
        
        for patron in sufijos_redundantes:
            titulo = re.sub(patron, '', titulo, flags=re.IGNORECASE)
        
        # 3. Si el título comienza con el nombre del autor (sin cargo), eliminarlo
        if speaker_name and titulo.lower().startswith(speaker_name.lower()):
            titulo = titulo[len(speaker_name):].lstrip(': ').strip()
        
        # 4. Limpiar espacios múltiples y caracteres sobrantes
        titulo = re.sub(r'\s+', ' ', titulo).strip()
        
        # 5. Eliminar puntuación redundante al inicio
        titulo = re.sub(r'^[,:;.\s]+', '', titulo)
        
        return titulo

    try:
        print("   📡 Solicitando discursos y transcripts del FMI a Coveo API...")
        response = requests.post(url, headers=headers, json=payload, timeout=15, verify=False)

        if response.status_code == 200:
            data = response.json()
            print(f"   ✅ Total resultados en API: {data.get('totalCount', 0)}")

            for item in data.get("results", []):
                titulo_raw = item.get("title", "").strip()
                link = item.get("clickUri", "")
                raw_date = item.get("raw", {}).get("date")
                speaker = item.get("raw", {}).get("imfspeaker", "")
                content_type = item.get("raw", {}).get("imftype", "")

                if isinstance(speaker, list) and len(speaker) > 0:
                    speaker = speaker[0]
                elif not speaker:
                    speaker = "IMF Staff"

                if not titulo_raw or not link or not raw_date:
                    continue

                try:
                    parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                except:
                    continue

                if start_date <= parsed_date <= end_date:
                    # Limpiar título usando el nombre del autor
                    titulo = limpiar_titulo(titulo_raw, speaker)
                    
                    # Construir título final
                    titulo_final = f"{speaker}: {titulo}"
                    
                    # Limpieza final
                    titulo_final = re.sub(r'\s+', ' ', titulo_final).strip()
                    
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo_final,
                        "Link": link,
                        "Organismo": "FMI"
                    })
                    print(f"      ✅ [{content_type}] {parsed_date.strftime('%d/%m/%Y')}: {speaker} - {titulo[:50]}...")

        else:
            print(f"   ❌ Error en API: {response.status_code}")

    except Exception as e:
        print(f"   ❌ Error: {e}")

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Title'], keep='first')
        df = df.drop_duplicates(subset=['Link'], keep='first')

    print(f"📊 FMI Discursos - Total: {len(df)}")
    return df

## Banco de Europa - ECB - Discursos 
@st.cache_data(show_spinner=False)
def load_data_ecb(start_date_str, end_date_str):
    """
    Extractor ECB (Europa) - Prioriza URLs con hash
    """
    import datetime
    import re
    import time
    from bs4 import BeautifulSoup
    import pandas as pd
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 ECB (Europa): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2025, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    seen_titles = set()
    
    year = start_date.year
    month = start_date.month
    
    meses = {
        'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
        'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12,
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
    }
    
    print("   🚀 Extrayendo discursos del ECB...")
    
    try:
        chrome_options = Options()
        chrome_options.add_argument("--headless=new")
        chrome_options.add_argument("--no-sandbox")
        chrome_options.add_argument("--disable-dev-shm-usage")
        chrome_options.add_argument("--window-size=1920,1080")
        chrome_options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36")
        
        driver = webdriver.Chrome(options=chrome_options)
        list_url = f"https://www.ecb.europa.eu/press/pubbydate/html/index.en.html?name_of_publication=Speech&year={year}"
        driver.get(list_url)
        
        time.sleep(8)
        for _ in range(5):
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(2)
        
        html = driver.page_source
        driver.quit()
        
        soup = BeautifulSoup(html, 'html.parser')
        
        # ========== PASO 1: Extraer TODAS las URLs con hash ==========
        url_hash_map = {}
        for link in soup.find_all('a', href=True):
            href = link['href']
            # Buscar URLs que contengan ~ (hash)
            if 'ecb.sp' in href and '~' in href:
                date_match = re.search(r'ecb\.sp(?:20)?(\d{2})(\d{2})(\d{2})', href)
                if date_match:
                    y = 2000 + int(date_match.group(1))
                    m = int(date_match.group(2))
                    d = int(date_match.group(3))
                    
                    if y == year and m == month:
                        fecha_key = f"{y}-{m}-{d}"
                        if fecha_key not in url_hash_map:
                            url_hash_map[fecha_key] = []
                        
                        full_url = href if href.startswith('http') else f"https://www.ecb.europa.eu{href}"
                        # Obtener el texto del enlace (título)
                        link_text = link.get_text(strip=True)
                        
                        url_hash_map[fecha_key].append({
                            'url': full_url,
                            'link_text': link_text
                        })
                        print(f"   📎 URL con hash encontrada para {d:02d}/{m:02d}: {full_url[:80]}...")
        
        print(f"\n   📊 URLs con hash por fecha: {sum(len(v) for v in url_hash_map.values())}")
        
        # ========== PASO 2: Extraer títulos y autores del texto ==========
        all_text = soup.get_text()
        date_pattern = r'(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})'
        
        for match in re.finditer(date_pattern, all_text):
            day = int(match.group(1))
            mes_str = match.group(2).lower()
            año = int(match.group(3))
            
            if año != year:
                continue
            
            mes_num = meses.get(mes_str, 0)
            if mes_num != month:
                continue
            
            start_pos = match.end()
            context = all_text[start_pos:start_pos + 800]
            
            if 'SPEECH' not in context.upper():
                continue
            
            speech_match = re.search(r'SPEECH\s+([^\n]+)', context, re.IGNORECASE)
            if not speech_match:
                continue
            
            titulo = speech_match.group(1).strip()
            titulo = re.sub(r'\s+', ' ', titulo).strip()
            
            # Filtrar solo basura obvia
            if titulo.lower() in ['select', 'topic', 'year', 'board member', 'jel code']:
                continue
            
            # Extraer autor
            autor = ""
            after_title = context[context.find(titulo) + len(titulo):]
            autor_match = re.search(r'\n\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})\s*\n', after_title)
            if autor_match:
                autor_raw = autor_match.group(1).strip()
                autor = re.sub(r'\s*Details.*$', '', autor_raw, flags=re.IGNORECASE)
                autor = autor.strip()
            
            parsed_date = datetime.datetime(año, mes_num, day)
            fecha_key = f"{año}-{mes_num}-{day}"
            
            # ========== BUSCAR URL CON HASH para este discurso ==========
            url_final = None
            
            if fecha_key in url_hash_map:
                # Intentar asociar por título
                for item in url_hash_map[fecha_key]:
                    # Si el texto del enlace contiene el título o el autor
                    if titulo.lower() in item['link_text'].lower() or autor.lower() in item['link_text'].lower():
                        url_final = item['url']
                        print(f"      🔗 Asociado por coincidencia: {titulo[:40]}... -> {item['link_text'][:40]}...")
                        break
                
                # Si no se encontró coincidencia, usar la primera URL con hash
                if not url_final and url_hash_map[fecha_key]:
                    url_final = url_hash_map[fecha_key][0]['url']
                    print(f"      🔗 Usando primera URL con hash para {day:02d}/{month:02d}")
            
            # Si no hay URL con hash, construir genérica (fallback)
            if not url_final:
                year_short = str(year)[2:]
                url_final = f"https://www.ecb.europa.eu/press/key/date/{year}/html/ecb.sp{year_short}{month:02d}{day:02d}.en.html"
                print(f"      ⚠️ Usando URL genérica para {day:02d}/{month:02d}")
            
            titulo_final = f"{autor}: {titulo}" if autor else titulo
            titulo_final = re.sub(r'\s+', ' ', titulo_final).strip()
            
            # Evitar duplicados por título
            if titulo_final not in seen_titles:
                seen_titles.add(titulo_final)
                rows.append({
                    "Date": parsed_date,
                    "Title": titulo_final,
                    "Link": url_final,
                    "Organismo": "ECB (Europa)"
                })
                print(f"      ✅ {parsed_date.strftime('%d/%m/%Y')}: {titulo_final[:60]}...")
        
    except Exception as e:
        print(f"   ❌ Error: {e}")
        import traceback
        traceback.print_exc()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    
    print(f"\n📊 ECB (Europa) - Total final: {len(df)} discursos")
    return df


## Discursos - BPI (BIS)
# ==========================================
# NOTA SOBRE FECHAS EN BPI (BIS)
# ==========================================
# El BIS publica discursos con dos fechas:
#   - Fecha de publicación (default): cuando el documento está disponible en la web
#   - Fecha del evento (opcional): cuando realmente se dio el discurso
#
# Por defecto usamos la fecha de publicación porque:
#   1. Es consistente con el resto de organismos
#   2. Un boletín mensual recopila PUBLICACIONES del mes
#   3. Usar fecha de evento puede incluir discursos de hasta un año atrás
#
# Para activar fechas de evento: usar checkbox en sidebar
# ==========================================
@st.cache_data(show_spinner=False)
def load_data_bis(use_event_date=False, target_year=None, target_month=None):
    """
    Extractor BIS (BPI) - Discursos
    
    Parámetros:
    -----------
    use_event_date : bool
        - False (default): Usa fecha de publicación del API
        - True: Intenta extraer fecha del evento desde el HTML
    target_year : int, optional
        Año objetivo para filtrar (ej: 2026)
    target_month : int, optional
        Mes objetivo para filtrar (ej: 5 para mayo)
    """
    import datetime
    import requests
    import pandas as pd
    import html
    import re
    from bs4 import BeautifulSoup
    from dateutil import parser
    
    print(f"\n{'='*60}")
    print(f"🔍 load_data_bis() llamada con:")
    print(f"   use_event_date = {use_event_date}")
    print(f"   target_year = {target_year}")
    print(f"   target_month = {target_month}")
    print(f"{'='*60}")
    
    urls = ["https://www.bis.org/api/document_lists/cbspeeches.json", 
            "https://www.bis.org/api/document_lists/bcbs_speeches.json", 
            "https://www.bis.org/api/document_lists/mgmtspeeches.json"]
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []
    current_year = datetime.datetime.now().year
    
    def extraer_fecha_evento_desde_html(url):
        """Descarga la página HTML y extrae la fecha del evento"""
        try:
            response = requests.get(url, headers=headers, timeout=10)
            if response.status_code != 200:
                return None
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Buscar en el div 'extratitle-div' (contiene ubicación y fecha del evento)
            extratitle = soup.find('div', id='extratitle-div')
            if extratitle:
                texto = extratitle.get_text()
            else:
                texto = soup.get_text()
            
            # Buscar patrón de fecha como "22 April 2026"
            patron = r'(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})'
            match = re.search(patron, texto, re.IGNORECASE)
            
            if match:
                dia = int(match.group(1))
                mes_str = match.group(2)
                año = int(match.group(3))
                
                meses = {
                    'january':1, 'february':2, 'march':3, 'april':4, 'may':5, 'june':6,
                    'july':7, 'august':8, 'september':9, 'october':10, 'november':11, 'december':12,
                    'jan':1, 'feb':2, 'mar':3, 'apr':4, 'may':5, 'jun':6,
                    'jul':7, 'aug':8, 'sep':9, 'oct':10, 'nov':11, 'dec':12
                }
                mes_num = meses.get(mes_str.lower()[:3], 1)
                from datetime import datetime
                return datetime(año, mes_num, min(dia, 28))
            return None
        except Exception as e:
            print(f"      ⚠️ Error extrayendo fecha de HTML: {e}")
            return None
    
    for url in urls:
        try:
            response = requests.get(url, headers=headers, timeout=15)
            if response.status_code != 200:
                continue
            data = response.json()
            
            for path, speech in data.get("list", {}).items():
                title = html.unescape(speech.get("short_title", ""))
                if not title:
                    continue
                    
                pub_date_str = speech.get("publication_start_date", "")
                if not pub_date_str:
                    continue
                    
                try:
                    pub_date = parser.parse(pub_date_str)
                except:
                    continue
                
                # ✅ FILTRO RÁPIDO: Solo procesar discursos del año/mes objetivo
                if target_year and target_month:
                    if pub_date.year != target_year or pub_date.month != target_month:
                        continue  # Saltar discursos que no son del mes solicitado
                
                link = "https://www.bis.org" + path + (".htm" if not path.endswith(".htm") else "")
                
                # Para depuración: mostrar el discurso que estamos procesando
                if "Doornbosch" in title or "Richard" in title:
                    print(f"\n   📍 PROCESANDO DISCURSO CLAVE:")
                    print(f"      Título: {title[:80]}...")
                    print(f"      Fecha publicación API: {pub_date.strftime('%Y-%m-%d')}")
                    print(f"      use_event_date = {use_event_date}")
                
                if use_event_date:
                    # Solo para discursos recientes descargamos el HTML
                    if pub_date.year >= current_year - 2:
                        evento_date = extraer_fecha_evento_desde_html(link)
                        if evento_date:
                            if "Doornbosch" in title or "Richard" in title:
                                print(f"      🎯 Fecha EVENTO encontrada: {evento_date.strftime('%Y-%m-%d')}")
                            rows.append({"Date": evento_date, "Title": title, "Link": link, "Organismo": "BPI"})
                            continue
                        else:
                            if "Doornbosch" in title or "Richard" in title:
                                print(f"      ⚠️ No se encontró fecha de evento en HTML")
                    else:
                        if "Doornbosch" in title or "Richard" in title:
                            print(f"      ⏭️ Discurso antiguo ({pub_date.year}), no se descarga HTML")
                    
                    # Fallback a fecha de publicación
                    if "Doornbosch" in title or "Richard" in title:
                        print(f"      📅 Usando FECHA PUBLICACIÓN: {pub_date.strftime('%Y-%m-%d')}")
                    rows.append({"Date": pub_date, "Title": title, "Link": link, "Organismo": "BPI"})
                else:
                    if "Doornbosch" in title or "Richard" in title:
                        print(f"      📅 Usando FECHA PUBLICACIÓN (checkbox off): {pub_date.strftime('%Y-%m-%d')}")
                    rows.append({"Date": pub_date, "Title": title, "Link": link, "Organismo": "BPI"})
                    
        except Exception as e:
            print(f"Error en {url}: {e}")
            continue
    
    df = pd.DataFrame(rows) if rows else pd.DataFrame()
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
        df = df.dropna(subset=['Date'])
        df = df.sort_values("Date", ascending=False)
    
    print(f"\n📊 load_data_bis() retorna {len(df)} discursos")
    if not df.empty:
        print("   Primeros 3 discursos:")
        for i, row in df.head(3).iterrows():
            print(f"      {row['Date'].strftime('%Y-%m-%d')}: {row['Title'][:50]}...")
    print(f"{'='*60}\n")
    
    return df


@st.cache_data(show_spinner=False)
def load_data_bbk(start_date_str, end_date_str):
    base_url = "https://www.bundesbank.de/action/en/730564/bbksearch"
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows, page = [], 0
    while True:
        params = {'sort': 'bbksortdate desc', 'dateFrom': start_date_str,
                  'dateTo': end_date_str, 'pageNumString': str(page)}
        try:
            response = requests.get(
                base_url, headers=headers, params=params, timeout=10)
        except:
            break
        soup = BeautifulSoup(response.text, 'html.parser')
        items = soup.find_all('li', class_='resultlist__item')
        if not items:
            break
        for item in items:
            fecha_tag = item.find('span', class_='metadata__date')
            fecha_str = fecha_tag.text.strip() if fecha_tag else ""
            author_tag = item.find('span', class_='metadata__authors')
            author_str = clean_author_name(
                author_tag.text) if author_tag else ""
            data_div = item.find('div', class_='teasable__data')
            link, titulo = "", ""
            if data_div and data_div.find('a'):
                a_tag = data_div.find('a')
                link = "https://www.bundesbank.de" + \
                    a_tag.get('href', '') if a_tag.get(
                        'href', '').startswith('/') else a_tag.get('href', '')
                if a_tag.find('span', class_='link__label'):
                    titulo = a_tag.find(
                        'span', class_='link__label').text.strip()
                    
            # ========== 🔽 FILTRO ANTI-ALEMÁN - VA AQUÍ 🔽 ==========
            # Excluir publicaciones en alemán por URL
            if '/de/' in link or 'german' in link.lower():
                print(f"  🚫 Excluido (URL en alemán): {titulo[:50] if titulo else 'sin título'}...")
                continue
            # ========== 🔼 FILTRO ANTI-ALEMÁN - VA AQUÍ 🔼 ==========
            if author_str and author_str not in titulo:
                titulo = f"{author_str}: {titulo}"
            if fecha_str and titulo:
                rows.append({"Date": fecha_str, "Title": titulo,
                            "Link": link, "Organismo": "BBk (Alemania)"})
        if len(items) < 10:
            break
        page += 1
        time.sleep(0.3)
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(
            df["Date"], format='%d.%m.%Y', errors='coerce')
        df = df.sort_values("Date", ascending=False)
    return df

## Discursos - Banco de China - PBoC
@st.cache_data(show_spinner=False)
def load_data_pboc(start_date_str, end_date_str):
    """
    Extractor PBoC (China) - Versión final con limpieza de título y cargos
    """
    import datetime
    import re
    import requests
    from bs4 import BeautifulSoup
    import pandas as pd
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 PBoC (China): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2025, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9',
    }
    
    year = start_date.year
    month = start_date.month
    
    print(f"   📡 Buscando discursos de {year}-{month:02d}...")
    
    # Intentar ambas URLs posibles
    urls_to_try = [
        "https://www.pbc.gov.cn/en/3688110/3688175/index.html",
        "https://www.pbc.gov.cn/en/3688110/3688175/index.html?page=1"
    ]
    
    for url in urls_to_try:
        try:
            response = requests.get(url, headers=headers, timeout=15, verify=False)
            
            if response.status_code != 200:
                continue
                
            # Corregir encoding
            response.encoding = 'utf-8'
            soup = BeautifulSoup(response.text, 'html.parser')
            items = soup.find_all('div', class_='prhhd1')
            
            if not items:
                # Fallback: buscar otros selectores
                items = soup.find_all('div', class_='ListR')
                if not items:
                    items = soup.find_all('li', class_='clearfix')
            
            for item in items:
                # Fecha
                date_span = item.find('span', class_='prhhdata')
                if not date_span:
                    date_span = item.find('span', class_='date')
                if not date_span:
                    continue
                
                fecha_texto = date_span.get_text(strip=True)
                
                try:
                    parsed_date = datetime.datetime.strptime(fecha_texto, '%Y-%m-%d')
                except:
                    continue
                
                if parsed_date.year != year or parsed_date.month != month:
                    continue
                
                # Enlace
                link_tag = item.find('a', href=True)
                if not link_tag:
                    continue
                
                # Extraer título
                listr_div = item.find('div', class_='ListR')
                if not listr_div:
                    listr_div = item.find('div', class_='listR')
                if not listr_div:
                    # Si no hay div específico, usar el texto del enlace
                    titulo_completo = link_tag.get_text(strip=True)
                else:
                    titulo_completo = listr_div.get_text(strip=True)
                
                # Eliminar la fecha del título
                titulo_completo = titulo_completo.replace(fecha_texto, '').strip()
                
                # ========== LIMPIEZA DEL TÍTULO ==========
                # 1. Eliminar todo después de "--Keynote" o "Keynote Speech by"
                titulo_limpio = re.split(r'--Keynote|Keynote Speech by', titulo_completo)[0].strip()
                
                # 2. Extraer autor (sin cargo)
                autor = ""
                
                # Patrón 1: "Governor Pan Gongsheng" o "Deputy Governor X"
                autor_match = re.search(r'(?:Governor|Deputy Governor|Administrator|Director|President)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)', titulo_completo)
                if autor_match:
                    autor = autor_match.group(1).strip()
                
                # Patrón 2: "by Pan Gongsheng"
                if not autor:
                    name_match = re.search(r'by\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', titulo_completo)
                    if name_match:
                        autor = name_match.group(1).strip()
                
                # Patrón 3: Si el título comienza con un nombre (ej. "Pan Gongsheng:")
                if not autor:
                    name_start_match = re.match(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\s*[:：]', titulo_limpio)
                    if name_start_match:
                        autor = name_start_match.group(1).strip()
                        # Eliminar el nombre del título
                        titulo_limpio = re.sub(rf'^{re.escape(autor)}[:：]\s*', '', titulo_limpio)
                
                # 4. Construir título final
                if autor:
                    # Limpiar espacios y caracteres extraños
                    autor = re.sub(r'\s+', ' ', autor).strip()
                    titulo_limpio = re.sub(r'\s+', ' ', titulo_limpio).strip()
                    titulo_final = f"{autor}: {titulo_limpio}"
                else:
                    titulo_final = titulo_limpio
                
                # Limpiar apóstrofes mal codificados
                titulo_final = titulo_final.replace('â', "'").replace('â€™', "'")
                # Eliminar espacios múltiples
                titulo_final = re.sub(r'\s+', ' ', titulo_final).strip()
                
                link = link_tag.get('href')
                if link and link.startswith('/'):
                    link = f"https://www.pbc.gov.cn{link}"
                elif not link:
                    continue
                
                # Verificar duplicados
                if not any(r['Link'] == link for r in rows):
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo_final,
                        "Link": link,
                        "Organismo": "PBoC (China)"
                    })
                    print(f"      ✅ {parsed_date.strftime('%d/%m/%Y')}: {titulo_final[:60]}...")
                    
        except Exception as e:
            print(f"   ⚠️ Error con URL {url}: {e}")
            continue
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Title'], keep='first')
        df = df.drop_duplicates(subset=['Link'], keep='first')
    
    print(f"📊 PBoC (China) - Total: {len(df)} discursos")
    return df

## FED - Discursos -
@st.cache_data(show_spinner=False)
def load_data_fed(anios_num):
    """
    Extractor Fed (Estados Unidos) - Usando API JSON oficial
    """
    import datetime
    import re
    import pandas as pd
    import requests
    import json
    
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
    
    rows = []
    
    speeches_url = "https://www.federalreserve.gov/json/ne-speeches.json"
    
    try:
        print(f"   📡 Cargando discursos de la Fed desde API...")
        response = requests.get(speeches_url, headers=headers, timeout=15)
        
        if response.status_code == 200:
            # Decodificar con utf-8-sig para eliminar BOM
            content = response.content.decode('utf-8-sig')
            speeches = json.loads(content)
            print(f"   📚 Total de discursos en API: {len(speeches)}")
            
            for speech in speeches:
                date_str = speech.get('d', '')
                if not date_str:
                    continue
                
                try:
                    date_part = date_str.split(' ')[0]
                    month, day, year = map(int, date_part.split('/'))
                    parsed_date = datetime.datetime(year, month, day)
                except Exception as e:
                    continue
                
                if parsed_date.year not in anios_num:
                    continue
                
                titulo = speech.get('t', '')
                speaker_raw = speech.get('s', '')
                link = speech.get('l', '')
                
                if not titulo or not link:
                    continue
                
                # ========== CORRECCIÓN: Extraer nombre completo ==========
                speaker_clean = speaker_raw
                
                # Patrón para extraer nombre completo (nombre + apellido)
                # Ejemplos:
                # "Vice Chair for Supervision Michelle W. Bowman" -> "Michelle W. Bowman"
                # "Governor Michael S. Barr" -> "Michael S. Barr"
                # "Chair Jerome H. Powell" -> "Jerome H. Powell"
                # "Vice Chair Philip N. Jefferson" -> "Philip N. Jefferson"
                # "Governor Lisa D. Cook" -> "Lisa D. Cook"
                # "Governor Stephen I. Miran" -> "Stephen I. Miran"
                # "Governor Christopher J. Waller" -> "Christopher J. Waller"
                
                # Buscar patrón: cargo + nombre (con posible inicial de segundo nombre)
                name_match = re.search(
                    r'(?:Chair|Vice Chair(?: for Supervision)?|Governor|President|Director)\s+([A-Z][a-z]+(?:\s+[A-Z]\.?(?:\s+[A-Z][a-z]+)?)?(?:\s+[A-Z][a-z]+)?)',
                    speaker_raw
                )
                
                if name_match:
                    speaker_clean = name_match.group(1).strip()
                else:
                    # Fallback: tomar las últimas 2-3 palabras que parezcan un nombre
                    words = speaker_raw.split()
                    # Buscar palabras que empiecen con mayúscula (posible nombre)
                    name_words = [w for w in words if re.match(r'^[A-Z][a-z]*\.?$', w)]
                    if len(name_words) >= 2:
                        speaker_clean = ' '.join(name_words[-2:])  # Nombre y apellido
                    elif len(name_words) == 1:
                        speaker_clean = name_words[-1]
                    else:
                        speaker_clean = speaker_raw
                
                # Limpiar puntos y espacios extra
                speaker_clean = re.sub(r'\s+', ' ', speaker_clean).strip()
                
                # Construir URL completa
                if link and not link.startswith('http'):
                    full_link = f"https://www.federalreserve.gov{link}"
                else:
                    full_link = link
                
                titulo_final = f"{speaker_clean}: {titulo}"
                
                rows.append({
                    "Date": parsed_date,
                    "Title": titulo_final,
                    "Link": full_link,
                    "Organismo": "Fed (Estados Unidos)"
                })
                print(f"      ✅ {parsed_date.strftime('%d/%m/%Y')}: {speaker_clean} - {titulo[:50]}...")
            
        else:
            print(f"   ❌ Error en API de la Fed: {response.status_code}")
            
    except Exception as e:
        print(f"   ❌ Error en load_data_fed: {e}")
        import traceback
        traceback.print_exc()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Title'], keep='first')
        df = df.drop_duplicates(subset=['Link'], keep='first')
    
    print(f"📊 Fed (Estados Unidos) - Total: {len(df)} discursos")
    return df

## Banco de Francia - BDF - Discursos 
@st.cache_data(show_spinner=False)
def load_data_bdf(start_date_str, end_date_str):
    """Extractor Banco de Francia (BdF) - Discursos del Gobernador (Versión Selenium)"""
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from bs4 import BeautifulSoup
    import datetime
    import time
    import re
    from dateutil import parser
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BdF (Francia) - Selenium: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")
    
    rows = []
    
    # URL principal con el filtro de discursos del Gobernador
    url = "https://www.banque-france.fr/en/governor-interventions?category%5B7052%5D=7052"
    
    # Configuración de Selenium
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
    chrome_options.add_argument("--disable-blink-features=AutomationControlled")
    chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
    chrome_options.add_experimental_option('useAutomationExtension', False)
    
    try:
        print(f"📡 Iniciando Selenium para BdF...")
        driver = webdriver.Chrome(options=chrome_options)
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        
        print(f"   Navegando a: {url}")
        driver.get(url)
        
        # Esperar a que cargue el contenido principal
        time.sleep(5)
        
        # Scroll para activar lazy loading si existe
        driver.execute_script("window.scrollTo(0, 1000);")
        time.sleep(2)
        driver.execute_script("window.scrollTo(0, 2000);")
        time.sleep(2)
        
        # Extraer el HTML ya renderizado
        html = driver.page_source
        soup = BeautifulSoup(html, 'html.parser')
        
        # Buscar los cards de discursos
        cards = soup.find_all('div', class_=lambda c: c and 'card' in c if c else False)
        
        # Si no encuentra cards, buscar directamente con selectores más específicos
        if not cards:
            cards = soup.find_all('div', class_='card')
        
        print(f"   📚 Cards encontrados: {len(cards)}")
        
        # Si aún no hay cards, buscar artículos
        if not cards:
            cards = soup.find_all('article')
            print(f"   📚 Artículos encontrados: {len(cards)}")
        
        # Mapeo de meses en inglés para fechas como "2nd of April 2026"
        meses_map = {
            'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
            'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12,
            'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
            'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
        }
        
        items_found = 0
        for card in cards:
            try:
                # === 1. EXTRAER TÍTULO Y ENLACE ===
                title_elem = None
                link = None
                
                # Buscar h3 con clase card__title o similar
                title_h3 = card.find('h3', class_=lambda c: c and 'card__title' in c if c else False)
                if not title_h3:
                    title_h3 = card.find('h3')
                
                if title_h3:
                    a_tag = title_h3.find('a')
                    if a_tag:
                        title_elem = a_tag
                        link = a_tag.get('href', '')
                
                if not title_elem:
                    # Buscar cualquier enlace con texto largo
                    for a in card.find_all('a', href=True):
                        texto = a.get_text(strip=True)
                        if len(texto) > 20:
                            title_elem = a
                            link = a.get('href', '')
                            break
                
                if not title_elem or not link:
                    continue
                
                titulo = title_elem.get_text(strip=True)
                
                # Limpiar título (eliminar saltos de línea y espacios extra)
                titulo = re.sub(r'\s+', ' ', titulo).strip()
                
                # === NUEVO: Eliminar comillas tipográficas del título original ===
                # Eliminar comillas dobles inglesas y españolas (apertura y cierre)
                titulo = titulo.replace('“', '').replace('”', '').replace('"', '').replace('«', '').replace('»', '')
                # Eliminar comillas simples si existen
                titulo = titulo.replace("'", "")

                # Construir URL absoluta
                if link.startswith('/'):
                    link = "https://www.banque-france.fr" + link
                
                # === 2. EXTRAER FECHA ===
                date_elem = None
                date_text = None
                
                # Buscar div con clase card__date
                date_div = card.find('div', class_=lambda c: c and 'card__date' in c if c else False)
                if date_div:
                    date_text = date_div.get_text(strip=True)
                else:
                    # Buscar cualquier elemento con clase que contenga 'date'
                    date_elem = card.find(class_=re.compile(r'date', re.I))
                    if date_elem:
                        date_text = date_elem.get_text(strip=True)
                
                if not date_text:
                    # Buscar en el texto del card
                    card_text = card.get_text()
                    date_match = re.search(r'(\d{1,2}(?:st|nd|rd|th)?\s+of\s+[A-Za-z]+\s+\d{4})', card_text, re.IGNORECASE)
                    if date_match:
                        date_text = date_match.group(1)
                
                if not date_text:
                    continue
                
                # Limpiar fecha: eliminar "st", "nd", "rd", "th" y "of"
                date_text = re.sub(r'(\d+)(st|nd|rd|th)\s+of\s+', r'\1 ', date_text, flags=re.IGNORECASE)
                date_text = re.sub(r'(\d+)(st|nd|rd|th)', r'\1', date_text)
                date_text = date_text.strip()
                
                # Parsear fecha
                parsed_date = None
                try:
                    # Intentar parsear formatos como "2 April 2026" o "April 2, 2026"
                    parsed_date = parser.parse(date_text)
                    if parsed_date.tzinfo is not None:
                        parsed_date = parsed_date.replace(tzinfo=None)
                except:
                    # Fallback: extraer manualmente
                    match = re.search(r'(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})', date_text, re.IGNORECASE)
                    if not match:
                        match = re.search(r'([A-Za-z]+)\s+(\d{1,2}),?\s+(\d{4})', date_text, re.IGNORECASE)
                    
                    if match:
                        groups = match.groups()
                        if len(groups) == 3:
                            # Determinar si el primer grupo es día o mes
                            if groups[0].isdigit():
                                dia = int(groups[0])
                                mes_str = groups[1].lower()
                                año = int(groups[2])
                            else:
                                mes_str = groups[0].lower()
                                dia = int(groups[1])
                                año = int(groups[2])
                            
                            mes_num = meses_map.get(mes_str, 1)
                            try:
                                parsed_date = datetime.datetime(año, mes_num, min(dia, 28))
                            except:
                                parsed_date = datetime.datetime(año, mes_num, 1)
                
                if not parsed_date:
                    continue
                
                # === 3. FILTRAR POR FECHA ===
                if parsed_date < start_date or parsed_date > end_date:
                    continue
                
                # === 4. VERIFICAR DUPLICADOS ===
                if not any(r['Link'] == link for r in rows):
                    # === NUEVO: Extraer autor desde la página del discurso ===
                    autor = None
                    # Solo intentar si el título no tiene ya formato "Nombre:"
                    if not re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+:', titulo):
                        try:
                            headers_page = {'User-Agent': 'Mozilla/5.0'}
                            page_response = requests.get(link, headers=headers_page, timeout=10)
                            if page_response.status_code == 200:
                                soup_page = BeautifulSoup(page_response.text, 'html.parser')
                                page_text = soup_page.get_text()
                                
                                # === CÓDIGO CORREGIDO ===
                                # Incluir letras acentuadas y cedilla: A-Za-zÀ-ÿç
                                match = re.search(r'Speech by ([A-Za-zÀ-ÿç\s]+?)(?:\s+Governor|\s+of|\s*$)', page_text)
                                if not match:
                                    # Fallback: capturar primeras palabras después de "Speech by"
                                    match = re.search(r'Speech by ([A-ZÀ-ÿ][a-zÀ-ÿç]+(?:\s+[A-Za-zÀ-ÿç]+)?(?:\s+[a-zÀ-ÿç]+)?(?:\s+[A-Za-zÀ-ÿç]+)?)', page_text)
                                
                                if match:
                                    autor = match.group(1).strip()
                                    # Limpiar espacios extra
                                    autor = re.sub(r'\s+', ' ', autor)
                                    print(f"      📝 Autor encontrado: {autor}")
                        except:
                            pass
                    
                    if autor:
                        # Limpiar título: eliminar comillas y espacios extra
                        titulo_limpio = titulo.strip()
                        # Eliminar comillas dobles inglesas y españolas (apertura y cierre)
                        for char in ['"', "'", '“', '”', '«', '»']:
                            if titulo_limpio.startswith(char) and titulo_limpio.endswith(char):
                                titulo_limpio = titulo_limpio[1:-1]
                                break
                        
                        # Verificar si el autor ya está al inicio del título (evitar duplicados)
                        if titulo_limpio.lower().startswith(autor.lower()):
                            titulo_final = titulo_limpio  # No añadir autor duplicado
                        else:
                            titulo_final = f"{autor}: {titulo_limpio}"
                    else:
                        titulo_final = titulo
                    
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo_final,
                        "Link": link,
                        "Organismo": "BdF (Francia)"
                    })
                    items_found += 1
                    print(f"   ✅ {parsed_date.strftime('%Y-%m-%d')}: {titulo_final[:60]}...")
                    items_found += 1
                
            except Exception as e:
                print(f"   ⚠️ Error procesando card: {e}")
                continue
        
        print(f"   📊 Documentos encontrados en BdF: {items_found}")
        driver.quit()
        
    except Exception as e:
        print(f"❌ Error en load_data_bdf: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"📊 BdF (Francia) - Total final: {len(df)}")
    return df


@st.cache_data(show_spinner=False)
def load_data_bm(start_date_str, end_date_str):
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 0
    while True:
        try:
            res = requests.get(base_url, headers=headers, params={
                               'scope': 'b6a50016-276d-56d3-bbe5-891c8d18db24', 'sort': 'dc.date.issued,DESC', 'page': page, 'size': 20}, timeout=12)
            objects = res.json().get('_embedded', {}).get(
                'searchResult', {}).get('_embedded', {}).get('objects', [])
            if not objects:
                break
            items_found = 0
            for obj in objects:
                item = obj.get('_embedded', {}).get('indexableObject', {})
                meta = item.get('metadata', {})
                title = meta.get('dc.title', [{'value': ''}])[
                    0].get('value', '')
                date_s = meta.get('dc.date.issued', [{'value': ''}])[
                    0].get('value', '')
                try:
                    parsed_date = parser.parse(date_s)
                except:
                    continue
                link = meta.get('dc.identifier.uri', [{'value': ''}])[0].get(
                    'value', '') or f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": title,
                                "Link": link, "Organismo": "BM"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date):
                break
            page += 1
            time.sleep(0.3)
        except:
            break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

## Banco de Canadá - Discrusos - boc
@st.cache_data(show_spinner=False)
def load_data_boc(start_date_str, end_date_str):
    """
    Extractor Banco de Canadá (BoC) - Versión con filtro de video/transcript
    """
    import requests
    from bs4 import BeautifulSoup
    import datetime
    import time
    import re
    from dateutil import parser

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BoC (Canadá): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")
    
    rows = []
    seen_links = set()
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
        'Accept-Language': 'en-US,en;q=0.5',
    }
    
    def es_discurso_legitimo(titulo, url):
        """
        Determina si un documento es un discurso legítimo.
        """
        titulo_lower = titulo.lower()
        url_lower = url.lower()
        
        # ========== 1. EXCLUIR POR URL ==========
        if '/multimedia/' in url_lower:
            return False
        
        # ========== 2. PALABRAS CLAVE DE EXCLUSIÓN ==========
        # Solo excluir "press conference" si NO tiene "opening statement"
        if 'press conference' in titulo_lower:
            # Si tiene "opening statement", es un discurso legítimo
            if 'opening statement' in titulo_lower:
                pass  # No excluir, continuar con otras verificaciones
            else:
                return False  # Excluir si es solo "press conference"
        
        # Otras exclusiones
        exclusiones = [
            'media availability',
            'response to',
            'labour congress',
            'cross-border coordination',
            'market participants survey',
            'notification by email',
            'summary of governing council',
            'release of the',
        ]
        
        for palabra in exclusiones:
            if palabra in titulo_lower:
                return False
        
        # ========== 3. PALABRAS CLAVE DE INCLUSIÓN ==========
        inclusiones = [
            'opening statement',
            'closing statement',
            'opening remarks',
            'closing remarks',
            'monetary policy report',
            'speech',
            'remarks',
            'address',
            'testimony',
            'fireside chat',
            'keynote',
            'inaugural',
        ]
        
        for palabra in inclusiones:
            if palabra in titulo_lower:
                return True
        
        # ========== 4. PATRONES DE TÍTULO ==========
        patrones = [
            r'^(?:Governor|Deputy Governor|Senior Deputy Governor)\s+[A-Z][a-z]+',
            r'^Adjusting for',
            r'^Global imbalances',
            r'^[A-Z][a-z]+ing\s+for',
        ]
        
        for patron in patrones:
            if re.search(patron, titulo, re.IGNORECASE):
                return True
        
        return False
    
    def tiene_transcript_o_texto(soup, url):
            """
            Verifica si la página contiene un transcript o texto del discurso.
            Retorna True si hay transcript, False si es solo video.
            """
            # ========== 1. EXCLUIR URLS DE MULTIMEDIA ==========
            if '/multimedia/' in url:
                print(f"      📹 URL de multimedia detectada, buscando transcript...")
                # Aún así, revisamos si tiene transcript
            
            # ========== 2. BUSCAR EL CONTENIDO PRINCIPAL ==========
            article = soup.find('article')
            if not article:
                article = soup.find('div', class_='entry-content')
            if not article:
                article = soup.find('main')
            if not article:
                article = soup.find('div', id='content')
            
            if not article:
                return False
            
            text = article.get_text()
            text_clean = re.sub(r'\s+', ' ', text).strip()
            
            # ========== 3. VERIFICAR SI HAY TRANSCRIPT ==========
            # Un transcript real tiene estas características:
            # - Más de 800 caracteres de texto
            # - Contiene frases típicas de discursos
            # - No está dominado por texto de descripción de video
            
            # Palabras que indican que es un transcript (no descripción)
            palabras_transcript = [
                'good morning', 'good afternoon', 'thank you',
                'governing council', 'monetary policy',
                'inflation', 'economy', 'growth',
                'interest rate', 'policy rate',
                'I am pleased', 'it is my pleasure',
                'let me', 'first', 'second', 'third',
                'governing council decided', 'we are maintaining',
                'the bank of canada today', 'our quarterly',
                'as i noted', 'as we have said',
                'let me turn to', 'i would like to',
            ]
            
            # Palabras que indican que es solo descripción de video
            palabras_video = [
                'press conference', 'media availability', 'webcast',
                'watch the video', 'video player', 'embed',
                'streaming', 'live stream', 'recorded',
                'youtube', 'vimeo', 'dailymotion',
            ]
            
            # ========== 4. ANÁLISIS DEL TEXTO ==========
            # Si el texto tiene menos de 500 caracteres, probablemente no es transcript
            if len(text_clean) < 500:
                print(f"      📄 Texto muy corto ({len(text_clean)} caracteres), no es transcript")
                return False
            
            # Contar cuántas palabras de transcript aparecen
            transcript_matches = 0
            for palabra in palabras_transcript:
                if palabra in text_clean.lower():
                    transcript_matches += 1
            
            # Contar cuántas palabras de video aparecen
            video_matches = 0
            for palabra in palabras_video:
                if palabra in text_clean.lower():
                    video_matches += 1
            
            print(f"      📊 Análisis: {transcript_matches} palabras de transcript, {video_matches} palabras de video")
            
            # ========== 5. DECISIÓN ==========
            # Si tiene al menos 3 palabras de transcript y más de 800 caracteres, es transcript
            if transcript_matches >= 3 and len(text_clean) > 800:
                return True
            
            # Si tiene más palabras de video que de transcript, es video
            if video_matches > transcript_matches:
                return False
            
            # Si tiene menos de 3 palabras de transcript, probablemente no es transcript
            if transcript_matches < 3:
                return False
            
            # Fallback: si el texto tiene más de 1000 caracteres y contiene frases de discurso
            if len(text_clean) > 1000:
                # Buscar frases típicas de inicio de discurso
                frases_inicio = ['good morning', 'thank you', 'i am pleased', 'it is my pleasure']
                for frase in frases_inicio:
                    if frase in text_clean.lower():
                        return True
            
            return False
    
    def extraer_autor_y_titulo_desde_pagina(url, titulo_lista):
        """Extrae el autor y el título de la página individual"""
        autor = None
        titulo_limpio = titulo_lista
        tiene_transcript = False
        
        try:
            response = requests.get(url, headers=headers, timeout=15)
            if response.status_code != 200:
                return autor, titulo_limpio, False
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # === VERIFICAR SI TIENE TRANSCRIPT ===
            tiene_transcript = tiene_transcript_o_texto(soup, url)
            
            # === OBTENER TÍTULO ===
            h1 = soup.find('h1', class_='entry-title')
            if h1:
                titulo_limpio = h1.get_text(strip=True)
                titulo_limpio = re.sub(r'\s+', ' ', titulo_limpio).strip()
            
            # === EXTRAER AUTOR ===
            article = soup.find('article')
            if article:
                text = article.get_text()
                match = re.search(r'(?:Governor|Senior Deputy Governor|Deputy Governor)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)', text)
                if match:
                    autor = match.group(1).strip()
                
                if not autor:
                    match = re.search(r'by\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', text)
                    if match:
                        autor = match.group(1).strip()
            
            if not autor:
                title_tag = soup.find('title')
                if title_tag:
                    title_text = title_tag.text
                    match = re.search(r'by\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)', title_text)
                    if match:
                        autor = match.group(1).strip()
            
            if not autor:
                page_text = soup.get_text()
                if 'Tiff Macklem' in page_text:
                    autor = 'Tiff Macklem'
                elif 'Carolyn Rogers' in page_text:
                    autor = 'Carolyn Rogers'
                elif 'Sharon Kozicki' in page_text:
                    autor = 'Sharon Kozicki'
            
            return autor, titulo_limpio, tiene_transcript
            
        except Exception as e:
            print(f"      ⚠️ Error extrayendo autor: {e}")
            return autor, titulo_limpio, False

    try:
        print("📡 Extrayendo discursos del Banco de Canadá...")
        
        url = "https://www.bankofcanada.ca/press/speeches/"
        print(f"📄 Procesando página principal: {url}")
        
        response = requests.get(url, headers=headers, timeout=15)
        if response.status_code != 200:
            print(f"   ❌ Error HTTP: {response.status_code}")
            return pd.DataFrame()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # ========== BUSCAR TODOS LOS ELEMENTOS ==========
        articles = []
        
        articles.extend(soup.find_all('div', class_='mtt-result'))
        articles.extend(soup.find_all('article'))
        
        main_content = soup.find('main') or soup.find('div', id='content') or soup.find('div', class_='content-area')
        if main_content:
            for div in main_content.find_all(['div', 'article', 'li'], recursive=True):
                if div.find('a', href=True):
                    has_date = div.find(string=re.compile(r'July|June|May|April|March|February|January|August|September|October|November|December', re.I))
                    if has_date:
                        articles.append(div)
        
        # Eliminar duplicados
        seen = set()
        unique_articles = []
        for art in articles:
            art_id = str(art)
            if art_id not in seen:
                seen.add(art_id)
                unique_articles.append(art)
        articles = unique_articles
        
        print(f"   📚 Elementos encontrados: {len(articles)}")
        
        items_found = 0
        for art in articles:
            try:
                # === EXTRAER TÍTULO Y ENLACE ===
                a_tag = None
                
                h3 = art.find('h3')
                if h3:
                    a_tag = h3.find('a')
                
                if not a_tag:
                    h2 = art.find('h2')
                    if h2:
                        a_tag = h2.find('a')
                
                if not a_tag:
                    for a in art.find_all('a', href=True):
                        if len(a.get_text(strip=True)) > 15:
                            a_tag = a
                            break
                
                if not a_tag:
                    continue
                
                titulo_raw = a_tag.get_text(strip=True)
                link = a_tag.get('href', '')
                
                if not link or not titulo_raw:
                    continue
                
                if link.startswith('/'):
                    link = "https://www.bankofcanada.ca" + link
                
                if link in seen_links:
                    continue
                seen_links.add(link)
                
                # === VERIFICAR SI ES UN DISCURSO LEGÍTIMO ===
                if not es_discurso_legitimo(titulo_raw, link):
                    print(f"      ⏭️ Excluido (no es discurso): {titulo_raw[:50]}...")
                    continue
                
                # === EXTRAER FECHA ===
                date_elem = art.find('span', class_='media-date')
                if not date_elem:
                    date_elem = art.find('time')
                if not date_elem:
                    date_elem = art.find(class_=re.compile(r'date', re.I))
                
                parsed_date = None
                if date_elem:
                    fecha_texto = date_elem.get_text(strip=True)
                    try:
                        parsed_date = parser.parse(fecha_texto)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        pass
                
                if not parsed_date:
                    art_text = art.get_text()
                    match = re.search(r'([A-Za-z]+)\s+(\d{1,2}),?\s+(\d{4})', art_text)
                    if match:
                        try:
                            parsed_date = parser.parse(f"{match.group(1)} {match.group(2)}, {match.group(3)}")
                            if parsed_date.tzinfo is not None:
                                parsed_date = parsed_date.replace(tzinfo=None)
                        except:
                            pass
                
                if not parsed_date:
                    continue
                
                if parsed_date < start_date or parsed_date > end_date:
                    continue
                
                print(f"   🔍 Procesando: {parsed_date.strftime('%Y-%m-%d')} - {titulo_raw[:50]}...")
                
                # === EXTRAER AUTOR, TÍTULO Y VERIFICAR TRANSCRIPT ===
                autor, titulo_limpio, tiene_transcript = extraer_autor_y_titulo_desde_pagina(link, titulo_raw)
                
                # ========== FILTRO CRÍTICO: SI NO TIENE TRANSCRIPT, NO ES DISCURSO ==========
                if not tiene_transcript:
                    print(f"      ⏭️ Excluido (sin transcript): {titulo_raw[:50]}...")
                    continue
                
                # === CONSTRUIR TÍTULO FINAL ===
                if autor and autor not in titulo_limpio:
                    titulo_final = f"{autor}: {titulo_limpio}"
                else:
                    titulo_final = titulo_limpio
                
                titulo_final = re.sub(r'\s+', ' ', titulo_final).strip()
                titulo_final = re.sub(r'([a-z])([A-Z])', r'\1 \2', titulo_final)
                
                rows.append({
                    "Date": parsed_date,
                    "Title": titulo_final,
                    "Link": link,
                    "Organismo": "BoC (Canadá)"
                })
                items_found += 1
                print(f"      ✅ Agregado: {titulo_final[:80]}...")
                
            except Exception as e:
                print(f"   ⚠️ Error procesando elemento: {e}")
                continue
        
        print(f"   📊 Discursos encontrados: {items_found}")
        
    except Exception as e:
        print(f"❌ Error en load_data_boc: {e}")
        import traceback
        traceback.print_exc()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"📊 BoC (Canadá) - Total final: {len(df)}")
    return df

## Conversor de Nombre (Nombre, Apellido) para Autores del Banco de Japón  
def convertir_nombre_japones(nombre):
    """
    Convierte nombre japonés (apellido primero) a formato occidental.
    
    Ejemplos:
    - "UEDA Kazuo" -> "Kazuo UEDA"
    - "UEDA Kazuo San" -> "Kazuo San UEDA"
    - "KURODA Haruhiko" -> "Haruhiko KURODA"
    - "AMAMIYA Masayoshi" -> "Masayoshi AMAMIYA"
    """
    if not nombre:
        return nombre
    
    partes = nombre.split()
    if len(partes) < 2:
        return nombre
    
    # La primera palabra es el apellido, el resto es el nombre
    apellido = partes[0]
    nombre_pila = " ".join(partes[1:])
    
    # Formato occidental: "Nombre Apellido"
    return f"{nombre_pila} {apellido}"

## Bank of Japan (BOJ - boj) - Discursos
@st.cache_data(show_spinner=False)
def load_data_boj(start_date_str, end_date_str):
    base_url = "https://www.boj.or.jp/en/about/press/index.htm"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BoJ (Japón): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")
    
    rows = []
    try:
        response = requests.get(base_url, headers=headers, timeout=12)
        soup = BeautifulSoup(response.text, 'html.parser')
        table = soup.find('table', class_='js-tbl')
        if table:
            for tr in table.find('tbody').find_all('tr'):
                tds = tr.find_all('td')
                if len(tds) < 3:
                    continue
                
                # === 1. EXTRAER FECHA ===
                try:
                    fecha_texto = tds[0].get_text(strip=True).replace('\xa0', ' ')
                    parsed_date = parser.parse(fecha_texto)
                except:
                    continue
                
                # Filtrar por rango de fechas
                if parsed_date < start_date or parsed_date > end_date:
                    continue
                
                # === 2. EXTRAER AUTOR (NUEVO) ===
                autor_raw = tds[1].get_text(strip=True)
                autor = None
                if autor_raw:
                    # Limpiar el autor: eliminar "Governor", "Deputy Governor", etc.
                    # Ejemplo: "UEDA Kazuo, Governor" -> "UEDA Kazuo"
                    autor = re.sub(r',\s*(Governor|Deputy Governor|Member of the Policy Board)$', '', autor_raw)
                    # Limpiar espacios extra
                    autor = autor.strip()
                    # === CONVERTIR NOMBRE JAPONÉS A FORMATO OCCIDENTAL ===
                    autor = convertir_nombre_japones(autor)
                
                # === 3. EXTRAER TÍTULO Y ENLACE ===
                a_tag = tds[2].find('a', href=True)
                if not a_tag:
                    continue
                
                titulo_raw = a_tag.get_text(strip=True).strip('"')
                link = "https://www.boj.or.jp" + a_tag['href'] if a_tag['href'].startswith('/') else a_tag['href']
                
                # === 4. CONSTRUIR TÍTULO FINAL CON AUTOR ===
                if autor:
                    # Limpiar título: eliminar el nombre del autor si está repetido
                    titulo_limpio = titulo_raw
                    # Si el título comienza con el nombre del autor, lo removemos
                    if titulo_limpio.startswith(autor.split(',')[0]):
                        titulo_limpio = re.sub(r'^[^:：]+[:：]\s*', '', titulo_limpio)
                    
                    titulo_final = f"{autor}: {titulo_limpio}"
                else:
                    titulo_final = titulo_raw
                
                rows.append({
                    "Date": parsed_date, 
                    "Title": titulo_final, 
                    "Link": link, 
                    "Organismo": "BoJ (Japón)"
                })
                print(f"   ✅ {parsed_date.strftime('%Y-%m-%d')}: {titulo_final[:60]}...")
    except Exception as e:
        print(f"⚠️ Error en load_data_boj: {e}")
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    print(f"📊 BoJ (Japón) - Total final: {len(df)}")
    return df

## --------------------------------------------------------------

@st.cache_data(show_spinner=False)
def load_data_cef(start_date_str, end_date_str):
    """
    Extractor CEF (FSB) - SOLO Discursos y Statements
    Con manejo robusto de timeouts y fallbacks para autor
    """
    import requests
    from bs4 import BeautifulSoup
    import datetime
    import time
    import re
    from dateutil import parser
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 CEF (FSB): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8'
    }
    
    rows = []
    page = 1
    
    def es_discurso(url, titulo):
        """Determina si una página es un discurso"""
        titulo_lower = titulo.lower()
        url_lower = url.lower()
        
        # Excluir comunicados de prensa puros
        if re.match(r'^(fsb publishes|fsb warns|fsb chair warns)(?!.*(speech|keynote|summit))', titulo_lower):
            return False
        
        # Incluir por URL
        if any(keyword in url_lower for keyword in ['/speech/', '/statement/', '/remarks/']):
            return True
        
        # Incluir por palabras clave en título
        if any(keyword in titulo_lower for keyword in [
            'speech', 'keynote', 'remarks', 'statement', 'foreword', 
            'address', 'testimony', 'opening remarks', 'closing remarks'
        ]):
            return True
        
        # Incluir si menciona autoridades del FSB
        if any(title_word in titulo_lower for title_word in ['fsb chair', 'secretary general', 'deputy governor']):
            return True
        
        return False
    
    def inferir_autor_desde_titulo(titulo):
        """Infiere el autor basándose en el título cuando no se puede acceder a la página"""
        titulo_lower = titulo.lower()
        
        # Palabras clave que indican quién es el autor
        if 'fsb chair' in titulo_lower or 'chair' in titulo_lower:
            return 'Andrew Bailey'
        if 'secretary general' in titulo_lower:
            return 'John Schindler'
        if 'deputy governor' in titulo_lower:
            # Podría ser varios, pero intentamos extraer del contexto
            if 'john schindler' in titulo_lower:
                return 'John Schindler'
            return 'FSB Deputy Governor'
        
        return None
    
    def extraer_autor_y_titulo_desde_pagina(url, titulo_lista):
        """Extrae el autor y el título limpio de la página individual con manejo de timeouts"""
        autor = None
        titulo_limpio = titulo_lista
        
        try:
            # Timeout más generoso y reintento
            time.sleep(0.5)
            response = requests.get(url, headers=headers, timeout=15)
            if response.status_code != 200:
                # Fallback: inferir autor del título
                autor = inferir_autor_desde_titulo(titulo_lista)
                return autor, titulo_limpio
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # === OBTENER TÍTULO CORRECTO DEL <h1> ===
            h1_tag = soup.find('h1')
            if h1_tag:
                titulo_limpio = h1_tag.get_text(strip=True)
                titulo_limpio = re.sub(r'\s+', ' ', titulo_limpio).strip()
            
            # === EXTRAER AUTOR ===
            # Método 1: Buscar en el bloque blockquote
            blockquote = soup.find('blockquote')
            if blockquote:
                texto = blockquote.get_text()
                match = re.search(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),?\s+(?:the\s+)?(?:Chair|Secretary General|Deputy Governor|Governor)', texto)
                if match:
                    autor = match.group(1).strip()
                
                if not autor:
                    match = re.search(r'by\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', texto)
                    if match:
                        autor = match.group(1).strip()
            
            # Método 2: Buscar en meta tags de perfil
            if not autor:
                meta_profile = soup.find('meta', attrs={'name': 'fsb_profile_post'})
                if meta_profile:
                    profile_value = meta_profile.get('content', '').lower()
                    nombres = {
                        'andrew-bailey': 'Andrew Bailey',
                        'john-schindler': 'John Schindler',
                        'klaas-knot': 'Klaas Knot',
                        'martin-moloney': 'Martin Moloney'
                    }
                    for key, name in nombres.items():
                        if key in profile_value:
                            autor = name
                            break
            
            # Método 3: Si el título contiene "FSB Chair", el autor es Andrew Bailey
            if not autor and ('FSB Chair' in titulo_limpio or 'Chair' in titulo_limpio):
                autor = 'Andrew Bailey'
            
            # Método 4: Buscar en el contenido del artículo
            if not autor:
                article = soup.find('article')
                if article:
                    text = article.get_text()
                    match = re.search(r'Speech\s+by\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', text, re.IGNORECASE)
                    if match:
                        autor = match.group(1).strip()
            
            # Si aún no hay autor, intentar inferir del título
            if not autor:
                autor = inferir_autor_desde_titulo(titulo_limpio)
            
            return autor, titulo_limpio
            
        except requests.exceptions.Timeout:
            print(f"      ⚠️ Timeout al acceder a {url}, infiriendo autor del título...")
            autor = inferir_autor_desde_titulo(titulo_lista)
            return autor, titulo_limpio
        except Exception as e:
            print(f"      ⚠️ Error: {e}")
            autor = inferir_autor_desde_titulo(titulo_lista)
            return autor, titulo_limpio
    
    while True:
        try:
            if page == 1:
                url = "https://www.fsb.org/press/speeches-and-statements/"
            else:
                url = f"https://www.fsb.org/press/speeches-and-statements/page/{page}/"
            
            print(f"📄 Procesando página {page}: {url}")
            
            response = requests.get(url, headers=headers, timeout=20)
            if response.status_code != 200:
                print(f"   ❌ Error HTTP: {response.status_code}")
                break
            
            soup = BeautifulSoup(response.text, 'html.parser')
            items = soup.find_all('div', class_='post-excerpt')
            
            if not items:
                items = soup.find_all('div', class_=lambda c: c and 'post-excerpt' in c if c else False)
            
            if not items:
                print(f"   📭 No se encontraron más elementos en página {page}")
                break
            
            print(f"   📚 Elementos encontrados: {len(items)}")
            items_found = 0
            
            for item in items:
                try:
                    title_elem = item.find('h3')
                    if not title_elem:
                        title_elem = item.find('div', class_='post-title')
                    
                    if not title_elem:
                        continue
                    
                    a_tag = title_elem.find('a')
                    if not a_tag:
                        continue
                    
                    titulo_raw = a_tag.get_text(strip=True)
                    link = a_tag.get('href', '')
                    
                    if not link:
                        continue
                    
                    date_elem = item.find('div', class_='post-date')
                    if not date_elem:
                        date_elem = item.find('span', class_='post-date')
                    
                    if not date_elem:
                        continue
                    
                    fecha_texto = date_elem.get_text(strip=True)
                    
                    try:
                        parsed_date = parser.parse(fecha_texto)
                        if parsed_date.tzinfo is not None:
                            parsed_date = parsed_date.replace(tzinfo=None)
                    except:
                        continue
                    
                    if parsed_date < start_date or parsed_date > end_date:
                        continue
                    
                    if any(r['Link'] == link for r in rows):
                        continue
                    
                    print(f"   🔍 Procesando: {parsed_date.strftime('%Y-%m-%d')} - {titulo_raw[:50]}...")
                    
                    if not es_discurso(link, titulo_raw):
                        print(f"      ⏭️ Excluido (no es discurso): {titulo_raw[:50]}...")
                        continue
                    
                    # === EXTRAER AUTOR Y TÍTULO ===
                    autor, titulo_limpio = extraer_autor_y_titulo_desde_pagina(link, titulo_raw)
                    
                    # === CONSTRUIR TÍTULO FINAL ===
                    if autor and titulo_limpio:
                        # Verificar si el autor ya está al inicio del título
                        if not titulo_limpio.lower().startswith(autor.lower()):
                            titulo_final = f"{autor}: {titulo_limpio}"
                        else:
                            titulo_final = titulo_limpio
                    else:
                        titulo_final = titulo_limpio
                    
                    # Limpieza mínima
                    titulo_final = re.sub(r'\s+', ' ', titulo_final).strip()
                    titulo_final = titulo_final.replace('â', "'").replace('â€™', "'")
                    
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo_final,
                        "Link": link,
                        "Organismo": "CEF"
                    })
                    items_found += 1
                    print(f"      ✅ Discurso: {titulo_final[:80]}...")
                    
                except Exception as e:
                    print(f"   ⚠️ Error procesando item: {e}")
                    continue
            
            print(f"   📊 Discursos en página {page}: {items_found}")
            
            # Si no encontramos discursos en 2 páginas consecutivas, paramos
            if items_found == 0 and page > 2:
                break
            
            page += 1
            time.sleep(1.5)  # Pausa más larga entre páginas
            
        except requests.exceptions.Timeout:
            print(f"   ⏱️ Timeout en página {page}, continuando...")
            page += 1
            time.sleep(3)
            continue
        except Exception as e:
            print(f"❌ Error en página {page}: {e}")
            break
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        df = df.drop_duplicates(subset=['Link'])
    
    print(f"\n📊 CEF (FSB) - Total final: {len(df)} discursos")
    return df

## - Discursos - Banco de España - 
@st.cache_data(show_spinner=False)
def load_data_bde(start_date_str, end_date_str):
    """
    Extractor Banco de España - Versión SIN Selenium con selectores flexibles y filtro de discursos
    """
    import requests
    from bs4 import BeautifulSoup
    import datetime
    import time
    import re
    import io
    from PyPDF2 import PdfReader

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BdE (España): {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2025, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    
    url = "https://www.bde.es/wbe/en/noticias-eventos/actualidad-banco-espana/intervenciones-publicas/"
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
        'Accept-Language': 'en-US,en;q=0.5',
        'Accept-Encoding': 'gzip, deflate, br',
        'Connection': 'keep-alive',
        'Upgrade-Insecure-Requests': '1',
    }

    def extraer_autor_y_cargo_desde_pdf(pdf_url):
        """Extrae el nombre y cargo del autor desde el PDF"""
        try:
            response = requests.get(pdf_url, headers=headers, timeout=15)
            if response.status_code != 200:
                return None, None
            
            pdf_file = io.BytesIO(response.content)
            reader = PdfReader(pdf_file)
            
            text = ""
            for i in range(min(3, len(reader.pages))):
                page_text = reader.pages[i].extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if not text:
                return None, None
            
            lineas = text.split('\n')
            nombre = None
            cargo = None
            
            for i, linea in enumerate(lineas):
                linea_limpia = linea.strip()
                
                if re.search(r'Governor|Gobernador', linea_limpia, re.IGNORECASE):
                    cargo = "Governor"
                    if i > 0 and lineas[i-1].strip() and len(lineas[i-1].strip().split()) >= 2:
                        nombre = lineas[i-1].strip()
                    elif i + 1 < len(lineas) and lineas[i+1].strip() and len(lineas[i+1].strip().split()) >= 2:
                        nombre = lineas[i+1].strip()
                    break
                elif re.search(r'Deputy Governor|Subgobernador', linea_limpia, re.IGNORECASE):
                    cargo = "Deputy Governor"
                    if i > 0 and lineas[i-1].strip() and len(lineas[i-1].strip().split()) >= 2:
                        nombre = lineas[i-1].strip()
                    elif i + 1 < len(lineas) and lineas[i+1].strip() and len(lineas[i+1].strip().split()) >= 2:
                        nombre = lineas[i+1].strip()
                    break
            
            if not nombre:
                for linea in lineas[:15]:
                    linea_limpia = linea.strip()
                    if re.match(r'^[A-ZÁÉÍÓÚÑ]{2,}(?:\s+[A-ZÁÉÍÓÚÑ]{2,}){1,3}$', linea_limpia):
                        if not any(palabra in linea_limpia for palabra in ['DIRECTOR', 'GENERAL', 'DEPARTAMENTO', 'SECRETARÍA', 'MINISTERIO', 'GOBIERNO', 'BANCO', 'ESPAÑA', 'MADRID']):
                            nombre = linea_limpia
                            break
            
            if nombre:
                nombre = ' '.join(nombre.split())
                nombre = nombre.title()
                nombre = re.sub(r'\bDe\b', 'de', nombre)
                nombre = re.sub(r'\bY\b', 'y', nombre)
                return nombre, cargo
            
            return None, None
            
        except Exception as e:
            print(f"      ⚠️ Error extrayendo del PDF: {e}")
            return None, None

    def es_discurso(titulo):
        """
        Determina si un documento es un discurso basado en su título.
        Retorna True si es discurso, False si no lo es.
        """
        titulo_lower = titulo.lower()
        
        # ========== 1. EXCEPCIONES ESPECÍFICAS ==========
        # Discursos que deben incluirse aunque no cumplan otros patrones
        excepciones = [
            'conference on the spanish economy',  # Soledad Núñez tiene PDF de discurso
            'deputy governor. conference on the',  # Mismo caso
            'subgobernadora. conferencia sobre',   # Versión en español
        ]
        for excepcion in excepciones:
            if excepcion in titulo_lower:
                return True
        
        # ========== 2. PALABRAS CLAVE DE EXCLUSIÓN ==========
        # Si el título contiene alguna de estas, NO es un discurso
        palabras_excluir = [
            # Presentaciones y reportes (que NO son discursos)
            'presentación', 'presentacion', 'presentation of',
            'annual report', 'memoria anual',
            'summary', 'resumen ejecutivo',
            'brochure', 'folleto',
            'infographic', 'infografía',
            'powerpoint', 'power point', 'ppt',
            'slides', 'diapositivas',
            'video', 'podcast',
            'press release', 'comunicado de prensa',
            'interview', 'entrevista',
            'article', 'artículo',
            'blog post',
            # Conferencias que NO son discursos (sin autor)
            '5th banco de españa',
            'banco de españa-cemfi-uimp',
        ]
        
        # ========== 3. VERIFICAR SI TIENE AUTOR (NOMBRE AL INICIO) ==========
        # Patrón: "Nombre Apellido: Título" o "Nombre Apellido - Título"
        tiene_autor = bool(re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+[:：\-–—]', titulo))
        
        # ========== 4. LÓGICA DE FILTRADO ==========
        
        # Caso especial: Si tiene autor Y es una conferencia, es un discurso
        # Ejemplo: "Soledad Núñez: Conference on the Spanish Economy"
        if tiene_autor and 'conference' in titulo_lower:
            # Verificar que no sea una presentación
            es_presentacion = any(p in titulo_lower for p in ['presentation of', 'presentación'])
            if not es_presentacion:
                return True
        
        # Caso especial: Si tiene autor Y el título contiene palabras de discurso
        if tiene_autor:
            palabras_discurso = ['speech', 'remarks', 'address', 'statement', 'intervención', 'discurso']
            for p in palabras_discurso:
                if p in titulo_lower:
                    return True
            # Si tiene autor pero no tiene palabras de discurso, lo incluimos
            # (asumimos que es un discurso a menos que sea una presentación)
            es_presentacion = any(p in titulo_lower for p in ['presentation of', 'presentación', 'slides', 'ppt'])
            if not es_presentacion:
                return True
        
        # ========== 5. PALABRAS CLAVE DE INCLUSIÓN ==========
        # Si el título contiene alguna de estas, ES un discurso
        palabras_incluir = [
            'speech', 'discurso',
            'remarks', 'palabras',
            'opening remarks', 'palabras de apertura',
            'closing remarks', 'palabras de clausura',
            'keynote', 'inaugural',
            'address', 'intervención',
            'statement', 'declaración',
            'testimony', 'testimonio',
        ]
        
        for palabra in palabras_incluir:
            if palabra in titulo_lower:
                return True
        
        # ========== 6. EXCLUIR POR PALABRAS DE EXCLUSIÓN ==========
        for palabra in palabras_excluir:
            if palabra in titulo_lower:
                return False
        
        # ========== 7. FALBACK ==========
        # Si no está claro, lo incluimos (mejor falso positivo que falso negativo)
        # Pero solo si tiene más de 4 palabras (para evitar títulos muy cortos)
        if len(titulo.split()) >= 4:
            return True
        
        return False

    try:
        print(f"📡 Solicitando página: {url}")
        response = requests.get(url, headers=headers, timeout=15)
        
        if response.status_code != 200:
            print(f"❌ Error al acceder a la página: {response.status_code}")
            return pd.DataFrame()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # ========== BUSCAR ELEMENTOS CON MÚLTIPLES SELECTORES ==========
        items = []
        
        selectores = [
            'div.block-search-result',
            'div.block-search-result--image',
            'article.search-result',
            'li.search-result',
            'div.search-result',
            'div.result-item',
            'div.teaser',
            'div.news-item',
            'article.news',
            'div.listing-item',
            'div.item',
            'div[class*="search-result"]',
            'div[class*="result"]',
            'div[class*="news"]',
            'div[class*="publication"]',
            '.search-results article',
            '.search-results li',
            '.results article',
            '.results li',
        ]
        
        for selector in selectores:
            encontrados = soup.select(selector)
            if encontrados:
                print(f"   ✅ Selector '{selector}' encontró {len(encontrados)} elementos")
                items = encontrados
                break
        
        if not items:
            print("   ⚠️ Buscando elementos por estructura genérica...")
            for div in soup.find_all(['div', 'article', 'li']):
                has_link = div.find('a', href=True)
                has_date = div.find(string=re.compile(r'\d{2}/\d{2}/\d{4}'))
                if has_link and has_date:
                    items.append(div)
            print(f"   📚 Encontrados {len(items)} elementos por estructura genérica")
        
        if not items:
            print("⚠️ No se encontraron elementos en la página")
            with open("bde_debug.html", "w", encoding="utf-8") as f:
                f.write(response.text)
            print("   💾 HTML guardado en bde_debug.html para depuración")
            return pd.DataFrame()
        
        print(f"   📚 Elementos encontrados: {len(items)}")
        
        for item in items:
            try:
                # === 1. EXTRAER TÍTULO ===
                title_elem = None
                
                for tag in ['h3', 'h4', 'h5']:
                    title_elem = item.find(tag)
                    if title_elem:
                        break
                
                if not title_elem:
                    title_elem = item.find('div', class_=re.compile(r'title', re.I))
                
                if not title_elem:
                    title_elem = item.find('a', class_=re.compile(r'title', re.I))
                
                if not title_elem:
                    for a in item.find_all('a', href=True):
                        texto = a.get_text(strip=True)
                        if len(texto) > 15:
                            title_elem = a
                            break
                
                if not title_elem:
                    continue
                
                a_tag = title_elem.find('a') if title_elem.name != 'a' else title_elem
                if not a_tag:
                    a_tag = title_elem
                
                raw_title = a_tag.get_text(strip=True)
                link = a_tag.get('href', '')
                
                if len(raw_title) < 10:
                    parent_text = item.get_text(strip=True)
                    raw_title = re.sub(r'\s+', ' ', parent_text).strip()
                    match = re.match(r'^(.*?)\s+\d{2}/\d{2}/\d{4}', raw_title)
                    if match:
                        raw_title = match.group(1).strip()
                
                if not raw_title or len(raw_title) < 5:
                    continue
                
                # === FILTRO: VERIFICAR SI ES UN DISCURSO ===
                if not es_discurso(raw_title):
                    print(f"      ⏭️ Excluido (no es discurso): {raw_title[:50]}...")
                    continue
                
                # === 2. EXTRAER FECHA ===
                date_elem = None
                
                date_elem = item.find('div', class_=re.compile(r'date', re.I))
                if not date_elem:
                    date_elem = item.find('span', class_=re.compile(r'date', re.I))
                if not date_elem:
                    date_elem = item.find('time')
                if not date_elem:
                    item_text = item.get_text()
                    match = re.search(r'(\d{2}/\d{2}/\d{4})', item_text)
                    if match:
                        raw_date_str = match.group(1)
                    else:
                        continue
                else:
                    raw_date_str = date_elem.get_text(strip=True)
                
                parsed_date = None
                try:
                    parsed_date = datetime.datetime.strptime(raw_date_str, '%d/%m/%Y')
                except:
                    match = re.search(r'(\d{2}/\d{2}/\d{4})', raw_date_str)
                    if match:
                        try:
                            parsed_date = datetime.datetime.strptime(match.group(1), '%d/%m/%Y')
                        except:
                            pass
                
                if not parsed_date:
                    continue
                
                if parsed_date < start_date or parsed_date > end_date:
                    continue
                
                # === 3. CONSTRUIR URL COMPLETA ===
                if link.startswith('/'):
                    link = "https://www.bde.es" + link
                elif not link.startswith('http'):
                    link = "https://www.bde.es" + '/' + link.lstrip('/')
                
                print(f"   🔍 Procesando: {parsed_date.strftime('%Y-%m-%d')} - {raw_title[:50]}...")
                
                # === 4. BUSCAR PDF EN LA PÁGINA INDIVIDUAL ===
                pdf_link = None
                autor = None
                titulo_final = raw_title
                
                try:
                    page_response = requests.get(link, headers=headers, timeout=10)
                    if page_response.status_code == 200:
                        page_soup = BeautifulSoup(page_response.text, 'html.parser')
                        
                        for a in page_soup.find_all('a', href=True):
                            href = a['href']
                            if href.endswith('.pdf') or '.pdf' in href.lower():
                                pdf_link = href
                                if pdf_link.startswith('/'):
                                    pdf_link = "https://www.bde.es" + pdf_link
                                elif not pdf_link.startswith('http'):
                                    pdf_link = "https://www.bde.es" + '/' + pdf_link.lstrip('/')
                                break
                        
                        # ========== FILTRO: Si no hay PDF, NO es un discurso ==========
                        if not pdf_link:
                            print(f"      ⏭️ Excluido (sin PDF, probablemente video): {raw_title[:50]}...")
                            continue
                        # ===================================================================
                        
                        if pdf_link:
                            print(f"      📄 PDF encontrado, extrayendo autor...")
                            autor, cargo = extraer_autor_y_cargo_desde_pdf(pdf_link)
                            if autor:
                                print(f"      📝 Autor extraído: {autor}")
                except Exception as e:
                    print(f"      ⚠️ Error accediendo a página individual: {e}")
                    continue
                
                # === 5. LIMPIAR TÍTULO ===
                titulo_limpio = raw_title
                
                patrones_cargo_principio = [
                    r'^Governor\.\s*', r'^Deputy\s*Governor\.\s*', 
                    r'^Subgobernador[a]?\.\s*', r'^Director\s*General\.\s*',
                    r'^D\.G\.\s*Econom[íi]a\.\s*', r'^D\.G\.\s*Economics\.\s*',
                    r'^Director\.\s*', r'^Chief\s*Economist\.\s*',
                    r'^Gerente\s*General\.\s*', r'^Vicepresident[ae]\.\s*',
                    r'^President[ae]\.\s*', r'^Head\s*of\s*\w+\.\s*',
                ]
                
                for patron in patrones_cargo_principio:
                    titulo_limpio = re.sub(patron, '', titulo_limpio, flags=re.IGNORECASE)
                
                titulo_limpio = re.sub(r'^[:.\-\s]+', '', titulo_limpio).strip()
                titulo_limpio = re.sub(r'\s+', ' ', titulo_limpio).strip()
                
                # === 6. CONSTRUIR TÍTULO FINAL ===
                if autor:
                    if autor.lower() not in titulo_limpio.lower():
                        titulo_final = f"{autor}: {titulo_limpio}"
                    else:
                        titulo_final = titulo_limpio
                else:
                    titulo_final = titulo_limpio
                
                titulo_final = re.sub(r'\s+', ' ', titulo_final).strip()
                titulo_final = titulo_final.strip('"').strip("'").strip()
                
                # === 7. AGREGAR A RESULTADOS ===
                if not any(r['Link'] == link for r in rows):
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo_final,
                        "Link": link,
                        "Organismo": "BdE (España)"
                    })
                    print(f"      ✅ Agregado: {titulo_final[:80]}...")
                
            except Exception as e:
                print(f"   ⚠️ Error procesando item: {e}")
                continue
                
    except Exception as e:
        print(f"❌ Error BDE: {e}")
        import traceback
        traceback.print_exc()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    
    print(f"📊 BdE (España) - Total final: {len(df)}")
    return df

# ==========================================
# NUEVAS FUNCIONES PARA BID (bypass Cloudflare)
# ==========================================

@st.cache_data(show_spinner=False)
def load_investigacion_bid_cloudscraper(start_date_str, end_date_str):
    """
    Extrae Working Papers usando cloudscraper (bypass Cloudflare)
    """
    try:
        import cloudscraper
    except ImportError:
        print("❌ cloudscraper no instalado. Ejecuta: pip install cloudscraper")
        return pd.DataFrame()
    
    from bs4 import BeautifulSoup
    import datetime
    import re
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BID Cloudscraper: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    
    # Crear scraper con configuraciones específicas
    scraper = cloudscraper.create_scraper(
        browser={
            'browser': 'chrome',
            'platform': 'windows',
            'mobile': False
        },
        delay=5
    )
    
    # URLs a probar
    urls_to_try = [
        "https://publications.iadb.org/en?f%5B0%5D=type%3AWorking%20Papers",
        "https://publications.iadb.org/es?f%5B0%5D=type%3A4633&f%5B1%5D=type%3ADocumentos%20de%20Trabajo"
    ]
    
    for url in urls_to_try:
        lang = "en" if "en?" in url else "es"
        try:
            print(f"📡 Accediendo a {url[:60]}...")
            response = scraper.get(url, timeout=30)
            
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Extraer artículos
                articles = soup.find_all('div', class_='views-row')
                print(f"   📚 Artículos encontrados: {len(articles)}")
                
                for article in articles:
                    # Extraer título y link
                    title_elem = article.find('div', class_='views-field-field-title')
                    if not title_elem:
                        continue
                    
                    a_tag = title_elem.find('a')
                    if not a_tag:
                        continue
                    
                    titulo = a_tag.get_text(strip=True)
                    link = a_tag.get('href')
                    if link and not link.startswith('http'):
                        link = "https://publications.iadb.org" + link
                    
                    # Extraer fecha
                    date_elem = article.find('div', class_='views-field-field-date-issued-text')
                    if date_elem:
                        date_text = date_elem.get_text(strip=True)
                        # Parsear fecha "Mar 2026"
                        match = re.search(r'([A-Za-z]{3})\s+(\d{4})', date_text)
                        if match:
                            mes_str, año = match.groups()
                            meses = {'Jan':1, 'Feb':2, 'Mar':3, 'Apr':4, 'May':5, 'Jun':6,
                                   'Jul':7, 'Aug':8, 'Sep':9, 'Oct':10, 'Nov':11, 'Dec':12}
                            mes = meses.get(mes_str, 1)
                            parsed_date = datetime.datetime(int(año), mes, 1)
                            
                            if start_date <= parsed_date <= end_date:
                                rows.append({
                                    "Date": parsed_date,
                                    "Title": titulo,
                                    "Link": link,
                                    "Organismo": f"BID ({'Inglés' if lang == 'en' else 'Español'})"
                                })
                                print(f"      ✅ {parsed_date.strftime('%Y-%m')}: {titulo[:50]}...")
            
        except Exception as e:
            print(f"⚠️ Error en {lang}: {e}")
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    
    print(f"📊 BID Cloudscraper - Total: {len(df)} documentos")
    return df


@st.cache_data(show_spinner=False)
def load_investigacion_bid_selenium_fallback(start_date_str, end_date_str):
    """
    Fallback: Extrae Working Papers con Selenium + delay largo
    """
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from bs4 import BeautifulSoup
    import datetime
    import time
    import re
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 BID Selenium Fallback: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
    
    rows = []
    
    chrome_options = Options()
    chrome_options.add_argument('--headless=new')
    chrome_options.add_argument('--no-sandbox')
    chrome_options.add_argument('--disable-dev-shm-usage')
    chrome_options.add_argument('--disable-blink-features=AutomationControlled')
    chrome_options.add_argument('--window-size=1920,1080')
    
    urls = [
        ("https://publications.iadb.org/en?f%5B0%5D=type%3AWorking%20Papers", "en"),
        ("https://publications.iadb.org/es?f%5B0%5D=type%3A4633&f%5B1%5D=type%3ADocumentos%20de%20Trabajo", "es")
    ]
    
    for url, lang in urls:
        driver = None
        try:
            print(f"📡 Accediendo con Selenium a {url[:60]}...")
            driver = webdriver.Chrome(options=chrome_options)
            driver.get(url)
            
            # ⚠️ CLAVE: Esperar a que Cloudflare resuelva
            print("   ⏳ Esperando 20 segundos para Cloudflare...")
            time.sleep(20)
            
            # Scroll para cargar contenido
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(3)
            
            # Extraer usando BeautifulSoup
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            articles = soup.find_all('div', class_='views-row')
            print(f"   📚 Artículos encontrados: {len(articles)}")
            
            for article in articles:
                title_elem = article.find('div', class_='views-field-field-title')
                if not title_elem:
                    continue
                
                a_tag = title_elem.find('a')
                if not a_tag:
                    continue
                
                titulo = a_tag.get_text(strip=True)
                link = a_tag.get('href')
                if link and not link.startswith('http'):
                    link = "https://publications.iadb.org" + link
                
                date_elem = article.find('div', class_='views-field-field-date-issued-text')
                if date_elem:
                    date_text = date_elem.get_text(strip=True)
                    match = re.search(r'([A-Za-z]{3})\s+(\d{4})', date_text)
                    if match:
                        mes_str, año = match.groups()
                        meses = {'Jan':1, 'Feb':2, 'Mar':3, 'Apr':4, 'May':5, 'Jun':6,
                               'Jul':7, 'Aug':8, 'Sep':9, 'Oct':10, 'Nov':11, 'Dec':12,
                               'ene':1, 'feb':2, 'mar':3, 'abr':4, 'may':5, 'jun':6,
                               'jul':7, 'ago':8, 'sep':9, 'oct':10, 'nov':11, 'dic':12}
                        mes = meses.get(mes_str, 1)
                        parsed_date = datetime.datetime(int(año), mes, 1)
                        
                        if start_date <= parsed_date <= end_date:
                            rows.append({
                                "Date": parsed_date,
                                "Title": titulo,
                                "Link": link,
                                "Organismo": f"BID ({'Inglés' if lang == 'en' else 'Español'})"
                            })
            
        except Exception as e:
            print(f"⚠️ Error Selenium en {lang}: {e}")
        finally:
            if driver:
                driver.quit()
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    
    print(f"📊 BID Selenium - Total: {len(df)} documentos")
    return df


def load_investigacion_bid_unified(start_date_str, end_date_str):
    """
    UNIFICADOR: Prueba cloudscraper primero, si falla usa Selenium
    """
    print("="*50)
    print("🔍 Iniciando extracción BID con estrategia unificada")
    print("="*50)
    
    # Intentar primero con cloudscraper
    try:
        print("\n🚀 Estrategia 1: Cloudscraper")
        df = load_investigacion_bid_cloudscraper(start_date_str, end_date_str)
        if not df.empty:
            print(f"✅ Cloudscraper exitoso: {len(df)} documentos")
            return df
        else:
            print("⚠️ Cloudscraper no obtuvo resultados")
    except Exception as e:
        print(f"⚠️ Cloudscraper falló: {e}")
    
    # Fallback a Selenium
    print("\n🚀 Estrategia 2: Selenium con delay largo")
    try:
        df = load_investigacion_bid_selenium_fallback(start_date_str, end_date_str)
        if not df.empty:
            print(f"✅ Selenium exitoso: {len(df)} documentos")
            return df
        else:
            print("⚠️ Selenium no obtuvo resultados")
    except Exception as e:
        print(f"⚠️ Selenium falló: {e}")
    
    print("\n❌ Ambas estrategias fallaron para BID")
    return pd.DataFrame()

# ==========================================
# EXPORTACIÓN A WORD
# ==========================================


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0000EE')
    rPr.append(c)
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)
    b = docx.oxml.shared.OxmlElement('w:b')
    rPr.append(b)

    for s in ['w:sz', 'w:szCs']:
        sz = docx.oxml.shared.OxmlElement(s)
        sz.set(docx.oxml.shared.qn('w:val'), '28')
        rPr.append(sz)

    rFonts = docx.oxml.shared.OxmlElement('w:rFonts')
    rFonts.set(docx.oxml.shared.qn('w:ascii'), 'Calibri')
    rFonts.set(docx.oxml.shared.qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)
    t = docx.oxml.shared.OxmlElement('w:t')
    t.text = text
    new_run.append(rPr)
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def generate_word(df, title="Boletín Mensual", subtitle=""):
    doc = Document()
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.name, run.font.size = 'Calibri', Pt(14)
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=len(df.columns)-1)
    table.style = 'Table Grid'

    cols = [c for c in df.columns if c != 'Link']

    for idx, name in enumerate(cols):
        p = table.rows[0].cells[idx].paragraphs[0]
        run = p.add_run(name)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.bold = True

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            p = cells[i].paragraphs[0]
            if col == 'Nombre de Documento':
                add_hyperlink(p, str(row[col]), str(row['Link']))
            else:
                run = p.add_run(str(row[col]))
                run.font.name = 'Calibri'
                run.font.size = Pt(14)
                run.bold = True

    if 'Tipo de Documento' in df.columns and 'Organismo' in df.columns:
        col_tipo = cols.index('Tipo de Documento')
        col_org = cols.index('Organismo')

        start_row = 1
        while start_row <= len(df):
            cat_val = df.iloc[start_row - 1]['Tipo de Documento']
            org_val = df.iloc[start_row - 1]['Organismo']
            end_row = start_row

            if cat_val == "Discursos":
                table.cell(start_row, col_org).text = ""
                while end_row < len(df) and df.iloc[end_row]['Tipo de Documento'] == "Discursos":
                    table.cell(end_row + 1, col_org).text = ""
                    end_row += 1

                if end_row > start_row:
                    target_cell = table.cell(start_row, col_org)
                    target_cell.merge(table.cell(end_row, col_org))

                start_row = end_row + 1
                continue

            while end_row < len(df) and df.iloc[end_row]['Tipo de Documento'] == cat_val and df.iloc[end_row]['Organismo'] == org_val:
                table.cell(end_row + 1, col_org).text = ""
                end_row += 1

            if end_row > start_row:
                target_cell = table.cell(start_row, col_org)
                target_cell.merge(table.cell(end_row, col_org))
                target_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            start_row = end_row + 1

        start_row = 1
        while start_row <= len(df):
            cat_val = df.iloc[start_row - 1]['Tipo de Documento']
            end_row = start_row

            while end_row < len(df) and df.iloc[end_row]['Tipo de Documento'] == cat_val:
                table.cell(end_row + 1, col_tipo).text = ""
                end_row += 1

            if end_row > start_row:
                target_cell = table.cell(start_row, col_tipo)
                target_cell.merge(table.cell(end_row, col_tipo))
                target_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            start_row = end_row + 1

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

## FUNCIÓN DE WORD AUDITADO 
def generate_word_auditado(df, title="Boletín Mensual", subtitle=""):
    """
    Genera un Word con auditoría de fuentes.
    Incluye: ID Enlace, Nombre de la Fuente, Origen, Estado, Título con hipervínculo,
    y agrupa los documentos por su fuente original.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import docx
    
    doc = Document()
    
    # Configurar márgenes más amplios para mejor visualización
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Título
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.name, run.font.size = 'Calibri', Pt(14)
    doc.add_paragraph()
    
    # Verificar qué columnas existen
    columnas_disponibles = df.columns.tolist()
    print(f"📋 Columnas disponibles en df_word: {columnas_disponibles}")
    
    # Usar 'Fecha' si existe, o 'Date' si no
    columna_fecha = 'Fecha' if 'Fecha' in columnas_disponibles else 'Date'
    
    # ========== NUEVO: Identificar si existe la columna 'Estado' ==========
    tiene_estado = 'Estado' in columnas_disponibles
    
    # Ordenar por categoría, organismo, fuente y fecha
    df = df.sort_values(['Categoría', 'Organismo', 'Nombre Fuente', columna_fecha], 
                        ascending=[True, True, True, False])
    
    # Agrupar por fuente original
    for (categoria, organismo, nombre_fuente, id_enlace, origen), group in df.groupby(
        ['Categoría', 'Organismo', 'Nombre Fuente', 'ID Enlace', 'Origen']
    ):
        # Título de la sección
        p = doc.add_paragraph()
        run = p.add_run(f"{categoria} | {organismo} | {id_enlace}")
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.color.rgb = docx.shared.RGBColor(0, 32, 91)  # Azul oscuro
        
        p = doc.add_paragraph()
        
        # ========== NUEVO: Mostrar estado si existe ==========
        if tiene_estado and 'Estado' in group.columns:
            estado = group['Estado'].iloc[0] if not group['Estado'].empty else "⚠️ Sin publicaciones en el mes"
            run = p.add_run(f"Fuente: {nombre_fuente} (Origen: {origen}) | {estado}")
        else:
            run = p.add_run(f"Fuente: {nombre_fuente} (Origen: {origen})")
        
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        
        # ========== NUEVO: Si no hay documentos, mostrar mensaje ==========
        # Verificar si hay documentos reales (Title no vacío)
        tiene_documentos = False
        for _, row in group.iterrows():
            titulo = str(row.get('Nombre de Documento', ''))
            if titulo and titulo.strip():
                tiene_documentos = True
                break
        
        if not tiene_documentos:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run("Sin documentos disponibles este mes")
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.italic = True
        else:
            # Lista de documentos bajo esta fuente
            for _, row in group.iterrows():
                p = doc.add_paragraph(style='List Bullet')
                titulo = str(row.get('Nombre de Documento', ''))
                link = str(row.get('Enlace', ''))
                
                # Si el título está vacío, saltar (es un enlace fijo sin documentos)
                if not titulo or not titulo.strip():
                    continue
                
                fecha = row.get(columna_fecha, '')
                
                # Formatear fecha
                if hasattr(fecha, 'strftime'):
                    fecha_str = fecha.strftime('%d/%m/%Y')
                else:
                    fecha_str = str(fecha)
                
                # Agregar el título con hipervínculo
                if link and link.startswith('http'):
                    run = p.add_run(f"{fecha_str} - ")
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                    
                    add_hyperlink(p, titulo, link)
                else:
                    run = p.add_run(f"{fecha_str} - {titulo}")
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
        
        doc.add_paragraph()  # Espacio entre secciones
    
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# ==========================================
# INTERFAZ DE USUARIO Y MAIN
# ==========================================
try:
    st.sidebar.image("logo_banxico.png", width="stretch")
except:
    st.sidebar.markdown("### 🏦 BANCO DE MÉXICO")

st.sidebar.markdown("---")
st.sidebar.header("Menú de Navegación")
modo_app = st.sidebar.radio(
    label="Menú Principal", 
    options= ["Boletín", "Categorías", "Carga Manual"], 
    key="menu_principal")
st.sidebar.markdown("---")

anios_str = ["2026", "2025", "2024", "2023", "2022"]
meses_dict = {
    "Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4, "Mayo": 5, "Junio": 6,
    "Julio": 7, "Agosto": 8, "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12
}

# --- LISTAS DINÁMICAS DE ORGANISMOS ---
orgs_discursos = ["BBk (Alemania)", "BdE (España)", "BdF (Francia)", "BM", "BoC (Canadá)", "BoE (Inglaterra)", "BoJ (Japón)", "BPI", "CEF", "ECB (Europa)", "Fed (Estados Unidos)", "FMI", "PBoC (China)"]
orgs_reportes = ["BID", "BM", "BPI", "CEF", "FEM", "OCDE"]
orgs_pub_inst = ["BM", "BPI", "CEF", "CEMLA", "FMI", "F&D", "G20", "OCDE", "OEI", "F&D Magazine"]
orgs_investigacion = ["BID", "BM", "BPI", "CEMLA", "FMI", "OCDE"]

st.sidebar.markdown("---")
st.sidebar.subheader("⚙️ Configuración avanzada")

# Opción para BPI (BIS) - Discursos
use_event_date_bis = st.sidebar.checkbox(
    "📅 BPI (BIS): Usar fecha del evento en lugar de publicación",
    value=False,
    help="⚠️ EXPERIMENTAL: Puede causar inconsistencias. Marcar solo si se necesita fecha exacta del discurso.\n\nEjemplo: '5 May 2025' (evento) vs '6 May 2026' (publicación)"
)

if use_event_date_bis:
    st.sidebar.info("📌 Modo activado: BPI usará fechas de evento (ej: '5 May 2025' en lugar de '6 May 2026')")
else:
    st.sidebar.caption("Modo estándar: usando fecha de publicación (recomendado)")

st.sidebar.markdown("---")

if modo_app == "Boletín":
    st.title("Generador de Boletín Mensual")
    st.markdown(
        "Extrae y unifica documentos de todas las categorías y organismos por mes.")
    st.markdown("---")

    c1, c2 = st.columns(2)
    m_sel = c1.multiselect("Mes(es)", options=list(meses_dict.keys()))
    a_sel = c2.multiselect("Año(s)", options=anios_str, default=["2026"])


    if st.button("📄 Generar Boletín Mensual", type="primary"):
        if not m_sel or not a_sel:
            st.warning("⚠️ Selecciona mes y año.")
        else:
            m_num = [meses_dict[m] for m in m_sel]
            a_num = [int(a) for a in a_sel]
            sd = f"01.{min(m_num):02d}.{min(a_num)}"
            ed = f"{calendar.monthrange(max(a_num), max(m_num))[1]:02d}.{max(m_num):02d}.{max(a_num)}"

            all_dfs = []
            prog = st.progress(0)
            txt = st.empty()

            total_pasos = len(orgs_discursos) + len(orgs_reportes) + \
                len(orgs_pub_inst) + len(orgs_investigacion)
            paso_actual = 0

            # 1. BARRIDO DE DISCURSOS
            for org in orgs_discursos:
                txt.text(f"Procesando Discursos: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BPI":
                        df = load_data_bis(
                            use_event_date=use_event_date_bis,
                            target_year=a_num[0] if a_num else None,
                            target_month=m_num[0] if m_num else None
                        )
                    elif org == "ECB (Europa)":
                        df = load_data_ecb(sd, ed)
                    elif org == "FMI":
                        df = load_discursos_fmi(sd, ed)
                    elif org == "BBk (Alemania)":
                        df = load_data_bbk(sd, ed)
                    elif org == "Fed (Estados Unidos)":
                        df = load_data_fed(a_num)
                    elif org == "BdF (Francia)":
                        df = load_data_bdf(sd, ed)
                    elif org == "BM":
                        df = load_data_bm(sd, ed)
                    elif org == "BoC (Canadá)":
                        df = load_data_boc(sd, ed)
                    elif org == "BoJ (Japón)":
                        df = load_data_boj(sd, ed)
                    elif org == "BoE (Inglaterra)":
                        df = load_discursos_boe(sd, ed)
                    elif org == "CEF":
                        df = load_data_cef(sd, ed)
                    elif org == "PBoC (China)":
                        df = load_data_pboc(sd, ed)
                    elif org == "BdE (España)":
                        df = load_data_bde(sd, ed)
                except Exception as e:
                    print(f"❌ Error en {org}: {e}")
                    pass

                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty:
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Discursos"
                        all_dfs.append(df_f)
                paso_actual += 1
                prog.progress(paso_actual / total_pasos)

            # 2. BARRIDO DE REPORTES
            for org in orgs_reportes:
                txt.text(f"Procesando Reportes: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BID":
                        df = load_reportes_bid_en(sd, ed)
                    elif org == "BM":
                        df = load_reportes_bm(sd, ed) 
                    elif org == "BPI":
                        df = load_reportes_bpi(sd, ed)
                    elif org == "CEF":
                        df = load_reportes_cef(sd, ed)
                    elif org == "OCDE":
                        df = load_reportes_ocde(sd, ed)
                    elif org == "FEM": 
                        df = load_reportes_fem(sd, ed)
                except Exception as e:
                    pass

                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (
                        df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty:
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Reportes"
                        all_dfs.append(df_f)
                paso_actual += 1
                prog.progress(paso_actual / total_pasos)

            # 3. BARRIDO DE PUBLICACIONES INSTITUCIONALES
            for org in orgs_pub_inst:
                txt.text(f"Procesando Pub. Institucionales: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BPI":
                        df = load_pub_inst_bpi(sd, ed)
                    elif org == "CEF":
                        df = load_pub_inst_cef(sd, ed)
                    elif org == "BM":
                        df = load_pub_inst_bm(sd, ed)
                    elif org == "OEI": 
                        df = load_pub_inst_oei(sd, ed)
                    elif org == "OCDE":
                        df = load_pub_inst_ocde(sd, ed)
                    elif org == "F&D":  
                        df = load_pub_inst_fandd(sd, ed)
                    elif org == "CEMLA":
                        df = load_pub_inst_cemla(sd, ed)
                    elif org == "FMI":
                        # 1. SSG - JSON Estático (WEO, Fiscal Monitor)
                        df_flagships = load_pub_inst_fmi(sd, ed)
                        print(f"📊 FMI - Flagships: {len(df_flagships)} documentos")

                        # 2. SSG - JSON Estático (Comunicados)
                        df_prs = load_press_releases_fmi(sd, ed)
                        print(f"📊 FMI - Press Releases: {len(df_prs)} documentos")

                        # 3. CSR API - Coveo (Country Reports)
                        df_crs = load_country_reports_fmi(sd, ed)
                        print(f"📊 FMI - Country Reports: {len(df_crs)} documentos")

                        # 4. NUEVO: Mission Concluding (de load_fmi_news_all)
                        df_mcs = load_fmi_news_all(sd, ed)
                        print(f"📊 FMI - Mission Concluding: {len(df_mcs)} documentos")

                        # Unión de todos
                        dfs_a_unir = [d for d in [df_flagships, df_prs, df_crs, df_mcs] if not d.empty]
                        if dfs_a_unir:
                            df = pd.concat(dfs_a_unir, ignore_index=True)
                            df = df.sort_values("Date", ascending=False)
                            print(f"📊 FMI - TOTAL combinado: {len(df)} documentos")
                    elif org == "G20":  
                        df = load_pub_inst_g20(sd, ed)
                except Exception as e:
                    print(f"Error en {org}: {e}")

                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (
                        df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty:
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Publicaciones Institucionales"
                        all_dfs.append(df_f)
                paso_actual += 1
                prog.progress(paso_actual / total_pasos)

            # 4. BARRIDO DE INVESTIGACIÓN
            for org in orgs_investigacion:
                txt.text(f"Procesando Investigación: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BID": 
                        df = load_investigacion_bid_unified(sd, ed)
                    elif org == "BPI": df = load_investigacion_bpi(sd, ed)
                    elif org == "BM": df = load_investigacion_bm(sd, ed)
                    elif org == "CEMLA":
                        df = load_investigacion_cemla(sd, ed)
                    elif org == "FMI": 
                        df_blogs = pd.DataFrame()
                        df_wp = pd.DataFrame()
                        
                        try:
                            df_blogs = load_investigacion_fmi(sd, ed)
                        except: pass
                        
                        try:
                            df_wp = load_working_papers_fmi(sd, ed)
                        except: pass
                        
                        # Unimos Blogs y Working Papers
                        dfs_a_unir = [d for d in [df_blogs, df_wp] if not d.empty]
                        if dfs_a_unir:
                            df = pd.concat(dfs_a_unir, ignore_index=True)
                            df = df.drop_duplicates(subset=['Link'])
                            df = df.sort_values("Date", ascending=False)
                    # ========== NUEVO BLOQUE PARA OCDE ==========
                    elif org == "OCDE":
                        print(f"🔍 === ENTRANDO A OCDE INVESTIGACIÓN ===")
                        print(f"   sd: {sd}, ed: {ed}")
                        try:
                            df = load_investigacion_ocde(sd, ed)
                            print(f"📊 OCDE Investigación: {len(df)} documentos encontrados")
                        except Exception as e:
                            print(f"⚠️ Error en OCDE Investigación: {e}")
                            import traceback
                            traceback.print_exc()
                    # ==========================================
                except Exception as e:
                    print(f"⚠️ Error general en {org}: {e}")
                    pass
            
                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (
                        df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty:
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Investigación"
                        all_dfs.append(df_f)
                paso_actual += 1
                prog.progress(paso_actual / total_pasos)

            # --- INYECCIÓN DE TEXTO MANUAL ---
            if 'df_extra' in st.session_state and not st.session_state.df_extra.empty:
                all_dfs.append(st.session_state.df_extra)
                txt.text("Inyectando reportes manuales...")

            txt.empty()
            prog.empty()

            # --- CONSOLIDACIÓN FINAL ---
            if all_dfs:
                f_df = pd.concat(all_dfs, ignore_index=True)

                            # ========== ELIMINACIÓN MEJORADA DE DUPLICADOS ==========
                print(f"📊 Total antes de desduplicar: {len(f_df)}")
                
                # 1. Eliminar duplicados exactos por Link
                f_df = f_df.drop_duplicates(subset=['Link'], keep='first')
                print(f"   Después de eliminar duplicados por Link: {len(f_df)}")
                
                # 2. Eliminar duplicados por Título (normalizado)
                # Normalizar títulos: quitar caracteres especiales, espacios múltiples, pasar a minúsculas
                f_df['Title_Normalized'] = f_df['Title'].str.lower()
                f_df['Title_Normalized'] = f_df['Title_Normalized'].str.replace(r'[^\w\s]', '', regex=True)
                f_df['Title_Normalized'] = f_df['Title_Normalized'].str.replace(r'\s+', ' ', regex=True).str.strip()
                
                # Eliminar duplicados por título normalizado, manteniendo el primero (el más reciente por fecha)
                f_df = f_df.sort_values('Date', ascending=False).drop_duplicates(subset=['Title_Normalized'], keep='first')
                print(f"   Después de eliminar duplicados por título: {len(f_df)}")
                
                # 3. Eliminar duplicados que sean casi idénticos (similitud de título > 90%)
                # Esto ayuda con títulos como "Preserving stability..." vs "Preserving Stability..."
                def is_similar(title1, title2, threshold=0.9):
                    """Compara similitud entre dos títulos usando secuencia de palabras"""
                    words1 = set(title1.lower().split())
                    words2 = set(title2.lower().split())
                    if not words1 or not words2:
                        return False
                    intersection = words1.intersection(words2)
                    union = words1.union(words2)
                    return len(intersection) / len(union) > threshold
                
                # Comparar títulos dentro de cada categoría y organismo
                indices_a_eliminar = set()
                for categoria in f_df['Categoría'].unique():
                    for organismo in f_df['Organismo'].unique():
                        mask = (f_df['Categoría'] == categoria) & (f_df['Organismo'] == organismo)
                        df_subset = f_df[mask].copy()
                        
                        for i in range(len(df_subset)):
                            if i in indices_a_eliminar:
                                continue
                            title_i = df_subset.iloc[i]['Title_Normalized']
                            for j in range(i + 1, len(df_subset)):
                                if j in indices_a_eliminar:
                                    continue
                                title_j = df_subset.iloc[j]['Title_Normalized']
                                if is_similar(title_i, title_j):
                                    # Mantener el más reciente
                                    date_i = df_subset.iloc[i]['Date']
                                    date_j = df_subset.iloc[j]['Date']
                                    if date_i >= date_j:
                                        indices_a_eliminar.add(df_subset.index[j])
                                    else:
                                        indices_a_eliminar.add(df_subset.index[i])
                
                # Eliminar duplicados similares
                f_df = f_df.drop(index=indices_a_eliminar, errors='ignore')
                print(f"   Después de eliminar duplicados similares: {len(f_df)}")
                
                # Eliminar columna temporal
                f_df = f_df.drop(columns=['Title_Normalized'], errors='ignore')
                
                print(f"📊 Total después de desduplicación: {len(f_df)}")

                # --- PREPARACIÓN PARA EL WORD (Orden Institucional) ---
                df_rep = f_df[f_df['Categoría'] == "Reportes"].copy()
                df_pub = f_df[f_df['Categoría'] ==
                              "Publicaciones Institucionales"].copy()
                df_inv = f_df[f_df['Categoría'] == "Investigación"].copy()
                df_disc = f_df[f_df['Categoría'] == "Discursos"].copy()

                if not df_rep.empty:
                    df_rep = df_rep.sort_values(
                        by=["Organismo", "Title"], ascending=[True, True])
                if not df_pub.empty:
                    df_pub = df_pub.sort_values(
                        by=["Organismo", "Title"], ascending=[True, True])
                if not df_inv.empty:
                    df_inv = df_inv.sort_values(
                        by=["Organismo", "Title"], ascending=[True, True])
                if not df_disc.empty:
                    df_disc = df_disc.sort_values(
                        by=["Title"], ascending=[True])

                f_df_word = pd.concat(
                    [df_rep, df_pub, df_inv, df_disc], ignore_index=True)
                f_df_word = f_df_word[['Categoría',
                                       'Organismo', 'Title', 'Link']]
                f_df_word = f_df_word.rename(
                    columns={"Categoría": "Tipo de Documento", "Title": "Nombre de Documento"})

                st.success(
                    f"Se consolidaron **{len(f_df)}** documentos en total.")
                word = generate_word(f_df_word, subtitle=", ".join(
                    m_sel) + " " + ", ".join(a_sel))
                
                # Botón de Descarga
                st.download_button("📄 Descargar Boletín en Word",
                                   word, f"Boletin_{'_'.join(m_sel)}.docx")

                # Limpiar cache manual después de generar el Word para no duplicar en el futuro
                if 'df_extra' in st.session_state:
                    del st.session_state.df_extra

                # --- PREPARACIÓN PARA LA VISTA PREVIA ---
                disp = f_df.copy()
                disp = disp.sort_values(
                    by="Date", ascending=False)  # Orden cronológico
                disp["Fecha"] = disp["Date"].dt.strftime('%d/%m/%Y')
                disp["Nombre de Documento"] = disp.apply(
                    lambda x: f"[{x['Title']}]({x['Link']})", axis=1)
                disp = disp.rename(columns={"Categoría": "Tipo de Documento"})

                st.markdown(disp[["Fecha", "Tipo de Documento", "Organismo", "Nombre de Documento"]].to_markdown(
                    index=False), unsafe_allow_html=True)
            else:
                st.warning(
                    "No se encontraron documentos para los criterios seleccionados.")

    # ==========================================
    # BOTÓN 2: CON AUDITORÍA
    # ==========================================
    if st.button("🔍 Generar Boletín con Auditoría", type="secondary"):
        if not m_sel or not a_sel:
            st.warning("⚠️ Selecciona mes y año.")
        else:
            with st.spinner("Generando boletín con auditoría..."):
                m_num = [meses_dict[m] for m in m_sel]
                a_num = [int(a) for a in a_sel]
                sd = f"01.{min(m_num):02d}.{min(a_num)}"
                ed = f"{calendar.monthrange(max(a_num), max(m_num))[1]:02d}.{max(m_num):02d}.{max(a_num)}"

                all_dfs = []
                
                                # ===== CÓDIGO DE EXTRACCIÓN COPIADO DEL BOTÓN 1 =====
                # (copias TODO desde "prog = st.progress(0)" hasta "txt.empty()")
                prog = st.progress(0)
                txt = st.empty()
                
                total_pasos = len(orgs_discursos) + len(orgs_reportes) + \
                    len(orgs_pub_inst) + len(orgs_investigacion)
                paso_actual = 0
                
                # 1. BARRIDO DE DISCURSOS
                for org in orgs_discursos:
                    txt.text(f"Procesando Discursos: {org}...")
                    df = pd.DataFrame()
                    try:
                        if org == "BPI":
                            df = load_data_bis(
                                use_event_date=use_event_date_bis,
                                target_year=a_num[0] if a_num else None,
                                target_month=m_num[0] if m_num else None
                            )
                        elif org == "ECB (Europa)":
                            df = load_data_ecb(sd, ed)
                        elif org == "FMI":
                            df = load_discursos_fmi(sd, ed)
                        elif org == "BBk (Alemania)":
                            df = load_data_bbk(sd, ed)
                        elif org == "Fed (Estados Unidos)":
                            df = load_data_fed(a_num)
                        elif org == "BdF (Francia)":
                            df = load_data_bdf(sd, ed)
                        elif org == "BM":
                            df = load_data_bm(sd, ed)
                        elif org == "BoC (Canadá)":
                            df = load_data_boc(sd, ed)
                        elif org == "BoJ (Japón)":
                            df = load_data_boj(sd, ed)
                        elif org == "BoE (Inglaterra)":
                            df = load_discursos_boe(sd, ed)
                        elif org == "CEF":
                            df = load_data_cef(sd, ed)
                        elif org == "PBoC (China)":
                            df = load_data_pboc(sd, ed)
                        elif org == "BdE (España)":
                            df = load_data_bde(sd, ed)
                    except Exception as e:
                        print(f"❌ Error en {org}: {e}")
                        pass
                    
                    if not df.empty:
                        df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                        df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                        if not df_f.empty:
                            df_f['Organismo'] = org
                            df_f['Categoría'] = "Discursos"
                            all_dfs.append(df_f)
                    paso_actual += 1
                    prog.progress(paso_actual / total_pasos)
                
                # 2. BARRIDO DE REPORTES
                for org in orgs_reportes:
                    txt.text(f"Procesando Reportes: {org}...")
                    df = pd.DataFrame()
                    try:
                        if org == "BID":
                            df = load_reportes_bid_en(sd, ed)
                        elif org == "BM":
                            df = load_reportes_bm(sd, ed)
                        elif org == "BPI":
                            df = load_reportes_bpi(sd, ed)
                        elif org == "CEF":
                            df = load_reportes_cef(sd, ed)
                        elif org == "OCDE":
                            df = load_reportes_ocde(sd, ed)
                        elif org == "FEM":
                            df = load_reportes_fem(sd, ed)
                    except Exception as e:
                        pass
                    
                    if not df.empty:
                        df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                        df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                        if not df_f.empty:
                            df_f['Organismo'] = org
                            df_f['Categoría'] = "Reportes"
                            all_dfs.append(df_f)
                    paso_actual += 1
                    prog.progress(paso_actual / total_pasos)
                
                # 3. BARRIDO DE PUBLICACIONES INSTITUCIONALES
                # 3. BARRIDO DE PUBLICACIONES INSTITUCIONALES (BLOQUE MEJORADO PARA FMI)
                for org in orgs_pub_inst:
                    txt.text(f"Procesando Pub. Institucionales: {org}...")
                    df = pd.DataFrame()
                    try:
                        if org == "BPI":
                            df = load_pub_inst_bpi(sd, ed)
                        elif org == "CEF":
                            df = load_pub_inst_cef(sd, ed)
                        elif org == "BM":
                            df = load_pub_inst_bm(sd, ed)
                        elif org == "OEI":
                            df = load_pub_inst_oei(sd, ed)
                        elif org == "OCDE":
                            df = load_pub_inst_ocde(sd, ed)
                        elif org == "F&D":
                            df = load_pub_inst_fandd(sd, ed)
                        elif org == "CEMLA":
                            df = load_pub_inst_cemla(sd, ed)
                        elif org == "FMI":
                            # --- MEJORA: Extraer TODOS los documentos ---
                            df_flagships = load_pub_inst_fmi(sd, ed)
                            df_prs = load_press_releases_fmi(sd, ed)
                            df_crs = load_country_reports_fmi(sd, ed)
                            df_mcs = load_fmi_news_all(sd, ed)

                            # Unir todos los documentos del FMI para Publicaciones Institucionales
                            dfs_a_unir = [d for d in [df_flagships, df_prs, df_crs, df_mcs] if not d.empty]
                            if dfs_a_unir:
                                # Este df contiene TODOS los documentos de FMI para el mes
                                df = pd.concat(dfs_a_unir, ignore_index=True)
                                df = df.sort_values("Date", ascending=False)
                                print(f"   📊 FMI - TOTAL documentos extraídos: {len(df)}")
                            else:
                                print(f"   ⚠️ FMI - No se encontraron documentos para el mes")
                        elif org == "G20":
                            df = load_pub_inst_g20(sd, ed)
                    except Exception as e:
                        print(f"Error en {org}: {e}")

                    if not df.empty:
                        df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                        df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                        if not df_f.empty:
                            df_f['Organismo'] = org
                            df_f['Categoría'] = "Publicaciones Institucionales"
                            all_dfs.append(df_f)
                    paso_actual += 1
                    prog.progress(paso_actual / total_pasos)
                
                # 4. BARRIDO DE INVESTIGACIÓN
                for org in orgs_investigacion:
                    txt.text(f"Procesando Investigación: {org}...")
                    df = pd.DataFrame()
                    try:
                        if org == "BID":
                            df = load_investigacion_bid_unified(sd, ed)
                        elif org == "BPI":
                            df = load_investigacion_bpi(sd, ed)
                        elif org == "BM":
                            df = load_investigacion_bm(sd, ed)
                        elif org == "CEMLA":
                            df = load_investigacion_cemla(sd, ed)
                        elif org == "FMI":
                            df_blogs = pd.DataFrame()
                            df_wp = pd.DataFrame()
                            try:
                                df_blogs = load_investigacion_fmi(sd, ed)
                            except:
                                pass
                            try:
                                df_wp = load_working_papers_fmi(sd, ed)
                            except:
                                pass
                            dfs_a_unir = [d for d in [df_blogs, df_wp] if not d.empty]
                            if dfs_a_unir:
                                df = pd.concat(dfs_a_unir, ignore_index=True)
                                df = df.drop_duplicates(subset=['Link'])
                                df = df.sort_values("Date", ascending=False)
                        elif org == "OCDE":
                            df = load_investigacion_ocde(sd, ed)
                    except Exception as e:
                        print(f"⚠️ Error general en {org}: {e}")
                        pass
                    
                    if not df.empty:
                        df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                        df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                        if not df_f.empty:
                            df_f['Organismo'] = org
                            df_f['Categoría'] = "Investigación"
                            all_dfs.append(df_f)
                    paso_actual += 1
                    prog.progress(paso_actual / total_pasos)
                
                # --- INYECCIÓN DE TEXTO MANUAL ---
                if 'df_extra' in st.session_state and not st.session_state.df_extra.empty:
                    all_dfs.append(st.session_state.df_extra)
                    txt.text("Inyectando reportes manuales...")
                
                txt.empty()
                prog.empty()
                # ===== FIN DEL CÓDIGO DE EXTRACCIÓN COPIADO =====
                
                # ===== CONSOLIDACIÓN CON AUDITORÍA =====
                if all_dfs:
                    f_df = pd.concat(all_dfs, ignore_index=True)

                    # Eliminar duplicados
                    f_df = f_df.drop_duplicates(subset=['Link'], keep='first')

                    # Añadir auditoría para FMI
                    def obtener_fuente(row):
                        if row['Organismo'] == 'FMI':
                            return identificar_fuente_fmi(row)
                        else:
                            return {"nombre": "No aplica", "origen": "No aplica", "url": "", "enlace_id": "N/A"}

                    fuentes = f_df.apply(obtener_fuente, axis=1)
                    f_df['Nombre Fuente'] = fuentes.apply(lambda x: x['nombre'])
                    f_df['Origen'] = fuentes.apply(lambda x: x['origen'])
                    f_df['URL Fuente'] = fuentes.apply(lambda x: x['url'])
                    f_df['Enlace ID'] = fuentes.apply(lambda x: x.get('enlace_id', 'N/A'))

                    # ===== NUEVA NUMERACIÓN FIJA PARA PUBLICACIONES INSTITUCIONALES =====
                    print("🔍 CATEGORÍAS - FMI Pub. Inst. - Reconstruyendo enlaces fijos...")

                    # 1. Crear un DataFrame con los 8 enlaces fijos de FMI
                    enlaces_fijos_list = []
                    for enlace in ENLACES_FMI_PUB_INST:
                        enlaces_fijos_list.append({
                            'Categoría': 'Publicaciones Institucionales',
                            'Organismo': 'FMI',
                            'ID Enlace': f"Enlace {enlace['id']} de 8 (FMI - Pub. Inst.)",
                            'Nombre Fuente': enlace['nombre'],
                            'Origen': enlace['origen'],
                            'URL Fuente': enlace['url'],
                            'enlace_id': str(enlace['id']),
                            'Title': '',  # Sin documentos por defecto
                            'Link': '',   # Sin documentos por defecto
                            'Date': None,  # Sin documentos por defecto
                            'Estado': '⚠️ Sin publicaciones en el mes'  # Estado por defecto
                        })

                    df_enlaces_fijos = pd.DataFrame(enlaces_fijos_list)

                    # 2. Identificar qué enlaces tienen documentos reales
                    mask_fmi_pub = (f_df['Categoría'] == 'Publicaciones Institucionales') & (f_df['Organismo'] == 'FMI')
                    df_fmi_pub = f_df[mask_fmi_pub].copy()

                    print(f"   📊 Documentos FMI Pub. Inst. encontrados: {len(df_fmi_pub)}")

                    # Agrupar documentos reales por Nombre Fuente
                    docs_por_fuente = {}
                    if not df_fmi_pub.empty:
                        for nombre_fuente in df_fmi_pub['Nombre Fuente'].unique():
                            docs = df_fmi_pub[df_fmi_pub['Nombre Fuente'] == nombre_fuente]
                            docs_por_fuente[nombre_fuente] = docs

                    # 3. Rellenar los enlaces fijos con los documentos reales (si existen)
                    for idx, row in df_enlaces_fijos.iterrows():
                        nombre_fuente = row['Nombre Fuente']
                        if nombre_fuente in docs_por_fuente:
                            docs = docs_por_fuente[nombre_fuente]
                            # Tomar el primer documento como ejemplo para el título
                            primer_doc = docs.iloc[0]
                            df_enlaces_fijos.loc[idx, 'Title'] = primer_doc['Title']
                            df_enlaces_fijos.loc[idx, 'Link'] = primer_doc['Link']
                            df_enlaces_fijos.loc[idx, 'Date'] = primer_doc['Date']
                            df_enlaces_fijos.loc[idx, 'Estado'] = f'✅ {len(docs)} documentos disponibles'
                            print(f"   ✅ Enlace {row['ID Enlace']} - {nombre_fuente}: {len(docs)} documentos")

                    print(f"   📊 Enlaces fijos reconstruidos: {len(df_enlaces_fijos)}")

                    # 4. Separar los datos que NO son FMI en Publicaciones Institucionales
                    f_df_resto = f_df[~mask_fmi_pub].copy()

                    # 5. Para el resto, mantener la numeración dinámica
                    if not f_df_resto.empty:
                        grupos_unicos = f_df_resto[['Categoría', 'Organismo', 'Nombre Fuente']].drop_duplicates()
                        grupos_unicos = grupos_unicos.sort_values(['Categoría', 'Organismo', 'Nombre Fuente'])
                        grupos_unicos['Grupo Fuente'] = grupos_unicos.groupby(['Categoría', 'Organismo']).cumcount() + 1
                        grupos_unicos['Total Fuentes'] = grupos_unicos.groupby(['Categoría', 'Organismo'])['Grupo Fuente'].transform('max')
                        f_df_resto = f_df_resto.merge(
                            grupos_unicos[['Categoría', 'Organismo', 'Nombre Fuente', 'Grupo Fuente', 'Total Fuentes']],
                            on=['Categoría', 'Organismo', 'Nombre Fuente'],
                            how='left'
                        )
                        f_df_resto['ID Enlace'] = f_df_resto.apply(
                            lambda row: f"Enlace {row['Grupo Fuente']} de {row['Total Fuentes']} ({row['Organismo']} - {row['Categoría']})",
                            axis=1
                        )

                    # 6. Combinar todo
                    df_enlaces_fijos_clean = df_enlaces_fijos.drop(columns=['Estado', 'documentos'], errors='ignore')

                    # Unir
                    f_df_final = pd.concat([df_enlaces_fijos_clean, f_df_resto], ignore_index=True)

                    # 7. Reemplazar f_df con f_df_final para que el resto del código funcione
                    f_df = f_df_final
                    # ===== FIN NUEVA NUMERACIÓN =====

                    # Preparar DataFrame para Word
                    df_word = f_df[['Categoría', 'Organismo', 'ID Enlace', 'Nombre Fuente', 'Origen', 'Enlace ID', 'Date', 'Title', 'Link']].copy()
                    df_word = df_word.rename(columns={
                        'Date': 'Fecha',
                        'Title': 'Nombre de Documento',
                        'Link': 'Enlace',
                        'Enlace ID': 'ID del Enlace Original'
                    })

                    # Guardar el DataFrame en session_state para depuración
                    st.session_state['df_word_auditado'] = df_word
                    st.session_state['f_df_final'] = f_df
                    print(f"📋 Columnas disponibles en df_word: {df_word.columns.tolist()}")
                    print(f"📊 Total de documentos en df_word: {len(df_word)}")

                    # Generar Word con auditoría
                    word = generate_word_auditado(
                        df_word, 
                        title=f"Boletín con Auditoría",
                        subtitle=f"{', '.join(m_sel)} {', '.join(a_sel)} - Fuentes originales identificadas"
                    )

                    st.success(f"✅ Boletín con auditoría generado con {len(df_word)} documentos.")
                    st.download_button(
                        "📄 Descargar Boletín con Auditoría",
                        word,
                        f"Boletin_Auditoria_{'_'.join(m_sel)}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                    # ========== VISTA PREVIA ==========
                    st.markdown("### 📋 Vista previa de datos")
                    st.dataframe(df_word)
                    # ============================================================
                    # ====== CÓDIGO DE DEPURACIÓN ======
                    # ============================================================
                    with st.expander("🔍 Ver datos de depuración (FMI)"):
                                        if 'f_df_final' in st.session_state:
                                            st.write("**DataFrame final con todos los documentos:**")
                                            st.dataframe(st.session_state['f_df_final'])
                                            
                                            st.write("**Documentos de FMI (Publicaciones Institucionales):**")
                                            fmi_pub = st.session_state['f_df_final'][
                                                (st.session_state['f_df_final']['Categoría'] == 'Publicaciones Institucionales') & 
                                                (st.session_state['f_df_final']['Organismo'] == 'FMI')
                                            ]
                                            st.dataframe(fmi_pub)
                                            
                                            # Información adicional útil
                                            st.write("**Conteo de documentos por fuente FMI:**")
                                            if not fmi_pub.empty:
                                                conteo_fuentes = fmi_pub['Nombre Fuente'].value_counts()
                                                st.dataframe(conteo_fuentes)
                                            else:
                                                st.info("No hay documentos de FMI en Publicaciones Institucionales")
                                    # ============================================================
                                            # ============================================================
                                            # ====== NUEVO: ENLACES FIJOS DE FMI (AQUÍ VA) ======
                                            # ============================================================
                                            st.write("**🔗 Enlaces fijos de FMI (8 enlaces):**")
                                            enlaces_fijos = st.session_state['f_df_final'][
                                                st.session_state['f_df_final']['ID Enlace'].str.contains('Enlace \d de 8', na=False)
                                            ]
                                            
                                            # Mostrar solo las columnas relevantes
                                            if not enlaces_fijos.empty:
                                                # Seleccionar columnas para mostrar
                                                columnas_mostrar = ['ID Enlace', 'Nombre Fuente', 'Estado', 'Nombre de Documento']
                                                # Verificar que las columnas existan
                                                columnas_existentes = [col for col in columnas_mostrar if col in enlaces_fijos.columns]
                                                
                                                # Si 'Estado' no existe, crearla con valor por defecto
                                                if 'Estado' not in enlaces_fijos.columns:
                                                    enlaces_fijos['Estado'] = '⚠️ Sin publicaciones en el mes'
                                                    columnas_existentes = ['ID Enlace', 'Nombre Fuente', 'Estado', 'Nombre de Documento']
                                                
                                                # Si 'Nombre de Documento' no existe, usar 'Title'
                                                if 'Nombre de Documento' not in enlaces_fijos.columns and 'Title' in enlaces_fijos.columns:
                                                    enlaces_fijos['Nombre de Documento'] = enlaces_fijos['Title']
                                                    if 'Nombre de Documento' not in columnas_existentes:
                                                        columnas_existentes.append('Nombre de Documento')
                                                
                                                st.dataframe(enlaces_fijos[columnas_existentes])
                                                
                                                # Resumen de estado de los enlaces
                                                st.write("**Resumen de estado:**")
                                                if 'Estado' in enlaces_fijos.columns:
                                                    resumen_estado = enlaces_fijos['Estado'].value_counts()
                                                    st.dataframe(resumen_estado)
                                            else:
                                                st.info("No se encontraron enlaces fijos en el DataFrame")
                                    # ============================================================

                else:
                    st.warning("No se encontraron documentos para los criterios seleccionados.")

    # ==========================================
    # BOTÓN 3: EXPORTAR A EXCEL
    # ==========================================
    if st.button("📊 Exportar a Excel (con auditoría)", type="secondary"):
        if not m_sel or not a_sel:
            st.warning("⚠️ Selecciona mes y año.")
        else:
            with st.spinner("Generando Excel con auditoría..."):
                # ===== EXTRAER DATOS (mismo código que los otros botones) =====
                m_num = [meses_dict[m] for m in m_sel]
                a_num = [int(a) for a in a_sel]
                sd = f"01.{min(m_num):02d}.{min(a_num)}"
                ed = f"{calendar.monthrange(max(a_num), max(m_num))[1]:02d}.{max(m_num):02d}.{max(a_num)}"

                all_dfs = []
                prog = st.progress(0)
                txt = st.empty()
                
                # ... (todo el código de extracción de los otros botones) ...
                # (copia el mismo código que usas para el botón de auditoría)
                
                # ===== CONSOLIDACIÓN CON AUDITORÍA =====
                if all_dfs:
                    f_df = pd.concat(all_dfs, ignore_index=True)
                    f_df = f_df.drop_duplicates(subset=['Link'], keep='first')
                    
                    # Añadir auditoría
                    def obtener_fuente(row):
                        if row['Organismo'] == 'FMI':
                            return identificar_fuente_fmi(row)
                        else:
                            return {"nombre": "No aplica", "origen": "No aplica", "url": "", "enlace_id": "N/A"}
                    
                    fuentes = f_df.apply(obtener_fuente, axis=1)
                    f_df['Nombre Fuente'] = fuentes.apply(lambda x: x['nombre'])
                    f_df['Origen'] = fuentes.apply(lambda x: x['origen'])
                    f_df['URL Fuente'] = fuentes.apply(lambda x: x['url'])
                    f_df['Enlace ID'] = fuentes.apply(lambda x: x.get('enlace_id', 'N/A'))
                    
                    # Preparar DataFrame para Excel
                    df_excel = f_df[['Date', 'Categoría', 'Organismo', 'Title', 'Link', 'Nombre Fuente', 'Origen', 'Enlace ID']].copy()
                    df_excel = df_excel.rename(columns={
                        'Date': 'Fecha',
                        'Title': 'Nombre de Documento',
                        'Link': 'Enlace',
                        'Nombre Fuente': 'Fuente de Consulta',
                        'Origen': 'Tipo de Fuente',
                        'Enlace ID': 'ID del Enlace Original'
                    })
                    
                    # Ordenar por fecha (más reciente primero)
                    df_excel = df_excel.sort_values('Fecha', ascending=False)
                    
                    # ===== CREAR ARCHIVO EXCEL =====
                    output = BytesIO()
                    with pd.ExcelWriter(output, engine='openpyxl') as writer:
                        df_excel.to_excel(writer, sheet_name='Boletín con Auditoría', index=False)
                        
                        # Ajustar ancho de columnas
                        worksheet = writer.sheets['Boletín con Auditoría']
                        for column in worksheet.columns:
                            max_length = 0
                            column_letter = column[0].column_letter
                            for cell in column:
                                try:
                                    if len(str(cell.value)) > max_length:
                                        max_length = len(str(cell.value))
                                except:
                                    pass
                            adjusted_width = min(max_length + 2, 50)
                            worksheet.column_dimensions[column_letter].width = adjusted_width
                    
                    output.seek(0)
                    
                    st.success(f"✅ Excel generado con {len(df_excel)} documentos.")
                    st.download_button(
                        "📊 Descargar Excel",
                        output,
                        f"Boletin_Auditoria_{'_'.join(m_sel)}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                    
                    # Mostrar vista previa
                    st.dataframe(df_excel)
                else:
                    st.warning("No se encontraron documentos para los criterios seleccionados.")

elif modo_app == "Categorías":
    st.title("Documentos de Organismos Internacionales")
    tipo_doc = st.sidebar.selectbox("Tipo de Documento", [
                                    "Discursos", "Reportes", "Investigación", "Publicaciones Institucionales"])

    # Construcción segura de las listas de interfaz
    if tipo_doc == "Discursos":
        orgs_list = ["Todos"] + sorted(orgs_discursos)
    elif tipo_doc == "Reportes":
        orgs_list = ["Todos"] + sorted(orgs_reportes)
    elif tipo_doc == "Investigación":
        orgs_list = ["Todos"] + sorted(orgs_investigacion)
    elif tipo_doc == "Publicaciones Institucionales":
        orgs_list = ["Todos"] + sorted(orgs_pub_inst)
    else:
        orgs_list = ["Todos"] + sorted(
            list(set(orgs_discursos + orgs_reportes + orgs_investigacion + orgs_pub_inst)))

    organismo_seleccionado = st.sidebar.selectbox("Organismo", orgs_list)

    c1, c2 = st.columns(2)
    m_sel = c1.multiselect("Mes(es)", options=list(meses_dict.keys()))
    a_sel = c2.multiselect("Año(s)", options=anios_str, default=["2026"])

    if st.button("🔍 Buscar", type="primary"):
        if not m_sel or not a_sel:
            st.warning("⚠️ Selecciona mes y año.")
        else:
            m_num = [meses_dict[m] for m in m_sel]
            a_num = [int(a) for a in a_sel]
            sd = f"01.{min(m_num):02d}.{min(a_num)}"
            ed = f"{calendar.monthrange(max(a_num), max(m_num))[1]:02d}.{max(m_num):02d}.{max(a_num)}"

            target_orgs = orgs_list[1:] if organismo_seleccionado == "Todos" else [
                organismo_seleccionado]
            dfs_comb = []
            progreso = st.progress(0)
            txt = st.empty()

            for i, o in enumerate(target_orgs):
                txt.text(f"Extrayendo: {o}...")
                df = pd.DataFrame()
                try:
                    # --- LÓGICA DE EXTRACCIÓN POR TIPO ---
                    if tipo_doc == "Discursos":
                        if o == "BPI":
                            df = load_data_bis(
                                use_event_date=use_event_date_bis,
                                target_year=a_num[0] if a_num else None,
                                target_month=m_num[0] if m_num else None
                            )
                        elif o == "ECB (Europa)":
                            df = load_data_ecb(sd, ed)
                        elif o == "BBk (Alemania)":
                            df = load_data_bbk(sd, ed)
                        elif o == "Fed (Estados Unidos)":
                            df = load_data_fed(a_num)
                        elif o == "BdF (Francia)":
                            df = load_data_bdf(sd, ed)
                        elif o == "BM":
                            df = load_data_bm(sd, ed)
                        elif o == "BoC (Canadá)":
                            df = load_data_boc(sd, ed)
                        elif o == "BoJ (Japón)":
                            df = load_data_boj(sd, ed)
                        elif o == "BoE (Inglaterra)": df = load_discursos_boe(sd, ed)
                        elif o == "CEMLA": 
                            print("🔴🔴🔴 LLAMANDO A CEMLA INVESTIGACIÓN 🔴🔴🔴")
                            df = load_investigacion_cemla(sd, ed)
                            print(f"🔴🔴🔴 RESULTADO CEMLA: {len(df)} documentos 🔴🔴🔴")
                        elif o == "CEF":
                            df = load_data_cef(sd, ed)
                        elif o == "FMI":
                            df = load_discursos_fmi(sd, ed)
                        elif o == "PBoC (China)":
                            df = load_data_pboc(sd, ed)
                        elif o == "BdE (España)":
                            df = load_data_bde(sd, ed)

                    elif tipo_doc == "Reportes":
                        if o == "BID":
                            dfs_bid = []
                            try:
                                dfs_bid.append(load_reportes_bid(sd, ed))
                            except:
                                pass
                            try:
                                dfs_bid.append(load_reportes_bid_en(sd, ed))
                            except:
                                pass
                            dfs_bid = [d for d in dfs_bid if not d.empty]
                            if dfs_bid:
                                df = pd.concat(dfs_bid, ignore_index=True).drop_duplicates(
                                    subset=['Link'])
                        elif o == "BM":
                            df = load_reportes_bm(sd, ed)
                        elif o == "BPI":
                            df = load_reportes_bpi(sd, ed)
                        elif o == "CEF":
                            df = load_reportes_cef(sd, ed)
                        elif o == "OCDE":
                            df = load_reportes_ocde(sd, ed)
                        elif o == "FEM": df = load_reportes_fem(sd, ed)

                    elif tipo_doc == "Investigación":
                        if o == "BID":
                            df = load_investigacion_bid_unified(sd, ed)
                        elif o == "BPI":
                            df = load_investigacion_bpi(sd, ed)
                        elif o == "BM":
                            df = load_investigacion_bm(sd, ed)
                        elif o == "CEMLA":   # <-- ESTA LÍNEA DEBE EXISTIR
                            print("🔴 LLAMANDO A CEMLA")
                            df = load_investigacion_cemla(sd, ed)
                        elif o == "BPI":
                            df = load_investigacion_bpi(sd, ed)
                        elif o == "BM":
                            df = load_investigacion_bm(sd, ed)
                            
                        elif o == "FMI": 
                            df_blogs, df_wp = pd.DataFrame(), pd.DataFrame()
                            try: df_blogs = load_investigacion_fmi(sd, ed)
                            except: pass
                            try: df_wp = load_working_papers_fmi(sd, ed)
                            except: pass
                            
                            dfs_fmi = [d for d in [df_blogs, df_wp] if not d.empty]
                            if dfs_fmi:
                                df = pd.concat(dfs_fmi, ignore_index=True).drop_duplicates(subset=['Link']).sort_values("Date", ascending=False)
                        elif o == "OCDE":
                            df = load_investigacion_ocde(sd, ed)

                    elif tipo_doc == "Publicaciones Institucionales":
                        if o == "BPI":
                            df = load_pub_inst_bpi(sd, ed)
                        elif o == "CEF":
                            df = load_pub_inst_cef(sd, ed)
                        elif o == "OEI": 
                            df = load_pub_inst_oei(sd, ed)
                        elif o == "OCDE":
                            df = load_pub_inst_ocde(sd, ed)
                        elif o == "BM":
                            df = load_pub_inst_bm(sd, ed)
                        elif o == "CEMLA":
                            df = load_pub_inst_cemla(sd, ed)
                        elif o == "FMI":
                            print(f"\n{'='*50}")
                            print(f"🔍 CATEGORÍAS - Procesando FMI para {m_sel} {a_sel}")
                            print(f"   Fechas: {sd} a {ed}")
                            print(f"{'='*50}")
                            
                            # 1. F&D Magazine (agregar explícitamente)
                            df_fandd = load_pub_inst_fandd(sd, ed)
                            print(f"   📊 F&D Magazine: {len(df_fandd)} documentos")

                            # 2. SSG - JSON Estático (WEO, Fiscal Monitor)
                            df_flagships = load_pub_inst_fmi(sd, ed)
                            print(f"   📊 Flagships: {len(df_flagships)} documentos")
                            
                            # 3. SSG - JSON Estático (Comunicados)
                            df_prs = load_press_releases_fmi(sd, ed)
                            print(f"   📊 Press Releases: {len(df_prs)} documentos")
                            
                            # 4. CSR API - Coveo (Country Reports)
                            df_crs = load_country_reports_fmi(sd, ed)
                            print(f"   📊 Country Reports: {len(df_crs)} documentos")
                            
                            # Unir todos
                            print(f"🔍 CATEGORÍAS - Flagships: {len(df_flagships)}, PRs: {len(df_prs)}, CRs: {len(df_crs)}")
                            dfs_a_unir = [d for d in [df_fandd, df_flagships, df_prs, df_crs] if not d.empty]
                            if dfs_a_unir:
                                df = pd.concat(dfs_a_unir, ignore_index=True)
                                df = df.sort_values("Date", ascending=False)
                                print(f"   📊 TOTAL combinado FMI: {len(df)} documentos")
                            else:
                                print(f"   ⚠️ Ninguna fuente retornó datos")
                except Exception as e:
                    pass

                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (
                        df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty:
                        df_f['Organismo'] = o
                        dfs_comb.append(df_f)
                progreso.progress((i+1)/len(target_orgs))

            txt.empty()
            progreso.empty()

            if dfs_comb:
                f_df = pd.concat(dfs_comb, ignore_index=True)
                f_df['Categoría'] = tipo_doc

                # ========== ELIMINACIÓN DE DUPLICADOS ==========
                print(f"📊 Total antes de desduplicar: {len(f_df)}")

                # 1. Eliminar duplicados exactos por Link
                f_df = f_df.drop_duplicates(subset=['Link'], keep='first')
                print(f"   Después de eliminar duplicados por Link: {len(f_df)}")

                # 2. Normalizar títulos para comparación
                f_df['Title_Norm'] = f_df['Title'].str.lower().str.replace(r'[^\w\s]', '', regex=True).str.replace(r'\s+', ' ', regex=True).str.strip()

                # 3. Eliminar duplicados por título normalizado (mismo título en diferentes categorías)
                f_df = f_df.sort_values('Date', ascending=False).drop_duplicates(subset=['Title_Norm'], keep='first')
                print(f"   Después de eliminar duplicados por título: {len(f_df)}")

                # 4. Eliminar columna temporal
                f_df = f_df.drop(columns=['Title_Norm'])

                print(f"📊 Total después de desduplicación: {len(f_df)}")

                # --- PREPARACIÓN PARA EL WORD (Orden Institucional) ---
                if tipo_doc == "Discursos":
                    f_df_word = f_df.sort_values(
                        by=["Title"], ascending=[True])
                else:
                    f_df_word = f_df.sort_values(
                        by=["Organismo", "Title"], ascending=[True, True])

                f_df_word = f_df_word[['Categoría',
                                       'Organismo', 'Title', 'Link']]
                f_df_word = f_df_word.rename(
                    columns={"Categoría": "Tipo de Documento", "Title": "Nombre de Documento"})

                st.success(f"Se encontraron **{len(f_df)}** documentos.")
                word_file = generate_word(
                    f_df_word, title=f"Explorador - {tipo_doc}")
                st.download_button(
                    "📄 Descargar en Word", data=word_file, file_name=f"Explorador_{tipo_doc}.docx")

                # --- PREPARACIÓN PARA LA VISTA PREVIA ---
                disp = f_df.copy()
                disp = disp.sort_values(
                    by="Date", ascending=False)  # Orden cronológico
                disp["Fecha"] = disp["Date"].dt.strftime(
                    '%d/%m/%Y')  # Formatear fecha
                disp["Nombre de Documento"] = disp.apply(
                    lambda x: f"[{x['Title']}]({x['Link']})", axis=1)
                disp = disp.rename(columns={"Categoría": "Tipo de Documento"})

                if organismo_seleccionado == "Todos":
                    cols_vis = ["Fecha", "Tipo de Documento",
                                "Organismo", "Nombre de Documento"]
                else:
                    cols_vis = ["Fecha", "Tipo de Documento",
                                "Nombre de Documento"]

                st.markdown(disp[cols_vis].to_markdown(index=False), unsafe_allow_html=True)
            else:
                # Verificar si CEMLA estaba en la búsqueda y no hay resultados
                if "CEMLA" in target_orgs and tipo_doc == "Investigación":
                    st.warning("⚠️ No se encontraron documentos para las fechas seleccionadas.")
                    st.info("📌 **CEMLA Investigación**: ScienceDirect bloquea el acceso automatizado. No se pueden extraer artículos.\n\n➡️ **Solución**: Utiliza la sección **'Carga Manual'** en el menú principal para agregar estos documentos al boletín mensual.")
                else:
                    st.warning(
                    "No se encontraron documentos para las fechas seleccionadas.")

elif modo_app == "Carga Manual":
    st.title("🛠️ Centro de Carga Manual")
    st.markdown("Pega el texto de las páginas que fallan. Previsualiza, valida y une todo en un solo documento.")
    
    if 'cargas_validadas' not in st.session_state:
        st.session_state.cargas_validadas = {
            "OCDE (Reportes)": pd.DataFrame(),
            "OCDE (Pub. Institucionales)": pd.DataFrame(),
            "OCDE (Investigación)": pd.DataFrame()
        }

    st.subheader("Estado de Carga")
    cols_estado = st.columns(3)
    claves_cajas = list(st.session_state.cargas_validadas.keys())
    
    for i, clave in enumerate(claves_cajas):
        estado = "✅ Listo" if not st.session_state.cargas_validadas[clave].empty else "❌ Pendiente"
        cols_estado[i].info(f"**{clave}**\n\n{estado}")

    st.markdown("---")
    
    c1, c2 = st.columns(2)
    mes_manual = c1.selectbox("Mes objetivo a filtrar:", [1,2,3,4,5,6,7,8,9,10,11,12], index=datetime.datetime.now().month-1, format_func=lambda x: calendar.month_name[x].capitalize())
    año_manual = c2.number_input("Año objetivo a filtrar:", min_value=2020, max_value=2030, value=datetime.datetime.now().year)

    st.markdown("---")
    st.subheader("Cajas de Extracción")

    def crear_caja_manual(titulo_caja, categoria_doc, organismo_nombre, url_fuente=None):
        with st.expander(f"📥 Cargar: {titulo_caja}", expanded=True):
            
            if url_fuente:
                st.markdown(f"👉 **[Haz clic aquí para abrir la página oficial de {titulo_caja}]({url_fuente})**")
                
            texto = st.text_area(f"Copia el texto de la página y pégalo aquí (Ctrl+A, Ctrl+C, Ctrl+V):", height=150, key=f"txt_{titulo_caja}")
            
            col_btn1, col_btn2 = st.columns([1, 1])
            
            if col_btn1.button(f"🔍 Previsualizar {titulo_caja}", key=f"btn_prev_{titulo_caja}"):
                if texto:
                    with st.spinner("Procesando y buscando links..."):
                        df_bruto = procesar_texto_pegado(texto, organismo_nombre)
                            
                        if not df_bruto.empty:
                            df_filtrado = df_bruto[
                                (df_bruto['Date'].dt.month == mes_manual) & 
                                (df_bruto['Date'].dt.year == año_manual)
                            ].copy()
                            
                            if not df_filtrado.empty:
                                for idx in df_filtrado.index:
                                    t = df_filtrado.loc[idx, "Title"]
                                    df_filtrado.loc[idx, "Link"] = buscar_link_inteligente(t, organismo_nombre)
                                
                                df_filtrado['Categoría'] = categoria_doc
                                st.session_state[f"temp_{titulo_caja}"] = df_filtrado
                                
                                st.success(f"Se encontraron {len(df_filtrado)} documentos de {mes_manual} {año_manual}.")
                                st.dataframe(df_filtrado, width="stretch")
                            else:
                                st.warning("No hay coincidencias con el mes y año seleccionados.")
                else:
                    st.error("Pega el texto primero.")
            
            if col_btn2.button(f"➕ Agregar a Descarga Final", type="primary", key=f"btn_add_{titulo_caja}"):
                if f"temp_{titulo_caja}" in st.session_state and not st.session_state[f"temp_{titulo_caja}"].empty:
                    st.session_state.cargas_validadas[titulo_caja] = st.session_state[f"temp_{titulo_caja}"]
                    st.success(f"¡{titulo_caja} guardado en memoria! ✅")
                    time.sleep(1)
                    st.rerun() 
                else:
                    st.error("Primero debes Previsualizar y obtener resultados.")

    link_ocde_rep = "https://www.oecd.org/en/search/publications.html?orderBy=mostRecent&page=0&facetTags=oecd-content-types%3Apublications%2Freports%2Coecd-languages%3Aen&minPublicationYear=2026&maxPublicationYear=2026"
    link_ocde_pub = "https://www.oecd.org/en/search.html?orderBy=mostRecent&page=0&facetTags=oecd-policy-subissues%3Apsi114%2Coecd-languages%3Aen"
    link_ocde_inv = "https://www.oecd.org/en/publications/reports.html?orderBy=mostRecent&page=0&facetTags=oecd-content-types%3Apublications%2Fworking-papers%2Coecd-languages%3Aen"
    
    crear_caja_manual("OCDE (Reportes)", "Reportes", "OCDE", link_ocde_rep)
    crear_caja_manual("OCDE (Pub. Institucionales)", "Publicaciones Institucionales", "OCDE", link_ocde_pub)
    crear_caja_manual("OCDE (Investigación)", "Investigación", "OCDE", link_ocde_inv)

    st.markdown("---")
    st.subheader("Exportación Final")
    
    tablas_listas = [df for df in st.session_state.cargas_validadas.values() if not df.empty]
    
    if tablas_listas:
        df_maestro = pd.concat(tablas_listas, ignore_index=True)
        num_cat = len(tablas_listas)
        st.info(f"Tienes **{num_cat}/3** categorías listas, sumando un total de **{len(df_maestro)}** documentos para exportar.")
        
        df_word_manual = df_maestro[['Categoría', 'Organismo', 'Title', 'Link']].copy()
        df_word_manual = df_word_manual.rename(columns={"Categoría": "Tipo de Documento", "Title": "Nombre de Documento"})
        
        word_manual = generate_word(df_word_manual, title="Boletín - Carga Manual", subtitle=f"Mes: {mes_manual} | Año: {año_manual}")
        
        c_down, c_clear = st.columns(2)
        with c_down:
            st.download_button(
                label=f"📄 Descargar Word ({num_cat}/3 Listas)", 
                data=word_manual, 
                file_name=f"Carga_Manual_{mes_manual}_{año_manual}.docx"
            )
        with c_clear:
            if st.button("🗑️ Reiniciar todo el módulo"):
                for clave in st.session_state.cargas_validadas.keys():
                    st.session_state.cargas_validadas[clave] = pd.DataFrame()
                st.rerun()
    else:
        st.warning("Aún no has agregado ninguna carga a la descarga final. Agrega al menos 1 para habilitar el botón de descarga.")

